"""
MCS Platform - AI Service Layer (Sprint 1 + CPM Review Enhancements)
Consolidated AI service using Anthropic Claude models with tiered routing:
  - Haiku:  Fast classification, simple lookups, flag categorisation
  - Sonnet: Contextual risk analysis, flag explanation, prioritisation
  - Opus:   Complex advisory, risk summary reports, multi-flag synthesis

Features:
  - Tier 3: Contextual Risk Analysis per flag
  - AI Risk Summary Report generation (narrative Word document)
  - AI Smart Flag Prioritisation (ATO Interest scoring)
  - Batch Risk Engine support
  - Materiality scaling based on entity revenue
  - Response caching via data hash to avoid redundant API calls
  - AI Feedback Loop — stores user corrections for future prompt improvement
"""
import hashlib
import json
import io
import logging
from decimal import Decimal

from django.conf import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Model Tier Configuration
# ---------------------------------------------------------------------------
MODEL_TIERS = {
    "haiku": "gpt-4.1-nano",       # Fast, cheap — classification tasks
    "sonnet": "gpt-4.1-mini",      # Balanced — analysis and explanation
    "opus": "gpt-4.1-mini",        # Complex — reports and synthesis (upgrade when available)
}

# When Anthropic API is configured directly, use these:
ANTHROPIC_MODEL_TIERS = {
    "haiku": "claude-3-5-haiku-20241022",
    "sonnet": "claude-3-5-sonnet-20241022",
    "opus": "claude-3-opus-20240229",
}


def _get_model(tier):
    """
    Get the appropriate model for the given tier.
    Uses Anthropic models if ANTHROPIC_API_KEY is set, otherwise falls back
    to OpenAI-compatible API with mapped models.
    """
    anthropic_key = getattr(settings, "ANTHROPIC_API_KEY", None)
    if not anthropic_key:
        import os
        anthropic_key = os.environ.get("ANTHROPIC_API_KEY")

    if anthropic_key:
        return ANTHROPIC_MODEL_TIERS.get(tier, ANTHROPIC_MODEL_TIERS["sonnet"])
    return MODEL_TIERS.get(tier, MODEL_TIERS["sonnet"])


def _use_anthropic():
    """Check if we should use Anthropic API directly."""
    anthropic_key = getattr(settings, "ANTHROPIC_API_KEY", None)
    if not anthropic_key:
        import os
        anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
    return bool(anthropic_key)


# ---------------------------------------------------------------------------
# LLM Client — supports both OpenAI-compatible and Anthropic native
# ---------------------------------------------------------------------------
def _call_llm(system_prompt, user_prompt, tier="sonnet", temperature=0.3, max_tokens=2000):
    """
    Call the LLM with system and user prompts using the appropriate tier.
    Supports both Anthropic native API and OpenAI-compatible fallback.
    Returns the response text.
    """
    model = _get_model(tier)

    if _use_anthropic():
        return _call_anthropic(system_prompt, user_prompt, model, temperature, max_tokens)
    else:
        return _call_openai_compat(system_prompt, user_prompt, model, temperature, max_tokens)


def _call_anthropic(system_prompt, user_prompt, model, temperature, max_tokens):
    """Call Anthropic API directly."""
    try:
        import anthropic
    except ImportError:
        raise ImportError("anthropic package not installed. Run: pip install anthropic")

    import os
    client = anthropic.Anthropic(
        api_key=os.environ.get("ANTHROPIC_API_KEY", getattr(settings, "ANTHROPIC_API_KEY", ""))
    )

    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        system=system_prompt,
        messages=[
            {"role": "user", "content": user_prompt},
        ],
    )
    return response.content[0].text.strip()


def _call_openai_compat(system_prompt, user_prompt, model, temperature, max_tokens):
    """Call OpenAI-compatible API (default fallback)."""
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai package not installed. Run: pip install openai")

    client = OpenAI()  # Uses OPENAI_API_KEY and base_url from env
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return response.choices[0].message.content.strip()


# ---------------------------------------------------------------------------
# Data Hash for Cache Invalidation
# ---------------------------------------------------------------------------
def _compute_flag_hash(flag):
    """Compute MD5 hash of flag data for cache invalidation."""
    data = {
        "rule_id": flag.rule_id,
        "severity": flag.severity,
        "description": flag.description,
        "calculated_values": _serialise_values(flag.calculated_values),
        "affected_accounts": flag.affected_accounts,
    }
    return hashlib.md5(json.dumps(data, sort_keys=True).encode()).hexdigest()


def _serialise_values(obj):
    """Convert Decimal values to float for JSON serialisation."""
    if isinstance(obj, dict):
        return {k: _serialise_values(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_serialise_values(v) for v in obj]
    elif isinstance(obj, Decimal):
        return float(obj)
    return obj


# ---------------------------------------------------------------------------
# Materiality Scaling
# ---------------------------------------------------------------------------
def _calculate_materiality(financial_year):
    """
    Calculate materiality thresholds based on entity revenue.
    Returns a dict with overall_materiality, performance_materiality,
    and trivial_threshold.

    Uses standard audit materiality benchmarks:
    - 1-2% of revenue for commercial entities
    - 0.5-1% for SMSFs (regulated entities)
    - 5% of gross profit for smaller entities
    """
    from core.models import TrialBalanceLine

    lines = financial_year.trial_balance_lines.all()
    entity_type = financial_year.entity.entity_type

    total_revenue = Decimal("0")
    total_assets = Decimal("0")
    for line in lines:
        balance = abs(line.adjusted_balance or line.original_balance or Decimal("0"))
        section = (line.section or "").lower()
        if "revenue" in section or "income" in section:
            total_revenue += balance
        elif "asset" in section:
            total_assets += balance

    # Determine base for materiality
    if total_revenue > 0:
        base = total_revenue
    elif total_assets > 0:
        base = total_assets
    else:
        base = Decimal("100000")  # Default minimum

    # Entity-type specific percentages
    if entity_type == "smsf":
        pct = Decimal("0.005")  # 0.5% for SMSFs
    elif entity_type in ("trust", "partnership"):
        pct = Decimal("0.01")   # 1% for trusts/partnerships
    elif total_revenue < Decimal("500000"):
        pct = Decimal("0.02")   # 2% for small entities
    else:
        pct = Decimal("0.015")  # 1.5% for larger entities

    overall = (base * pct).quantize(Decimal("1"))
    performance = (overall * Decimal("0.75")).quantize(Decimal("1"))
    trivial = (overall * Decimal("0.05")).quantize(Decimal("1"))

    # Minimum floors
    overall = max(overall, Decimal("500"))
    performance = max(performance, Decimal("375"))
    trivial = max(trivial, Decimal("25"))

    return {
        "overall_materiality": overall,
        "performance_materiality": performance,
        "trivial_threshold": trivial,
        "base_amount": base,
        "base_type": "revenue" if total_revenue > 0 else "total_assets",
        "percentage": float(pct * 100),
    }


# ---------------------------------------------------------------------------
# Context Builder — assembles financial data for AI prompts
# ---------------------------------------------------------------------------
def _build_entity_context(financial_year, include_materiality=True):
    """
    Build a rich context string about the entity and financial year
    for use in AI prompts. Includes materiality thresholds.
    """
    entity = financial_year.entity
    lines = financial_year.trial_balance_lines.all().order_by("section", "account_code")

    # Basic info
    ctx = f"Entity: {entity.entity_name}\n"
    ctx += f"Type: {entity.get_entity_type_display()}\n"
    ctx += f"ABN: {entity.abn or 'Not provided'}\n"
    ctx += f"Financial Year: {financial_year.year_label} ({financial_year.start_date} to {financial_year.end_date})\n"
    ctx += f"Period Type: {financial_year.get_period_type_display()}\n"
    ctx += f"Status: {financial_year.get_status_display()}\n\n"

    # Materiality thresholds
    if include_materiality:
        mat = _calculate_materiality(financial_year)
        ctx += "=== MATERIALITY ===\n"
        ctx += f"Overall Materiality: ${mat['overall_materiality']:,.0f} "
        ctx += f"({mat['percentage']:.1f}% of {mat['base_type']})\n"
        ctx += f"Performance Materiality: ${mat['performance_materiality']:,.0f}\n"
        ctx += f"Trivial Threshold: ${mat['trivial_threshold']:,.0f}\n\n"

    # Trial Balance Summary
    ctx += "=== TRIAL BALANCE ===\n"
    total_revenue = Decimal("0")
    total_expenses = Decimal("0")
    total_assets = Decimal("0")
    total_liabilities = Decimal("0")

    for line in lines:
        balance = line.adjusted_balance or line.original_balance or Decimal("0")
        prior = line.prior_year_balance or Decimal("0")
        variance = balance - prior if prior else None

        ctx += f"  {line.account_code} | {line.account_name} | "
        ctx += f"Current: ${balance:,.2f}"
        if prior:
            ctx += f" | Prior: ${prior:,.2f}"
        if variance is not None and prior:
            pct = (variance / abs(prior) * 100) if prior != 0 else Decimal("0")
            ctx += f" | Var: ${variance:,.2f} ({pct:,.1f}%)"
        ctx += f" | Section: {line.section}\n"

        section = (line.section or "").lower()
        if "revenue" in section or "income" in section:
            total_revenue += balance
        elif "expense" in section or "cost" in section:
            total_expenses += abs(balance)
        elif "asset" in section:
            total_assets += balance
        elif "liabilit" in section:
            total_liabilities += abs(balance)

    ctx += f"\nSummary: Revenue=${total_revenue:,.2f}, Expenses=${total_expenses:,.2f}, "
    ctx += f"Net={total_revenue - total_expenses:,.2f}\n"
    ctx += f"Assets=${total_assets:,.2f}, Liabilities=${total_liabilities:,.2f}\n"

    return ctx


def _build_flag_context(flag):
    """Build context string for a specific risk flag."""
    ctx = f"Risk Flag: {flag.title}\n"
    ctx += f"Rule ID: {flag.rule_id}\n"
    ctx += f"Tier: {flag.tier}\n"
    ctx += f"Severity: {flag.severity}\n"
    ctx += f"Description: {flag.description}\n"
    ctx += f"Recommended Action: {flag.recommended_action}\n"

    if flag.legislation_ref:
        ctx += f"Legislation: {flag.legislation_ref}\n"

    if flag.affected_accounts:
        ctx += f"Affected Accounts: {', '.join(str(a) for a in flag.affected_accounts)}\n"

    if flag.calculated_values:
        ctx += "Calculated Values:\n"
        for k, v in flag.calculated_values.items():
            ctx += f"  {k}: {v}\n"

    return ctx


# ---------------------------------------------------------------------------
# AI Feedback Loop — store corrections for prompt improvement
# ---------------------------------------------------------------------------
def record_feedback(flag, feedback_type, user_notes, user):
    """
    Record user feedback on an AI analysis result.
    feedback_type: 'correct', 'partially_correct', 'incorrect', 'irrelevant'
    This data is stored for future prompt refinement.
    """
    from core.models import AuditLog

    feedback_data = {
        "flag_id": str(flag.pk),
        "rule_id": flag.rule_id,
        "severity": flag.severity,
        "ai_explanation": flag.ai_explanation or "",
        "ai_suggested_action": flag.ai_suggested_action or "",
        "ato_interest_score": flag.ato_interest_score,
        "feedback_type": feedback_type,
        "user_notes": user_notes,
    }

    # Store as audit log entry
    AuditLog.objects.create(
        user=user,
        action="ai_feedback",
        description=f"AI feedback ({feedback_type}) on flag: {flag.title}",
        affected_object_type="RiskFlag",
        affected_object_id=str(flag.pk),
        metadata=feedback_data,
    )

    # Update flag with feedback
    flag.ai_feedback = feedback_type
    flag.ai_feedback_notes = user_notes
    flag.save(update_fields=["ai_feedback", "ai_feedback_notes"])

    logger.info(f"AI feedback recorded: {feedback_type} for flag {flag.pk}")
    return True


def _get_feedback_context():
    """
    Build a context string from recent AI feedback to improve prompts.
    Returns examples of corrections to guide the AI.
    """
    from core.models import AuditLog

    recent_feedback = AuditLog.objects.filter(
        action="ai_feedback",
        metadata__feedback_type__in=["incorrect", "partially_correct"],
    ).order_by("-timestamp")[:10]

    if not recent_feedback.exists():
        return ""

    ctx = "\n=== LEARNING FROM PREVIOUS CORRECTIONS ===\n"
    ctx += "The following are examples of previous analyses that were corrected by accountants. "
    ctx += "Use these to calibrate your responses:\n\n"

    for log in recent_feedback:
        meta = log.metadata or {}
        ctx += f"- Rule: {meta.get('rule_id', 'N/A')}, "
        ctx += f"Feedback: {meta.get('feedback_type', 'N/A')}\n"
        if meta.get("user_notes"):
            ctx += f"  Correction: {meta['user_notes']}\n"

    return ctx


# ---------------------------------------------------------------------------
# Tier 3: AI Contextual Risk Analysis (uses Sonnet)
# ---------------------------------------------------------------------------
ANALYSIS_SYSTEM_PROMPT = """You are an expert Australian tax and accounting advisor working for MC & S, 
a Melbourne-based accounting practice. You analyse risk flags raised by an automated audit engine 
against financial statements prepared under Australian Accounting Standards (AASB).

Your role is to:
1. Explain the risk in plain English that a CPA-qualified accountant would understand
2. Provide specific, actionable recommendations
3. Reference relevant Australian legislation (ITAA 1936/1997, Corporations Act 2001, SIS Act 1993) where applicable
4. Consider the ATO's current compliance focus areas
5. Consider materiality — flag whether the amount is above or below materiality thresholds
6. Be concise but thorough — aim for 3-5 sentences for explanation, 2-4 bullet points for actions

IMPORTANT: You are assisting accountants, not clients. Use professional language.
Do NOT include disclaimers about seeking professional advice — the reader IS the professional."""


def analyse_risk_flag(flag):
    """
    Run AI contextual analysis on a single risk flag using Sonnet tier.
    Returns dict with success, explanation, suggested_action, data_hash.
    """
    data_hash = _compute_flag_hash(flag)

    # Check cache — skip if already analysed with same data
    if flag.ai_explanation and flag.ai_data_hash == data_hash:
        return {
            "success": True,
            "explanation": flag.ai_explanation,
            "suggested_action": flag.ai_suggested_action,
            "data_hash": data_hash,
            "cached": True,
        }

    entity_context = _build_entity_context(flag.financial_year, include_materiality=True)
    flag_context = _build_flag_context(flag)
    feedback_context = _get_feedback_context()

    user_prompt = f"""Analyse the following risk flag for this entity:

{entity_context}

{flag_context}
{feedback_context}

Provide your response in the following JSON format:
{{
    "explanation": "Plain English explanation of what this risk means and why it matters. Include whether the amount is material.",
    "suggested_action": "Specific steps the accountant should take to address this risk",
    "materiality_assessment": "above_materiality" or "below_materiality" or "borderline"
}}

Return ONLY the JSON, no other text."""

    try:
        response_text = _call_llm(
            ANALYSIS_SYSTEM_PROMPT,
            user_prompt,
            tier="sonnet",
            max_tokens=2000,
        )

        # Parse JSON response
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()

        result = json.loads(response_text)

        return {
            "success": True,
            "explanation": result.get("explanation", ""),
            "suggested_action": result.get("suggested_action", ""),
            "materiality_assessment": result.get("materiality_assessment", ""),
            "data_hash": data_hash,
            "cached": False,
        }

    except json.JSONDecodeError as e:
        logger.warning(f"AI response was not valid JSON: {e}")
        return {
            "success": True,
            "explanation": response_text[:1000] if response_text else "Analysis completed but response format was unexpected.",
            "suggested_action": "Please review the explanation above and determine appropriate action.",
            "data_hash": data_hash,
            "cached": False,
        }
    except Exception as e:
        logger.exception(f"AI analysis failed for flag {flag.pk}")
        return {
            "success": False,
            "error": str(e),
        }


# ---------------------------------------------------------------------------
# AI Smart Flag Prioritisation (ATO Interest Scoring) — uses Sonnet
# ---------------------------------------------------------------------------
PRIORITISATION_SYSTEM_PROMPT = """You are an expert on ATO (Australian Taxation Office) compliance 
and audit selection criteria. You score risk flags based on how likely they are to attract ATO attention.

Consider:
- ATO's published compliance focus areas (Division 7A, work-related expenses, rental properties, 
  crypto, SMSF compliance, trust distributions, cash economy businesses)
- The materiality of the amounts involved relative to the entity's materiality thresholds
- Whether the issue represents a pattern vs one-off
- The entity type and its typical ATO audit risk profile
- Current ATO data-matching capabilities
- Whether the amount exceeds the trivial threshold

Score from 1-10 where:
1-3: Low ATO interest — routine, immaterial, or well-documented
4-6: Moderate ATO interest — may trigger data-matching or random review
7-8: High ATO interest — likely to attract attention if ATO reviews this entity
9-10: Very high ATO interest — matches known ATO compliance targets

Be realistic and calibrated. Most flags should score 3-6."""


def prioritise_flags(flags, financial_year):
    """
    Score a batch of flags by ATO interest likelihood using Sonnet tier.
    Processes in batches to manage token limits.
    """
    if not flags:
        return {"success": True, "scored": 0}

    entity_context = _build_entity_context(financial_year, include_materiality=True)
    feedback_context = _get_feedback_context()
    scored = 0

    # Process in batches of 10
    batch_size = 10
    for i in range(0, len(flags), batch_size):
        batch = flags[i:i + batch_size]

        flags_text = ""
        for idx, flag in enumerate(batch):
            flags_text += f"\n--- Flag {idx + 1} (ID: {flag.pk}) ---\n"
            flags_text += _build_flag_context(flag)

        user_prompt = f"""Score the following risk flags for ATO interest likelihood.

{entity_context}

{flags_text}
{feedback_context}

Return a JSON array with one object per flag:
[
    {{
        "flag_id": "<uuid>",
        "score": <1-10>,
        "reasoning": "Brief explanation of why this score"
    }}
]

Return ONLY the JSON array, no other text."""

        try:
            response_text = _call_llm(
                PRIORITISATION_SYSTEM_PROMPT,
                user_prompt,
                tier="sonnet",
                max_tokens=3000,
            )

            # Parse JSON
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()

            results = json.loads(response_text)

            # Map results back to flags
            result_map = {str(r["flag_id"]): r for r in results}

            for flag in batch:
                result = result_map.get(str(flag.pk))
                if result:
                    flag.ato_interest_score = max(1, min(10, int(result["score"])))
                    flag.ato_interest_reasoning = result.get("reasoning", "")
                    flag.save(update_fields=["ato_interest_score", "ato_interest_reasoning"])
                    scored += 1

        except Exception as e:
            logger.exception(f"AI prioritisation batch error: {e}")
            continue

    return {"success": True, "scored": scored}


# ---------------------------------------------------------------------------
# Batch Risk Engine — run AI analysis on all open flags for a financial year
# ---------------------------------------------------------------------------
def batch_analyse_flags(financial_year, force=False):
    """
    Run AI analysis on all open flags for a financial year.
    Uses Sonnet for individual analysis, then Haiku for quick classification.
    If force=True, re-analyses even cached results.
    Returns summary dict.
    """
    flags = list(financial_year.risk_flags.filter(status="open").order_by("-severity"))

    if not flags:
        return {"success": True, "analysed": 0, "skipped": 0, "errors": 0}

    analysed = 0
    skipped = 0
    errors = 0

    for flag in flags:
        data_hash = _compute_flag_hash(flag)

        # Skip if cached and not forced
        if not force and flag.ai_explanation and flag.ai_data_hash == data_hash:
            skipped += 1
            continue

        result = analyse_risk_flag(flag)

        if result.get("success"):
            if not result.get("cached"):
                flag.ai_explanation = result["explanation"]
                flag.ai_suggested_action = result["suggested_action"]
                flag.ai_data_hash = result["data_hash"]
                flag.save(update_fields=[
                    "ai_explanation", "ai_suggested_action", "ai_data_hash"
                ])
                analysed += 1
            else:
                skipped += 1
        else:
            errors += 1

    # After individual analysis, run batch prioritisation
    unscored = [f for f in flags if not f.ato_interest_score or force]
    if unscored:
        prioritise_flags(unscored, financial_year)

    return {
        "success": True,
        "analysed": analysed,
        "skipped": skipped,
        "errors": errors,
        "total": len(flags),
    }


# ---------------------------------------------------------------------------
# Quick Classification — uses Haiku for fast categorisation
# ---------------------------------------------------------------------------
def quick_classify_flag(flag):
    """
    Use Haiku tier for fast classification of a flag into action categories.
    Returns: 'investigate', 'document', 'adjust', 'disclose', 'no_action'
    """
    prompt = f"""Classify this accounting risk flag into ONE action category:

Flag: {flag.title}
Severity: {flag.severity}
Description: {flag.description}

Categories:
- investigate: Needs further investigation or information gathering
- document: Needs documentation or working paper support
- adjust: Likely needs a journal adjustment
- disclose: Needs a note disclosure in financial statements
- no_action: Appears to be a false positive or immaterial

Return ONLY the category name, nothing else."""

    try:
        result = _call_llm(
            "You are a concise accounting classifier. Return only the category name.",
            prompt,
            tier="haiku",
            max_tokens=20,
            temperature=0.1,
        )
        category = result.strip().lower().replace('"', '').replace("'", "")
        valid = ["investigate", "document", "adjust", "disclose", "no_action"]
        return category if category in valid else "investigate"
    except Exception:
        return "investigate"


# ---------------------------------------------------------------------------
# AI Risk Summary Report (Narrative Word Document) — uses Opus
# ---------------------------------------------------------------------------
REPORT_SYSTEM_PROMPT = """You are a senior Australian accounting professional writing an internal 
risk assessment report for a financial year. The report is for the review partner/manager at MC & S.

Write in a professional but accessible tone. You MUST structure the report EXACTLY as follows:

## 1. Executive Summary
2-3 sentences summarising the overall risk posture.

## 2. Materiality Assessment
State the materiality thresholds used and their basis.

## 3. Critical and High-Priority Findings
Detail each CRITICAL and HIGH severity flag with:
- What was found
- The dollar amount and whether it exceeds materiality
- The relevant legislation
- Specific recommended action

## 4. Medium and Low-Priority Findings
Summarise MEDIUM and LOW flags grouped by category.

## 5. ATO Compliance Posture
Overall assessment of the entity's ATO compliance risk, referencing the ATO interest scores.

## 6. Recommended Actions (Prioritised)
Numbered list of actions in priority order.

## 7. Conclusion
Final assessment and sign-off recommendation.

Use Australian accounting terminology. Reference specific account balances and variances.
Be specific — don't just repeat the flag descriptions, synthesise them into a coherent narrative.
Include dollar amounts and percentages where relevant."""


def generate_risk_summary_report(financial_year):
    """
    Generate a narrative AI Risk Summary Report as a Word document.
    Uses Opus tier for complex synthesis.
    Returns bytes of the .docx file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    entity = financial_year.entity
    flags = financial_year.risk_flags.all().order_by("-severity", "-ato_interest_score")

    # Build the AI prompt with materiality
    entity_context = _build_entity_context(financial_year, include_materiality=True)
    feedback_context = _get_feedback_context()

    flags_text = ""
    for flag in flags:
        flags_text += f"\n--- [{flag.severity}] {flag.title} ---\n"
        flags_text += f"Status: {flag.status}\n"
        flags_text += f"Description: {flag.description}\n"
        if flag.ai_explanation:
            flags_text += f"AI Analysis: {flag.ai_explanation}\n"
        if flag.ato_interest_score:
            flags_text += f"ATO Interest Score: {flag.ato_interest_score}/10\n"
        if flag.ato_interest_reasoning:
            flags_text += f"ATO Reasoning: {flag.ato_interest_reasoning}\n"
        if flag.resolution_notes:
            flags_text += f"Resolution: {flag.resolution_notes}\n"

    user_prompt = f"""Write a risk assessment report for this entity and financial year.

{entity_context}

=== RISK FLAGS ({flags.count()} total) ===
{flags_text}

Open flags: {flags.filter(status='open').count()}
Resolved flags: {flags.filter(status__in=['resolved', 'auto_resolved']).count()}
Reviewed flags: {flags.filter(status='reviewed').count()}
{feedback_context}

Write the full report now following the EXACT structure specified. Use ## for section headers."""

    try:
        report_text = _call_llm(
            REPORT_SYSTEM_PROMPT,
            user_prompt,
            tier="opus",
            max_tokens=4000,
            temperature=0.4,
        )
    except Exception as e:
        logger.exception("AI report generation failed")
        report_text = (
            f"## 1. Executive Summary\n\n"
            f"AI report generation failed: {str(e)}\n\n"
            f"Please review the {flags.count()} risk flags manually.\n\n"
            f"## 2. Materiality Assessment\n\nUnable to generate.\n\n"
            f"## 3. Critical and High-Priority Findings\n\nUnable to generate.\n\n"
            f"## 4. Medium and Low-Priority Findings\n\nUnable to generate.\n\n"
            f"## 5. ATO Compliance Posture\n\nUnable to generate.\n\n"
            f"## 6. Recommended Actions\n\nUnable to generate.\n\n"
            f"## 7. Conclusion\n\nManual review required."
        )

    # Build Word document
    doc = Document()

    # Title
    title = doc.add_heading(f"Risk Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{entity.entity_name} — {financial_year.year_label}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Entity details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run(f"Entity Type: {entity.get_entity_type_display()} | ABN: {entity.abn or 'N/A'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from django.utils import timezone
    run = meta.add_run(f"Generated: {timezone.now().strftime('%d %B %Y at %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # Spacer

    # Materiality summary table
    mat = _calculate_materiality(financial_year)
    mat_table = doc.add_table(rows=2, cols=4)
    mat_table.style = "Light Grid Accent 1"
    hdr = mat_table.rows[0].cells
    hdr[0].text = "Overall Materiality"
    hdr[1].text = "Performance Materiality"
    hdr[2].text = "Trivial Threshold"
    hdr[3].text = "Basis"
    row = mat_table.rows[1].cells
    row[0].text = f"${mat['overall_materiality']:,.0f}"
    row[1].text = f"${mat['performance_materiality']:,.0f}"
    row[2].text = f"${mat['trivial_threshold']:,.0f}"
    row[3].text = f"{mat['percentage']:.1f}% of {mat['base_type']}"

    doc.add_paragraph()

    # Summary stats table
    table = doc.add_table(rows=2, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Total Flags"
    hdr[1].text = "Open"
    hdr[2].text = "Resolved"
    hdr[3].text = "Avg ATO Score"
    hdr[4].text = "Critical"

    row = table.rows[1].cells
    row[0].text = str(flags.count())
    row[1].text = str(flags.filter(status="open").count())
    row[2].text = str(flags.filter(status__in=["resolved", "auto_resolved"]).count())

    scored_flags = flags.exclude(ato_interest_score=None)
    if scored_flags.exists():
        avg_score = sum(f.ato_interest_score for f in scored_flags) / scored_flags.count()
        row[3].text = f"{avg_score:.1f}/10"
    else:
        row[3].text = "N/A"

    row[4].text = str(flags.filter(severity="CRITICAL").count())

    doc.add_paragraph()  # Spacer

    # Parse and add the AI-generated report text
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith(("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. ")):
            doc.add_paragraph(line[3:], style="List Number")
        else:
            # Handle bold text
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("This report was generated by StatementHub AI Risk Engine")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run = footer.add_run("\nMC & S Accounting — Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
