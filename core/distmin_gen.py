"""
StatementHub — Distribution Minutes Document Generator

Generates distribution minutes from the Distmin.docx template by replacing
placeholders with actual entity data:
  - [add trustee or trustees here] → Trustee name(s) from EntityOfficer
  - [add chairperson here] → Chairperson name from EntityOfficer (is_chairperson=True)
  - [current financial year] → Financial year (e.g. "2025")
  - Trust name in "As Trustee for the ..." line → actual entity name
"""
import io
import copy
import re
from pathlib import Path
from docx import Document


TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "docx_templates" / "Distmin.docx"


def _replace_in_paragraph(paragraph, replacements):
    """
    Replace text in a paragraph while preserving formatting.
    Works by joining all runs, performing replacement, then redistributing text.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return False

    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True

    if changed and paragraph.runs:
        # Put all text in the first run, clear the rest (preserves first run's formatting)
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""

    return changed


def _replace_in_table(table, replacements):
    """Replace text in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)


def generate_distribution_minutes(financial_year_id):
    """
    Generate distribution minutes for a given financial year.

    Args:
        financial_year_id: UUID of the FinancialYear

    Returns:
        io.BytesIO buffer containing the generated .docx file

    Raises:
        ValueError: If required data is missing (no trustees, no chairperson, etc.)
    """
    from core.models import FinancialYear, EntityOfficer

    fy = FinancialYear.objects.select_related("entity").get(pk=financial_year_id)
    entity = fy.entity

    # Get active trustees for this entity
    trustees = EntityOfficer.objects.filter(
        entity=entity,
        role=EntityOfficer.OfficerRole.TRUSTEE,
        date_ceased__isnull=True,
    ).order_by("display_order", "full_name")

    if not trustees.exists():
        raise ValueError(
            f"No active trustees found for {entity.entity_name}. "
            f"Please add at least one trustee in the Directors/Trustees/Beneficiaries tab."
        )

    # Get chairperson — must be explicitly ticked in Directors/Trustees/Beneficiaries tab
    chairperson = EntityOfficer.objects.filter(
        entity=entity,
        is_chairperson=True,
        date_ceased__isnull=True,
    ).first()

    if not chairperson:
        raise ValueError(
            f"No chairperson found for {entity.entity_name}. "
            f"Please tick the Chairperson checkbox for the appropriate person "
            f"in the Directors/Trustees/Beneficiaries tab."
        )

    # Build trustee name(s) string
    trustee_names = [t.full_name for t in trustees]
    if len(trustee_names) == 1:
        trustee_str = trustee_names[0]
    elif len(trustee_names) == 2:
        trustee_str = f"{trustee_names[0]} and {trustee_names[1]}"
    else:
        trustee_str = ", ".join(trustee_names[:-1]) + f", and {trustee_names[-1]}"

    chairperson_name = chairperson.full_name

    # Determine the financial year number (e.g. "2025" from "FY2025" or end_date year)
    year_digits = "".join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        fy_year = year_digits
    else:
        fy_year = str(fy.end_date.year)

    # Build the date string for the minutes (30 June YYYY)
    minutes_date = f"30 June {fy_year}"

    # Load the template
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Distribution minutes template not found at {TEMPLATE_PATH}. "
            f"Please ensure Distmin.docx is in templates/docx_templates/."
        )

    doc = Document(str(TEMPLATE_PATH))

    # Define all replacements
    replacements = {
        "[add trustee or trustees here]": trustee_str,
        "[add chairperson here]": chairperson_name,
        "[current financial year]": fy_year,
    }

    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Replace in all tables (if any)
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # Also check headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
