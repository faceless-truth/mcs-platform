"""
Table-based financial statement layout helpers.

Replaces the old tab-stop-based _add_amount_line approach with proper Word tables.
This ensures that thin lines (borders) only appear under the amount columns,
not spanning the full page width — matching the CORRECT reference document.

Each financial report section creates a table with invisible borders.
Subtotals get thin top borders on amount cells only.
Grand totals get bold text + thin top border + double bottom border on amount cells.
"""

from decimal import Decimal
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_SUBHEADING = Pt(12)


def _set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    Usage: _set_cell_border(cell, top={"sz": 4, "val": "single", "color": "000000"})
    Supported keys: top, bottom, start (left), end (right)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        edge_tag = edge  # top, bottom, start, end
        element = tcBorders.find(qn(f'w:{edge_tag}'))
        if element is None:
            element = parse_xml(
                f'<w:{edge_tag} {nsdecls("w")} '
                f'w:val="{attrs.get("val", "single")}" '
                f'w:sz="{attrs.get("sz", 4)}" '
                f'w:space="0" '
                f'w:color="{attrs.get("color", "000000")}"/>'
            )
            tcBorders.append(element)
        else:
            element.set(qn('w:val'), attrs.get('val', 'single'))
            element.set(qn('w:sz'), str(attrs.get('sz', 4)))
            element.set(qn('w:color'), attrs.get('color', '000000'))


def _clear_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    # Set all borders to none explicitly
    _set_cell_border(cell,
                     top={"val": "none", "sz": 0, "color": "auto"},
                     bottom={"val": "none", "sz": 0, "color": "auto"},
                     start={"val": "none", "sz": 0, "color": "auto"},
                     end={"val": "none", "sz": 0, "color": "auto"})


def _set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, name=FONT_NAME):
    """Apply font formatting to a run."""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)
    return run


def _fmt(amount, show_cents=False):
    """Format a Decimal as Australian currency string without $ sign."""
    if amount is None:
        return "-"
    from decimal import ROUND_HALF_UP
    d = Decimal(str(amount))
    if show_cents:
        val = d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    else:
        val = d.quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    if val == 0:
        return "-"
    if show_cents:
        if val < 0:
            return f"({abs(val):,.2f})"
        return f"{val:,.2f}"
    else:
        if val < 0:
            return f"({abs(val):,.0f})"
        return f"{val:,.0f}"


# =============================================================================
# Financial Statement Table
# =============================================================================

class FinancialTable:
    """
    A table-based layout for financial statement data.
    
    Creates a Word table with columns:
    - Label (wide, left-aligned)
    - Note (optional, narrow, right-aligned)
    - Current Year amount (right-aligned)
    - Prior Year amount (right-aligned, optional)
    
    All cell borders are invisible by default.
    Subtotals get thin top border on amount cells.
    Grand totals get thin top border + double bottom border on amount cells, bold text.
    """
    
    # Column widths in cm
    # For has_prior=True, include_note=True:  Label(9.5) + Note(1.5) + Current(2.5) + Prior(2.5) = 16cm
    # For has_prior=True, include_note=False: Label(11) + Current(2.5) + Prior(2.5) = 16cm
    # For has_prior=False, include_note=True: Label(11) + Note(1.5) + Current(3.5) = 16cm
    # For has_prior=False, include_note=False: Label(12.5) + Current(3.5) = 16cm
    
    def __init__(self, doc, has_prior=False, include_note=False, show_cents=False):
        self.doc = doc
        self.has_prior = has_prior
        self.include_note = include_note
        self.show_cents = show_cents
        
        # Calculate column count and widths
        self.num_cols = 2  # label + current
        if include_note:
            self.num_cols += 1
        if has_prior:
            self.num_cols += 1
        
        # Available width = page width - margins = 21cm - 2*2.54cm = 15.92cm ≈ 16cm
        available = 16.0
        amt_width = 2.5 if has_prior else 3.5
        note_width = 1.5 if include_note else 0
        label_width = available - (amt_width * (2 if has_prior else 1)) - note_width
        
        self.col_widths = [label_width]
        if include_note:
            self.col_widths.append(note_width)
        self.col_widths.append(amt_width)
        if has_prior:
            self.col_widths.append(amt_width)
        
        # Column indices
        self.label_idx = 0
        self.note_idx = 1 if include_note else None
        self.current_idx = (2 if include_note else 1)
        self.prior_idx = (self.current_idx + 1) if has_prior else None
        
        # Create the table
        self.table = doc.add_table(rows=0, cols=self.num_cols)
        self.table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.table.autofit = False
        
        # Remove all table borders
        tbl = self.table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        # Set table borders to none
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        # Remove existing borders if any
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblBorders)
        
        # Fixed layout
        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tblPr.append(tblLayout)
    
    def _allow_row_split(self, row):
        """Ensure a row is allowed to split across pages (cantSplit=false)."""
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
            tr.insert(0, trPr)
        # Remove any existing cantSplit
        existing = trPr.find(qn('w:cantSplit'))
        if existing is not None:
            trPr.remove(existing)
        # Explicitly set cantSplit to false
        cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="false"/>')
        trPr.append(cantSplit)
    
    def _keep_with_next(self, row):
        """Set keep-with-next on all paragraphs in a row so it stays with the following row."""
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                existing = pPr.find(qn('w:keepNext'))
                if existing is None:
                    keepNext = parse_xml(f'<w:keepNext {nsdecls("w")}/>')
                    pPr.append(keepNext)
    
    def _set_cell_width(self, cell, width_cm):
        """Set cell width."""
        cell.width = Cm(width_cm)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>')
            tcPr.append(tcW)
        else:
            tcW.set(qn('w:w'), str(int(width_cm * 567)))
            tcW.set(qn('w:type'), 'dxa')
    
    def _format_cell(self, cell, text, align=WD_ALIGN_PARAGRAPH.LEFT,
                     size=FONT_SIZE_BODY, bold=False, italic=False):
        """Format a cell with text and styling."""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if text:
            run = p.add_run(str(text))
            _set_run_font(run, size=size, bold=bold, italic=italic)
        # Clear default borders
        _clear_cell_borders(cell)
    
    def add_section_heading(self, label, size=FONT_SIZE_SUBHEADING, bold=True,
                           space_before=10, keep_with_next=False):
        """Add a section heading row (e.g., 'Income', 'Current Assets')."""
        row = self.table.add_row()
        self._allow_row_split(row)
        # Merge all cells for heading
        if self.num_cols > 1:
            merged = row.cells[0].merge(row.cells[self.num_cols - 1])
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(label)
            _set_run_font(run, size=size, bold=bold)
            _clear_cell_borders(merged)
        else:
            self._format_cell(row.cells[0], label, size=size, bold=bold)
        if keep_with_next:
            self._keep_with_next(row)
    
    def add_sub_heading(self, label, size=FONT_SIZE_BODY, bold=True, italic=False,
                        space_before=6):
        """Add a sub-heading row (e.g., 'Cash Assets', 'Payables')."""
        row = self.table.add_row()
        self._allow_row_split(row)
        if self.num_cols > 1:
            merged = row.cells[0].merge(row.cells[self.num_cols - 1])
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(label)
            _set_run_font(run, size=size, bold=bold, italic=italic)
            _clear_cell_borders(merged)
        else:
            self._format_cell(row.cells[0], label, size=size, bold=bold, italic=italic)
    
    def add_line(self, label, current=None, prior=None, note_ref="",
                 bold=False, indent=0, size=FONT_SIZE_BODY, keep_with_next=False):
        """Add a regular data line with label and amounts."""
        row = self.table.add_row()
        self._allow_row_split(row)
        
        # Label cell
        cell = row.cells[self.label_idx]
        self._set_cell_width(cell, self.col_widths[self.label_idx])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if indent > 0:
            p.paragraph_format.left_indent = Cm(indent * 0.5)
        run = p.add_run(label)
        _set_run_font(run, size=size, bold=bold)
        _clear_cell_borders(cell)
        
        # Note cell (if applicable)
        if self.note_idx is not None:
            cell = row.cells[self.note_idx]
            self._set_cell_width(cell, self.col_widths[self.note_idx])
            self._format_cell(cell, note_ref, align=WD_ALIGN_PARAGRAPH.RIGHT, size=size)
        
        # Current amount cell
        cell = row.cells[self.current_idx]
        self._set_cell_width(cell, self.col_widths[self.current_idx])
        current_str = _fmt(current, self.show_cents) if current is not None else ""
        self._format_cell(cell, current_str, align=WD_ALIGN_PARAGRAPH.RIGHT,
                         size=size, bold=bold)
        
        # Prior amount cell
        if self.prior_idx is not None:
            cell = row.cells[self.prior_idx]
            self._set_cell_width(cell, self.col_widths[self.prior_idx])
            prior_str = _fmt(prior, self.show_cents) if prior is not None else ""
            self._format_cell(cell, prior_str, align=WD_ALIGN_PARAGRAPH.RIGHT,
                             size=size, bold=bold)
        if keep_with_next:
            self._keep_with_next(row)
    
    def add_subtotal(self, label, current, prior=None, note_ref="",
                     bold=False, size=FONT_SIZE_BODY):
        """
        Add a subtotal line with thin top border on amount cells only.
        The label can be empty for inline subtotals.
        """
        row = self.table.add_row()
        self._allow_row_split(row)
        
        # Label cell
        cell = row.cells[self.label_idx]
        self._set_cell_width(cell, self.col_widths[self.label_idx])
        self._format_cell(cell, label, size=size, bold=bold)
        
        # Note cell
        if self.note_idx is not None:
            cell = row.cells[self.note_idx]
            self._set_cell_width(cell, self.col_widths[self.note_idx])
            self._format_cell(cell, note_ref, align=WD_ALIGN_PARAGRAPH.RIGHT, size=size)
        
        # Current amount cell — thin top border
        cell = row.cells[self.current_idx]
        self._set_cell_width(cell, self.col_widths[self.current_idx])
        current_str = _fmt(current, self.show_cents)
        self._format_cell(cell, current_str, align=WD_ALIGN_PARAGRAPH.RIGHT,
                         size=size, bold=bold)
        _clear_cell_borders(cell)
        _set_cell_border(cell, top={"val": "single", "sz": 4, "color": "000000"})
        
        # Prior amount cell — thin top border
        if self.prior_idx is not None:
            cell = row.cells[self.prior_idx]
            self._set_cell_width(cell, self.col_widths[self.prior_idx])
            prior_str = _fmt(prior, self.show_cents) if prior is not None else ""
            self._format_cell(cell, prior_str, align=WD_ALIGN_PARAGRAPH.RIGHT,
                             size=size, bold=bold)
            _clear_cell_borders(cell)
            _set_cell_border(cell, top={"val": "single", "sz": 4, "color": "000000"})
    
    def add_total(self, label, current, prior=None, note_ref="",
                  size=FONT_SIZE_BODY, is_grand_total=False):
        """
        Add a total line: bold, thin top border on amount cells.
        If is_grand_total=True, also add double bottom border (=) on amount cells.
        """
        row = self.table.add_row()
        self._allow_row_split(row)
        
        # Label cell — bold
        cell = row.cells[self.label_idx]
        self._set_cell_width(cell, self.col_widths[self.label_idx])
        self._format_cell(cell, label, size=size, bold=True)
        
        # Note cell
        if self.note_idx is not None:
            cell = row.cells[self.note_idx]
            self._set_cell_width(cell, self.col_widths[self.note_idx])
            self._format_cell(cell, note_ref, align=WD_ALIGN_PARAGRAPH.RIGHT, size=size)
        
        # Current amount cell — thin top border, bold, optional double bottom
        cell = row.cells[self.current_idx]
        self._set_cell_width(cell, self.col_widths[self.current_idx])
        current_str = _fmt(current, self.show_cents)
        self._format_cell(cell, current_str, align=WD_ALIGN_PARAGRAPH.RIGHT,
                         size=size, bold=True)
        _clear_cell_borders(cell)
        borders = {"top": {"val": "single", "sz": 4, "color": "000000"}}
        if is_grand_total:
            borders["bottom"] = {"val": "double", "sz": 4, "color": "000000"}
        _set_cell_border(cell, **borders)
        
        # Prior amount cell
        if self.prior_idx is not None:
            cell = row.cells[self.prior_idx]
            self._set_cell_width(cell, self.col_widths[self.prior_idx])
            prior_str = _fmt(prior, self.show_cents) if prior is not None else ""
            self._format_cell(cell, prior_str, align=WD_ALIGN_PARAGRAPH.RIGHT,
                             size=size, bold=True)
            _clear_cell_borders(cell)
            borders = {"top": {"val": "single", "sz": 4, "color": "000000"}}
            if is_grand_total:
                borders["bottom"] = {"val": "double", "sz": 4, "color": "000000"}
            _set_cell_border(cell, **borders)
    
    def add_spacer(self, keep_with_next=False):
        """Add an empty row for spacing between sections."""
        row = self.table.add_row()
        self._allow_row_split(row)
        for i in range(self.num_cols):
            cell = row.cells[i]
            self._set_cell_width(cell, self.col_widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            _clear_cell_borders(cell)
        if keep_with_next:
            self._keep_with_next(row)
