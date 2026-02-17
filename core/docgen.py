"""
StatementHub — Document Generation Engine v3

Generates financial statements matching the exact format of the Access Ledger
PDF output. Supports Company, Trust, Partnership, and Sole Trader entity types.

Key features:
- Entity-type-aware section ordering, wording, and declarations
- Trading Account page for entities with COGS
- Sole trader: equity at top of Balance Sheet ("Proprietors' Funds")
- Trust: "Undistributed income", "Trustee's Declaration"
- Configurable cents/rounding per entity (accountant's discretion)
- Prior year comparative columns when data exists
- Summary P&L with income tax and dividends (companies only)
- Conditional accounting policy notes based on data present
"""
import io
from decimal import Decimal, ROUND_HALF_UP
from datetime import date
from pathlib import Path
from collections import OrderedDict

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .models import (
    Entity, FinancialYear, TrialBalanceLine, AccountMapping,
    EntityOfficer, NoteTemplate, DepreciationAsset,
)
from .table_helpers import FinancialTable

# =============================================================================
# Constants
# =============================================================================
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_HEADING = Pt(14)
FONT_SIZE_SUBHEADING = Pt(12)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_FOOTER = Pt(8)

FIRM_NAME = "M C & S Pty Ltd"
FIRM_ADDRESS_1 = "PO Box 4440"
FIRM_ADDRESS_2 = "Dandenong South VIC 3164"
FIRM_PHONE = "Phone: (03) 9794 0000"
FIRM_EMAIL = "Email: info@mcands.com.au"
FIRM_WEBSITE = "Website: www.mcands.com.au"


# =============================================================================
# Formatting Helpers
# =============================================================================

def _round_aud(amount, show_cents=False):
    """Round to nearest whole dollar or keep cents."""
    if amount is None:
        return Decimal("0")
    d = Decimal(str(amount))
    if show_cents:
        return d.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return d.quantize(Decimal("1"), rounding=ROUND_HALF_UP)


def _fmt(amount, show_cents=False):
    """Format a Decimal as Australian currency string without $ sign.
    Negatives in brackets. Zero as dash."""
    if amount is None:
        return "-"
    val = _round_aud(amount, show_cents)
    if val == 0:
        return "-"
    if show_cents:
        if val < 0:
            return f"({abs(val):,.2f})"
        return f"{val:,.2f}"
    else:
        if val < 0:
            return f"({abs(val):,.0f})"
        return f"{val:,.0f}"


def _set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, name=FONT_NAME):
    """Apply font formatting to a run."""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)
    return run


def _add_paragraph(doc, text="", size=FONT_SIZE_BODY, bold=False, italic=False,
                   underline=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before=0, space_after=Pt(4),
                   first_line_indent=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before) if isinstance(space_before, (int, float)) else space_before
    pf.space_after = space_after if isinstance(space_after, Emu) else Pt(space_after) if isinstance(space_after, (int, float)) else space_after
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold, italic=italic)
        if underline:
            run.font.underline = True
    return p


def _add_centered_heading(doc, text, size=FONT_SIZE_HEADING, bold=True, space_after=2):
    """Add a centered heading."""
    return _add_paragraph(doc, text, size=size, bold=bold,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=space_after)


def _add_horizontal_line(doc):
    """Add a horizontal line (thick rule)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_thin_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_header_block(doc, entity, title, date_text=None):
    """Add the standard header block: entity name, ABN, title, optional date.
    Used for pages that DON'T use section-based repeating headers (cover, contents, declaration, compilation)."""
    _add_centered_heading(doc, entity.entity_name, size=FONT_SIZE_HEADING, bold=True, space_after=0)
    if entity.trading_as:
        _add_centered_heading(doc, f"Trading As", size=Pt(11), bold=False, space_after=0)
    if entity.abn:
        _add_centered_heading(doc, f"ABN {entity.abn}", size=Pt(11), bold=True, space_after=0)
    _add_centered_heading(doc, title, size=FONT_SIZE_SUBHEADING, bold=True, space_after=0)
    if date_text:
        _add_centered_heading(doc, date_text, size=Pt(11), bold=True, space_after=2)
    _add_horizontal_line(doc)


def _add_header_para(header, text, size=FONT_SIZE_BODY, bold=False, italic=False):
    """Add a centered paragraph to a Word section header."""
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def _start_report_section(doc, entity, report_title, footer_type="statement",
                          year=None, prior_year=None, has_prior=False,
                          show_column_headers=True, include_note=False,
                          show_cents=False, landscape=False):
    """
    Start a new Word Section with repeating header and footer.
    
    The header contains: entity name (ALL CAPS, bold), Trading As (if set),
    ABN, report title, and optionally column headers (year/$).
    
    The footer contains the appropriate disclaimer text.
    
    This ensures that when content spans multiple pages, the header and
    footer repeat automatically on every page.
    """
    section = doc.add_section()
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Use explicit A4 dimensions to avoid swap-based bugs when consecutive
    # sections share the same orientation (e.g., multiple landscape depreciation pages).
    A4_SHORT = Cm(21.0)   # 210mm
    A4_LONG  = Cm(29.7)   # 297mm
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width  = A4_LONG
        section.page_height = A4_SHORT
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width  = A4_SHORT
        section.page_height = A4_LONG
    
    # Different first page = False (same header on all pages)
    section.different_first_page_header_footer = False
    
    # ---- Build the header ----
    header = section.header
    header.is_linked_to_previous = False
    
    # Clear any existing content
    for p in header.paragraphs:
        p.clear()
    
    # Entity name - bold, normal case (matching reference PDF)
    _add_header_para(header, entity.entity_name,
                     size=FONT_SIZE_HEADING, bold=True)
    
    # Trading As (only if set)
    if entity.trading_as:
        _add_header_para(header, "Trading As",
                         size=Pt(11), bold=False)
    
    # ABN
    if entity.abn:
        _add_header_para(header, f"ABN {entity.abn}",
                         size=Pt(11), bold=True)
    
    # Report title
    _add_header_para(header, report_title,
                     size=FONT_SIZE_SUBHEADING, bold=True)
    
    # Column headers (year / $) if requested
    if show_column_headers and year:
        # Year line
        p = header.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        tab_stops = p.paragraph_format.tab_stops
        if has_prior:
            if include_note:
                tab_stops.add_tab_stop(Cm(12), WD_ALIGN_PARAGRAPH.RIGHT)
            tab_stops.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT)
            tab_stops.add_tab_stop(Cm(16.5), WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            if include_note:
                tab_stops.add_tab_stop(Cm(12), WD_ALIGN_PARAGRAPH.RIGHT)
            tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
        
        if include_note:
            run = p.add_run("\tNote")
            _set_run_font(run, size=FONT_SIZE_BODY, bold=True)
        run = p.add_run(f"\t{year}")
        _set_run_font(run, size=FONT_SIZE_BODY, bold=True)
        if has_prior and prior_year:
            run = p.add_run(f"\t{prior_year}")
            _set_run_font(run, size=FONT_SIZE_BODY, bold=True)
        
        # Dollar sign line
        p2 = header.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        tab_stops2 = p2.paragraph_format.tab_stops
        if has_prior:
            tab_stops2.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT)
            tab_stops2.add_tab_stop(Cm(16.5), WD_ALIGN_PARAGRAPH.RIGHT)
            run = p2.add_run(f"\t$\t$")
        else:
            tab_stops2.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
            run = p2.add_run(f"\t$")
        _set_run_font(run, size=FONT_SIZE_BODY)
        
        # Horizontal line in header (thin)
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>' 
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        p2._element.get_or_add_pPr().append(pBdr)
    else:
        # Just add a horizontal line after the title (thin)
        p_line = header.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(0)
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>' 
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        p_line._element.get_or_add_pPr().append(pBdr)
    
    # ---- Build the footer ----
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Clear existing
    for p in footer.paragraphs:
        p.clear()
    
    # Horizontal line (thin)
    p_line = footer.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(2)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>' 
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    p_line._element.get_or_add_pPr().append(pBdr)
    
    if footer_type == "statement":
        text = (
            "These financial statements are unaudited. They must be read in conjunction "
            "with the attached Accountant's Compilation Report and Notes which form part "
            "of these financial statements."
        )
    elif footer_type == "notes":
        text = (
            f"These notes should be read in conjunction with the attached financial "
            f"statements and compilation report of {FIRM_NAME}."
        )
    else:
        text = ""
    
    if text:
        p_footer = footer.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer.paragraph_format.space_before = Pt(0)
        p_footer.paragraph_format.space_after = Pt(0)
        run = p_footer.add_run(text)
        _set_run_font(run, size=FONT_SIZE_FOOTER, bold=True)
    
    return section


def _get_period_text(fy):
    """
    Get the period description based on the financial year's period_type.
    Annual:   'For the year ended 30 June 2025'
    Half:     'For the half-year ended 31 December 2024'
    Quarter:  'For the quarter ended 30 September 2024'
    Monthly:  'For the month ended 31 January 2025'
    Interim:  'For the period ended 31 March 2025'
    """
    end_str = fy.end_date.strftime('%-d %B %Y')
    period_type = getattr(fy, 'period_type', 'annual') or 'annual'

    period_labels = {
        'annual': 'year',
        'half_year': 'half-year',
        'quarterly': 'quarter',
        'monthly': 'month',
        'interim': 'period',
    }
    label = period_labels.get(period_type, 'year')
    return f"For the {label} ended {end_str}"


def _get_period_label(fy):
    """Get just the period label word (year, quarter, month, etc.)."""
    period_type = getattr(fy, 'period_type', 'annual') or 'annual'
    period_labels = {
        'annual': 'year',
        'half_year': 'half-year',
        'quarterly': 'quarter',
        'monthly': 'month',
        'interim': 'period',
    }
    return period_labels.get(period_type, 'year')


def _get_as_at_text(fy):
    """Get 'as at DD Month YYYY'."""
    return f"as at {fy.end_date.strftime('%-d %B %Y')}"


def _add_statement_footer(doc):
    """Add the standard footer text for P&L and BS pages."""
    _add_horizontal_line(doc)
    text = (
        "These financial statements are unaudited. They must be read in conjunction "
        "with the attached Accountant's Compilation Report and Notes which form part "
        "of these financial statements."
    )
    _add_paragraph(doc, text, size=FONT_SIZE_FOOTER, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)


def _add_notes_footer(doc):
    """Add the footer text for Notes pages."""
    _add_horizontal_line(doc)
    text = (
        f"These notes should be read in conjunction with the attached financial "
        f"statements and compilation report of {FIRM_NAME}."
    )
    _add_paragraph(doc, text, size=FONT_SIZE_FOOTER, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)


def _entity_label(entity_type, plural=False):
    """Get the responsible party label for an entity type."""
    labels = {
        "company": ("the director", "the directors"),
        "trust": ("the trustee", "the trustee"),
        "partnership": ("the partners", "the partners"),
        "sole_trader": ("the owner", "the owner"),
    }
    pair = labels.get(entity_type, ("the director", "the directors"))
    return pair[1] if plural else pair[0]


def _entity_ref(entity_type):
    """Get 'the company' / 'the trust' / 'the partnership' / 'the business'."""
    refs = {
        "company": "the company",
        "trust": "the trust",
        "partnership": "the partnership",
        "sole_trader": "the business",
    }
    return refs.get(entity_type, "the company")


# =============================================================================
# Trial Balance Data Extraction
# =============================================================================

def _get_tb_sections(fy):
    """
    Extract trial balance lines grouped into financial statement sections.
    Returns dict with keys: trading_income, cogs, income, expenses,
    current_assets, noncurrent_assets, current_liabilities,
    noncurrent_liabilities, equity.
    """
    lines = fy.trial_balance_lines.order_by("account_code").all()
    sections = {
        "trading_income": [],
        "cogs": [],
        "income": [],
        "expenses": [],
        "current_assets": [],
        "noncurrent_assets": [],
        "current_liabilities": [],
        "noncurrent_liabilities": [],
        "equity": [],
    }

    for line in lines:
        try:
            code_num = int(line.account_code.split('.')[0])
        except (ValueError, TypeError):
            continue

        # Calculate current year amount: debit - credit gives net movement
        # For income (credit balances): debit=0, credit=X -> net = -X (negative = income)
        # For expenses (debit balances): debit=X, credit=0 -> net = X (positive = expense)
        # For assets (debit balances): debit=X, credit=0 -> net = X (positive = asset)
        # For liabilities (credit balances): debit=0, credit=X -> net = -X (negative = liability)
        current_amount = line.debit - line.credit
        prior_amount = line.prior_debit - line.prior_credit
        entry = (line.account_code, line.account_name, current_amount, prior_amount)

        # Check for COGS/trading accounts (code range 5000-5999 or specific patterns)
        name_lower = line.account_name.lower()
        is_cogs = (
            "cost of" in name_lower or "opening stock" in name_lower or
            "closing stock" in name_lower or "purchases" in name_lower or
            "stock on hand" in name_lower
        )

        if code_num < 1000:
            # 0000-0999: Income accounts
            # Determine if this is trading income or other income
            is_trading = (
                "sales" in name_lower or "income" in name_lower or
                "takings" in name_lower or "revenue" in name_lower or
                "accommodation" in name_lower or "conference" in name_lower or
                "meals" in name_lower or "bar" in name_lower or
                "trading" in name_lower
            )
            is_other_income = (
                "interest" in name_lower or "other" in name_lower or
                "fbt" in name_lower or "contribution" in name_lower or
                "dividend" in name_lower or "sundry" in name_lower
            )
            if is_other_income:
                sections["income"].append(entry)
            elif is_trading:
                sections["trading_income"].append(entry)
            else:
                sections["income"].append(entry)
        elif code_num < 1200:
            # 1000-1199: COGS / Cost of Sales accounts
            sections["cogs"].append(entry)
        elif code_num < 2000:
            # 1200-1999: Expense accounts
            if is_cogs:
                sections["cogs"].append(entry)
            else:
                sections["expenses"].append(entry)
        elif code_num < 2500:
            # 2000-2499: Current assets
            sections["current_assets"].append(entry)
        elif code_num < 3000:
            # 2500-2999: Non-current assets (PPE, loans receivable, etc.)
            sections["noncurrent_assets"].append(entry)
        elif code_num < 3500:
            # 3000-3499: Current liabilities
            sections["current_liabilities"].append(entry)
        elif code_num < 4000:
            # 3500-3999: Non-current liabilities
            sections["noncurrent_liabilities"].append(entry)
        elif code_num < 5000:
            # 4000-4999: Equity accounts
            sections["equity"].append(entry)
        elif code_num < 6000:
            # 5000-5999 range: COGS/trading (alternative code range)
            sections["cogs"].append(entry)

    return sections


def _get_prior_balance(fy, account_code):
    """Get the prior year closing balance for an account code."""
    if not fy.prior_year:
        return Decimal("0")
    try:
        prior_line = fy.prior_year.trial_balance_lines.get(account_code=account_code)
        return prior_line.closing_balance
    except TrialBalanceLine.DoesNotExist:
        return Decimal("0")


def _has_prior_year(fy):
    """Check if there is prior year data."""
    if not fy.prior_year:
        return False
    return fy.prior_year.trial_balance_lines.exists()


def _has_cogs(sections):
    """Check if the entity has COGS/trading accounts."""
    return len(sections["cogs"]) > 0


# =============================================================================
# Note Registry — Automatic note assignment per AASB 101 para 113
# =============================================================================

class NoteRegistry:
    """
    Automatically assigns note numbers based on what data exists in the
    trial balance sections. Notes are numbered sequentially starting from 1.
    
    Per AASB 101 paragraph 113, each item in the financial statements shall
    be cross-referenced to any related information in the notes.
    
    Standard note structure for MC&S special purpose financial statements:
    - Note 1: Summary of Significant Accounting Policies (ALWAYS)
    - Note 2: Revenue (if trading_income or income data exists)
    - Note 3: Profit from Ordinary Activities (if depreciation, borrowing costs, COGS, or bad debts)
    - Note 4: Retained Profits / Undistributed Income (companies and trusts with equity data)
    """
    
    def __init__(self, entity, sections):
        self.entity_type = entity.entity_type
        self.notes = {}  # key -> note number
        self._next_num = 1
        
        # Note 1: Accounting Policies — always present
        self._assign("accounting_policies")
        
        # Note 2: Revenue — if there's any income data
        has_revenue = (len(sections["trading_income"]) > 0 or
                       len(sections["income"]) > 0)
        if has_revenue:
            self._assign("revenue")
        
        # Note 3: Profit from Ordinary Activities — if depreciation, borrowing, COGS, or bad debts
        has_depreciation = any("depreciation" in n.lower() or "amortisation" in n.lower()
                               for _, n, _, _ in sections["expenses"])
        has_borrowing = any("interest" in n.lower() and
                           ("loan" in n.lower() or "australia" in n.lower() or "mortgage" in n.lower())
                           for _, n, _, _ in sections["expenses"])
        has_cogs = len(sections["cogs"]) > 0
        has_bad_debts = any("bad" in n.lower() and "debt" in n.lower()
                           for _, n, _, _ in sections["expenses"])
        if has_depreciation or has_borrowing or has_cogs or has_bad_debts:
            self._assign("profit_ordinary")
        
        # Note 4: Retained Profits / Undistributed Income
        # For companies and trusts that have equity data
        if self.entity_type in ("company", "trust"):
            has_equity = len(sections["equity"]) > 0
            if has_equity:
                self._assign("retained_profits")
    
    def _assign(self, key):
        """Assign the next sequential note number to a key."""
        self.notes[key] = self._next_num
        self._next_num += 1
    
    def get(self, key):
        """Get the note number for a key, or empty string if not assigned."""
        num = self.notes.get(key)
        return str(num) if num else ""
    
    def has(self, key):
        """Check if a note exists for a key."""
        return key in self.notes
    
    def get_num(self, key):
        """Get the raw note number (int) or None."""
        return self.notes.get(key)


def _build_note_registry(entity, sections):
    """Build a NoteRegistry based on entity type and trial balance data."""
    return NoteRegistry(entity, sections)


# =============================================================================
# Financial Statement Line Helpers
# =============================================================================

def _add_amount_line(doc, label, current, prior=None, has_prior=False,
                     bold=False, indent=0, size=FONT_SIZE_BODY, note_ref="",
                     is_section_heading=False, heading_size=None,
                     show_cents=False, is_subtotal=False, is_total=False):
    """Add a single line to a financial statement using tab stops.
    
    Formatting matches the trust.docx reference:
    - is_subtotal: thin top border on the paragraph (line above the amount)
    - is_total: bold text, thin top border on the paragraph
    - No underlines on individual amounts ever
    - Section headings are bold and larger
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)

    if is_section_heading:
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)

    # Add top border for subtotals and totals (thin line above)
    if is_subtotal or is_total:
        pf.space_before = Pt(4)
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        p._element.get_or_add_pPr().append(pBdr)

    # Force bold for totals
    if is_total:
        bold = True

    # Tab stops for alignment
    tab_stops = pf.tab_stops
    if has_prior:
        tab_stops.add_tab_stop(Cm(12), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops.add_tab_stop(Cm(16.5), WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        tab_stops.add_tab_stop(Cm(12), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)

    # Indent
    if indent > 0:
        pf.left_indent = Cm(indent * 0.5)

    # Label
    use_size = heading_size if heading_size else size
    run = p.add_run(label)
    _set_run_font(run, size=use_size, bold=bold)

    # Only add amounts for non-section-heading lines
    if not is_section_heading and current is not None:
        if note_ref:
            run = p.add_run(f"\t{note_ref}")
            _set_run_font(run, size=size)

        current_str = _fmt(current, show_cents) if current is not None else ""
        run = p.add_run(f"\t{current_str}")
        _set_run_font(run, size=size, bold=bold)

        if has_prior:
            prior_str = _fmt(prior, show_cents) if prior is not None else ""
            run = p.add_run(f"\t{prior_str}")
            _set_run_font(run, size=size, bold=bold)

    return p


def _add_column_headers(doc, year, has_prior=False, prior_year=None, include_note=False,
                        show_cents=False):
    """Add the year column headers (e.g., 'Note    2025    2024')."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)

    tab_stops = pf.tab_stops
    if has_prior:
        if include_note:
            tab_stops.add_tab_stop(Cm(12), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops.add_tab_stop(Cm(16.5), WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        if include_note:
            tab_stops.add_tab_stop(Cm(12), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)

    if include_note:
        run = p.add_run("\tNote")
        _set_run_font(run, size=FONT_SIZE_BODY, bold=True)

    run = p.add_run(f"\t{year}")
    _set_run_font(run, size=FONT_SIZE_BODY, bold=True)

    if has_prior and prior_year:
        run = p.add_run(f"\t{prior_year}")
        _set_run_font(run, size=FONT_SIZE_BODY, bold=True)

    # Dollar sign line
    p2 = doc.add_paragraph()
    pf2 = p2.paragraph_format
    pf2.space_after = Pt(0)
    tab_stops2 = pf2.tab_stops
    if has_prior:
        tab_stops2.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT)
        tab_stops2.add_tab_stop(Cm(16.5), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p2.add_run(f"\t$\t$")
    else:
        tab_stops2.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p2.add_run(f"\t$")
    _set_run_font(run, size=FONT_SIZE_BODY)

    _add_horizontal_line(doc)


# =============================================================================
# Cover Page
# =============================================================================

def _get_logo_path():
    """Get the path to the MC&S logo file."""
    from django.conf import settings
    logo_path = getattr(settings, 'MCS_LOGO_PATH', None)
    if logo_path and Path(str(logo_path)).exists():
        return str(logo_path)
    # Fallback: check common locations
    fallbacks = [
        Path(__file__).resolve().parent.parent / 'static' / 'MCSlogo.png',
        Path('/home/ubuntu/upload/MCSlogo.png'),
    ]
    for fb in fallbacks:
        if fb.exists():
            return str(fb)
    return None


def _add_cover_page(doc, entity, fy):
    """Add the cover page with MC&S logo matching the reference PDF format."""
    # Small spacing before logo
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # Add MC&S logo — centered, approximately 7cm wide
    logo_path = _get_logo_path()
    if logo_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(7))

    # Spacing after logo
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(36)

    # Entity name — bold, centered, normal case (matching reference PDF)
    _add_centered_heading(doc, entity.entity_name, size=Pt(16), bold=True, space_after=4)

    # Trading As line
    if entity.trading_as:
        _add_centered_heading(doc, f"Trading As {entity.trading_as}", size=Pt(14),
                              bold=False, space_after=4)

    # ABN
    if entity.abn:
        _add_centered_heading(doc, f"ABN {entity.abn}", size=Pt(11), bold=False, space_after=12)

    # Spacing before Financial Statements title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)

    # "Financial Statements" title
    _add_centered_heading(doc, "Financial Statements", size=Pt(12), bold=False, space_after=2)

    # Period text
    _add_centered_heading(doc, _get_period_text(fy), size=Pt(11), bold=False, space_after=0)

    # Spacing before firm details — push to bottom of page
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    _add_centered_heading(doc, FIRM_NAME, size=Pt(10), bold=False, space_after=0)
    _add_centered_heading(doc, FIRM_ADDRESS_1, size=Pt(10), bold=False, space_after=0)
    _add_centered_heading(doc, FIRM_ADDRESS_2, size=Pt(10), bold=False, space_after=4)
    _add_centered_heading(doc, FIRM_PHONE, size=Pt(10), bold=False, space_after=0)
    _add_centered_heading(doc, FIRM_EMAIL, size=Pt(10), bold=False, space_after=0)
    _add_centered_heading(doc, FIRM_WEBSITE, size=Pt(10), bold=False, space_after=0)

    doc.add_page_break()


# =============================================================================
# Contents Page
# =============================================================================

def _get_section_order(entity, sections, fy=None):
    """Determine the section order based on entity type and data."""
    entity_type = entity.entity_type
    has_trading = _has_cogs(sections)
    # Only show depreciation schedule if DepreciationAsset records exist
    has_dep_schedule = False
    if fy:
        has_dep_schedule = DepreciationAsset.objects.filter(financial_year=fy).exists()
    has_ppe = has_dep_schedule  # Only show if actual depreciation assets are entered

    if entity_type == "company":
        items = []
        if has_trading:
            # Complex company: compilation report first
            items.append("Compilation Report")
            items.append("Trading Account")
        items.append("Detailed Profit and Loss Statement")
        items.append("Detailed Balance Sheet")
        if not has_trading:
            # Simple company: no summary P&L unless we have it
            pass
        else:
            items.append("Profit and Loss Statement")
        if has_ppe:
            items.append("Depreciation Schedule")
        items.append("Notes to the Financial Statements")
        items.append("Director's Declaration")
        if not has_trading:
            items.append("Compilation Report")
        return items

    elif entity_type == "trust":
        items = []
        items.append("Detailed Profit and Loss Statement")
        items.append("Detailed Balance Sheet")
        items.append("Notes to the Financial Statements")
        if has_ppe:
            items.append("Depreciation Schedule")
        items.append("Trustee's Declaration")
        items.append("Compilation Report")
        return items

    elif entity_type == "partnership":
        items = []
        if has_trading:
            items.append("Trading Account")
        items.append("Detailed Profit and Loss Statement")
        items.append("Detailed Balance Sheet")
        items.append("Partners' Profit Distribution Summary")
        if has_ppe:
            items.append("Depreciation Schedule")
        items.append("Notes to the Financial Statements")
        items.append("Partner Declaration")
        items.append("Compilation Report")
        return items

    else:  # sole_trader
        items = []
        if has_trading:
            items.append("Trading Account")
        items.append("Detailed Profit and Loss Statement")
        items.append("Detailed Balance Sheet")
        items.append("Notes to the Financial Statements")
        if has_ppe:
            items.append("Depreciation Schedule")
        items.append("Compilation Report")
        items.append("Proprietor Declaration")
        return items


def _add_contents_page(doc, entity, fy, sections):
    """Add the table of contents."""
    _add_centered_heading(doc, entity.entity_name, size=FONT_SIZE_HEADING, bold=True, space_after=0)
    if entity.abn:
        _add_centered_heading(doc, f"ABN {entity.abn}", size=Pt(11), bold=True, space_after=12)

    _add_paragraph(doc, "Contents", size=FONT_SIZE_HEADING, bold=True, space_after=12)

    items = _get_section_order(entity, sections, fy=fy)
    for item in items:
        p = _add_paragraph(doc, item, size=Pt(11), space_after=6)
        for run in p.runs:
            run.underline = True

    doc.add_page_break()


# =============================================================================
# Trading Account
# =============================================================================

def _add_trading_account(doc, entity, fy, sections, show_cents=False):
    """Add the Trading Account page (for entities with COGS)."""
    has_prior = _has_prior_year(fy)
    year = str(fy.end_date.year)
    prior_year_str = str(fy.end_date.year - 1) if has_prior else None

    _start_report_section(doc, entity,
                          f"Trading Account\n{_get_period_text(fy)}",
                          footer_type="statement",
                          year=year, prior_year=prior_year_str,
                          has_prior=has_prior, show_column_headers=True,
                          include_note=False, show_cents=show_cents)

    ft = FinancialTable(doc, has_prior=has_prior, include_note=False, show_cents=show_cents)

    # Trading Income
    total_trading_income = Decimal("0")
    total_trading_income_prior = Decimal("0")

    ft.add_section_heading("Trading Income")

    for code, name, balance, prior in sections["trading_income"]:
        val = abs(balance)
        prior_val = abs(prior) if prior else Decimal("0")
        total_trading_income += val
        total_trading_income_prior += prior_val
        ft.add_line(name, val, prior_val, indent=1)

    ft.add_total("Total Trading Income", total_trading_income,
                 total_trading_income_prior)

    ft.add_spacer()

    # Cost of Sales
    ft.add_section_heading("Cost of Sales")

    total_cogs = Decimal("0")
    total_cogs_prior = Decimal("0")

    # Separate opening stock, purchases, and closing stock
    opening_stock = []
    closing_stock = []
    other_cogs = []

    for code, name, balance, prior in sections["cogs"]:
        name_lower = name.lower()
        if "opening" in name_lower:
            opening_stock.append((code, name, balance, prior))
        elif "closing" in name_lower:
            closing_stock.append((code, name, balance, prior))
        else:
            other_cogs.append((code, name, balance, prior))

    # Add: Opening Stock + Purchases
    add_items = opening_stock + other_cogs
    if add_items:
        ft.add_sub_heading("Add:")

    add_subtotal = Decimal("0")
    add_subtotal_prior = Decimal("0")
    for code, name, balance, prior in add_items:
        val = abs(balance) if balance else Decimal("0")
        prior_val = abs(prior) if prior else Decimal("0")
        add_subtotal += val
        add_subtotal_prior += prior_val
        total_cogs += val
        total_cogs_prior += prior_val
        ft.add_line(name, val, prior_val, indent=1)

    # Show add subtotal if there are multiple add items
    if len(add_items) > 1:
        ft.add_subtotal("", add_subtotal, add_subtotal_prior)

    # Less: Closing Stock
    if closing_stock:
        ft.add_sub_heading("Less:")
        for code, name, balance, prior in closing_stock:
            val = abs(balance) if balance else Decimal("0")
            prior_val = abs(prior) if prior else Decimal("0")
            total_cogs -= val  # Closing stock reduces COGS
            total_cogs_prior -= prior_val
            ft.add_subtotal(name, val, prior_val)

    ft.add_spacer()

    ft.add_total("Cost of Sales", total_cogs, total_cogs_prior)

    ft.add_spacer()

    # Gross Profit — grand total with double underline
    gross_profit = total_trading_income - total_cogs
    gross_profit_prior = total_trading_income_prior - total_cogs_prior

    ft.add_total("Gross Profit from Trading", gross_profit, gross_profit_prior,
                 is_grand_total=True)

    return gross_profit, gross_profit_prior


# =============================================================================
# Detailed Profit and Loss Statement
# =============================================================================

def _add_detailed_pnl(doc, entity, fy, sections, show_cents=False,
                      gross_profit=None, gross_profit_prior=None,
                      note_registry=None):
    """Add the detailed P&L."""
    has_prior = _has_prior_year(fy)
    year = str(fy.end_date.year)
    prior_year_str = str(fy.end_date.year - 1) if has_prior else None
    nr = note_registry

    _start_report_section(doc, entity,
                          f"Detailed Profit and Loss Statement\n{_get_period_text(fy)}",
                          footer_type="statement",
                          year=year, prior_year=prior_year_str,
                          has_prior=has_prior, show_column_headers=True,
                          include_note=True, show_cents=show_cents)

    ft = FinancialTable(doc, has_prior=has_prior, include_note=True, show_cents=show_cents)

    # Income section
    total_income = Decimal("0")
    total_income_prior = Decimal("0")

    ft.add_section_heading("Income")

    # If we have a trading account, first line is "Trading profit"
    if gross_profit is not None:
        ft.add_line("Trading profit", gross_profit, gross_profit_prior, indent=1)
        total_income += gross_profit
        total_income_prior += gross_profit_prior
    else:
        # Show all trading income as regular income
        for code, name, balance, prior in sections["trading_income"]:
            val = abs(balance)
            prior_val = abs(prior) if prior else Decimal("0")
            total_income += val
            total_income_prior += prior_val
            ft.add_line(name, val, prior_val, indent=1)

    # Other income
    for code, name, balance, prior in sections["income"]:
        val = abs(balance)
        prior_val = abs(prior) if prior else Decimal("0")
        total_income += val
        total_income_prior += prior_val
        ft.add_line(name, val, prior_val, indent=1)

    # Note ref for revenue
    revenue_note = nr.get("revenue") if nr else ""
    ft.add_subtotal("Total income", total_income, total_income_prior,
                    note_ref=revenue_note)

    ft.add_spacer()

    # Expenses section
    total_expenses = Decimal("0")
    total_expenses_prior = Decimal("0")

    ft.add_section_heading("Expenses")

    for code, name, balance, prior in sections["expenses"]:
        val = abs(balance)
        prior_val = abs(prior) if prior else Decimal("0")
        total_expenses += val
        total_expenses_prior += prior_val
        ft.add_line(name, val, prior_val, indent=1)

    ft.add_subtotal("Total expenses", total_expenses, total_expenses_prior)

    ft.add_spacer()

    # Net Profit/Loss — grand total with double underline
    net_profit = total_income - total_expenses
    net_profit_prior = total_income_prior - total_expenses_prior

    # Wording varies by entity type
    entity_type = entity.entity_type
    if entity_type in ("trust", "sole_trader"):
        profit_label = "Net Profit from Ordinary Activities before income tax"
    else:
        profit_label = "Profit (Loss) from Ordinary Activities before income tax"

    # Note ref for profit from ordinary activities
    profit_note = nr.get("profit_ordinary") if nr else ""
    ft.add_total(profit_label, net_profit, net_profit_prior,
                 is_grand_total=True, note_ref=profit_note)

    return net_profit, net_profit_prior


# =============================================================================
# Detailed Balance Sheet
# =============================================================================

def _add_detailed_balance_sheet(doc, entity, fy, sections, show_cents=False,
                                net_profit=Decimal("0"), net_profit_prior=Decimal("0"),
                                note_registry=None):
    """Add the detailed balance sheet."""
    has_prior = _has_prior_year(fy)
    year = str(fy.end_date.year)
    prior_year_str = str(fy.end_date.year - 1) if has_prior else None
    entity_type = entity.entity_type
    nr = note_registry

    _start_report_section(doc, entity,
                          f"Detailed Balance Sheet {_get_as_at_text(fy)}",
                          footer_type="statement",
                          year=year, prior_year=prior_year_str,
                          has_prior=has_prior, show_column_headers=True,
                          include_note=True, show_cents=show_cents)

    # TABLE 1: Assets (and sole trader equity at top)
    ft = FinancialTable(doc, has_prior=has_prior, include_note=True, show_cents=show_cents)

    # ---- SOLE TRADER: Equity at TOP ----
    if entity_type == "sole_trader":
        ft.add_section_heading("Proprietors' Funds")

        # Calculate proprietors' funds
        opening_balance = Decimal("0")
        opening_balance_prior = Decimal("0")
        drawings = Decimal("0")
        drawings_prior = Decimal("0")

        for code, name, balance, prior in sections["equity"]:
            name_lower = name.lower()
            if "drawing" in name_lower:
                drawings = abs(balance) if balance else Decimal("0")
                drawings_prior = abs(prior) if prior else Decimal("0")
            elif "opening" in name_lower or "capital" in name_lower or "retained" in name_lower:
                opening_balance = abs(balance) if balance < 0 else balance
                opening_balance_prior = abs(prior) if prior and prior < 0 else (prior or Decimal("0"))

        if opening_balance == 0 and not any("opening" in n.lower() or "capital" in n.lower()
                                             for _, n, _, _ in sections["equity"]):
            pass

        ft.add_line("Opening balance", opening_balance, opening_balance_prior)
        ft.add_line("Net profit / (loss)", net_profit, net_profit_prior)
        if drawings > 0 or drawings_prior > 0:
            ft.add_line("Less: Drawings", drawings, drawings_prior)

        total_prop_funds = opening_balance + net_profit - drawings
        total_prop_funds_prior = opening_balance_prior + net_profit_prior - drawings_prior

        ft.add_total("Total Proprietors' Funds", total_prop_funds,
                     total_prop_funds_prior, is_grand_total=True)

        ft.add_spacer()
        ft.add_sub_heading("Represented by:")

    # ---- Current Assets ----
    total_ca = Decimal("0")
    total_ca_prior = Decimal("0")

    if sections["current_assets"]:
        ft.add_section_heading("Current Assets")

        # Sub-categorise current assets
        cash_items = []
        receivable_items = []
        inventory_items = []
        other_ca_items = []

        for code, name, balance, prior in sections["current_assets"]:
            code_num = int(code)
            name_lower = name.lower()
            if "cash" in name_lower or "bank" in name_lower or "petty" in name_lower or code_num < 2100:
                cash_items.append((code, name, balance, prior))
            elif "debtor" in name_lower or "receivable" in name_lower or "trade" in name_lower:
                receivable_items.append((code, name, balance, prior))
            elif "stock" in name_lower or "inventor" in name_lower:
                inventory_items.append((code, name, balance, prior))
            else:
                other_ca_items.append((code, name, balance, prior))

        # Cash Assets
        if cash_items:
            ft.add_sub_heading("Cash Assets")
            sub_total = Decimal("0")
            sub_total_prior = Decimal("0")
            for code, name, balance, prior in cash_items:
                val = abs(balance) if balance > 0 else balance
                prior_val = abs(prior) if prior and prior > 0 else (prior or Decimal("0"))
                sub_total += val
                sub_total_prior += prior_val
                total_ca += val
                total_ca_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)
            if len(cash_items) > 1:
                ft.add_subtotal("", sub_total, sub_total_prior)

        # Receivables
        if receivable_items:
            ft.add_sub_heading("Receivables")
            for code, name, balance, prior in receivable_items:
                val = abs(balance) if balance > 0 else balance
                prior_val = abs(prior) if prior and prior > 0 else (prior or Decimal("0"))
                total_ca += val
                total_ca_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # Inventories
        if inventory_items:
            ft.add_sub_heading("Inventories")
            for code, name, balance, prior in inventory_items:
                val = abs(balance) if balance > 0 else balance
                prior_val = abs(prior) if prior and prior > 0 else (prior or Decimal("0"))
                total_ca += val
                total_ca_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # Other current assets
        for code, name, balance, prior in other_ca_items:
            val = abs(balance) if balance > 0 else balance
            prior_val = abs(prior) if prior and prior > 0 else (prior or Decimal("0"))
            total_ca += val
            total_ca_prior += prior_val
            ft.add_line(name, val, prior_val, indent=1)

        ft.add_subtotal("Total Current Assets", total_ca, total_ca_prior, bold=True)

    # ---- Non-Current Assets ----
    total_nca = Decimal("0")
    total_nca_prior = Decimal("0")

    if sections["noncurrent_assets"]:
        ft.add_section_heading("Non-Current Assets")

        # Sub-categorise non-current assets
        ppe_items = []
        investment_items = []
        receivable_nca_items = []
        inventory_nca_items = []
        other_nca_items = []

        for code, name, balance, prior in sections["noncurrent_assets"]:
            name_lower = name.lower()
            if ("equipment" in name_lower or "vehicle" in name_lower or "furniture" in name_lower or
                "building" in name_lower or "fixture" in name_lower or "plant" in name_lower or
                "motor" in name_lower or "computer" in name_lower or "office" in name_lower or
                "accumulated" in name_lower or "amortisation" in name_lower or
                "depreciation" in name_lower or "less:" in name_lower):
                ppe_items.append((code, name, balance, prior))
            elif "investment" in name_lower or "unit" in name_lower or "share" in name_lower or "financial asset" in name_lower:
                investment_items.append((code, name, balance, prior))
            elif "loan" in name_lower or "receivable" in name_lower or "debtor" in name_lower:
                receivable_nca_items.append((code, name, balance, prior))
            elif "land" in name_lower or "inventor" in name_lower or "stock" in name_lower:
                inventory_nca_items.append((code, name, balance, prior))
            else:
                other_nca_items.append((code, name, balance, prior))

        # NCA Receivables
        if receivable_nca_items:
            ft.add_sub_heading("Receivables")
            for code, name, balance, prior in receivable_nca_items:
                val = balance
                prior_val = prior or Decimal("0")
                total_nca += val
                total_nca_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # NCA Inventories (e.g., land held for resale)
        if inventory_nca_items:
            ft.add_sub_heading("Inventories")
            for code, name, balance, prior in inventory_nca_items:
                val = abs(balance) if balance > 0 else balance
                prior_val = abs(prior) if prior and prior > 0 else (prior or Decimal("0"))
                total_nca += val
                total_nca_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # Other Financial Assets
        if investment_items:
            ft.add_sub_heading("Other Financial Assets")
            for code, name, balance, prior in investment_items:
                val = abs(balance) if balance > 0 else balance
                prior_val = abs(prior) if prior and prior > 0 else (prior or Decimal("0"))
                total_nca += val
                total_nca_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # PPE
        if ppe_items:
            ft.add_sub_heading("Property, Plant and Equipment")
            ppe_total = Decimal("0")
            ppe_total_prior = Decimal("0")
            for code, name, balance, prior in ppe_items:
                name_lower = name.lower()
                if "accumulated" in name_lower or "amortisation" in name_lower or "less:" in name_lower:
                    val = -abs(balance) if balance else Decimal("0")
                    prior_val = -abs(prior) if prior else Decimal("0")
                else:
                    val = abs(balance) if balance else Decimal("0")
                    prior_val = abs(prior) if prior else Decimal("0")
                ppe_total += val
                ppe_total_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

            ft.add_subtotal("", ppe_total, ppe_total_prior)
            total_nca += ppe_total
            total_nca_prior += ppe_total_prior

        # Other NCA
        for code, name, balance, prior in other_nca_items:
            val = balance
            prior_val = prior or Decimal("0")
            total_nca += val
            total_nca_prior += prior_val
            ft.add_line(name, val, prior_val, indent=1)

        ft.add_subtotal("Total Non-Current Assets", total_nca, total_nca_prior, bold=True)

    # Total Assets — grand total with double underline
    total_assets = total_ca + total_nca
    total_assets_prior = total_ca_prior + total_nca_prior
    ft.add_spacer()
    ft.add_total("Total Assets", total_assets, total_assets_prior, is_grand_total=True)

    # TABLE 2: Liabilities (separate table for better pagination)
    ft = FinancialTable(doc, has_prior=has_prior, include_note=True, show_cents=show_cents)

    # ---- Liabilities ----
    total_cl = Decimal("0")
    total_cl_prior = Decimal("0")

    if sections["current_liabilities"]:
        ft.add_section_heading("Current Liabilities")

        payable_items = []
        tax_items = []
        provision_items = []
        other_cl_items = []

        for code, name, balance, prior in sections["current_liabilities"]:
            name_lower = name.lower()
            if "gst" in name_lower or "tax" in name_lower or "payg" in name_lower or "super" in name_lower:
                tax_items.append((code, name, balance, prior))
            elif "creditor" in name_lower or "credit card" in name_lower or "payable" in name_lower:
                payable_items.append((code, name, balance, prior))
            elif "provision" in name_lower or "leave" in name_lower or "lsl" in name_lower:
                provision_items.append((code, name, balance, prior))
            else:
                other_cl_items.append((code, name, balance, prior))

        # Payables
        if payable_items:
            ft.add_sub_heading("Payables")
            secured = [i for i in payable_items if "secured" in i[1].lower()]
            unsecured = [i for i in payable_items if "secured" not in i[1].lower()]
            if secured:
                ft.add_sub_heading("Secured:", italic=True)
                for code, name, balance, prior in secured:
                    val = abs(balance)
                    prior_val = abs(prior) if prior else Decimal("0")
                    total_cl += val
                    total_cl_prior += prior_val
                    ft.add_line(name, val, prior_val, indent=1)
            if unsecured:
                if secured:
                    ft.add_sub_heading("Unsecured:", italic=True)
                for code, name, balance, prior in unsecured:
                    val = abs(balance)
                    prior_val = abs(prior) if prior else Decimal("0")
                    total_cl += val
                    total_cl_prior += prior_val
                    ft.add_line(name, val, prior_val, indent=1)

        # Current Tax Liabilities
        if tax_items:
            ft.add_sub_heading("Current Tax Liabilities")
            for code, name, balance, prior in tax_items:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_cl += val
                total_cl_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # Provisions
        if provision_items:
            ft.add_sub_heading("Provisions")
            for code, name, balance, prior in provision_items:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_cl += val
                total_cl_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        # Other CL
        if other_cl_items:
            for code, name, balance, prior in other_cl_items:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_cl += val
                total_cl_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        ft.add_subtotal("Total Current Liabilities", total_cl, total_cl_prior, bold=True)

    # ---- Non-Current Liabilities ----
    total_ncl = Decimal("0")
    total_ncl_prior = Decimal("0")

    if sections["noncurrent_liabilities"]:
        ft.add_section_heading("Non-Current Liabilities")

        loan_items = []
        other_ncl_items = []

        for code, name, balance, prior in sections["noncurrent_liabilities"]:
            name_lower = name.lower()
            if "loan" in name_lower or "mortgage" in name_lower or "borrowing" in name_lower:
                loan_items.append((code, name, balance, prior))
            else:
                other_ncl_items.append((code, name, balance, prior))

        if loan_items:
            ft.add_sub_heading("Financial Liabilities")

            secured_loans = [i for i in loan_items if "mortgage" in i[1].lower() or "secured" in i[1].lower()]
            unsecured_loans = [i for i in loan_items if "mortgage" not in i[1].lower() and "secured" not in i[1].lower()]

            if unsecured_loans:
                ft.add_sub_heading("Unsecured:", italic=True)
                for code, name, balance, prior in unsecured_loans:
                    val = abs(balance)
                    prior_val = abs(prior) if prior else Decimal("0")
                    total_ncl += val
                    total_ncl_prior += prior_val
                    ft.add_line(name, val, prior_val, indent=1)

            if secured_loans:
                ft.add_sub_heading("Secured:", italic=True)
                for code, name, balance, prior in secured_loans:
                    val = abs(balance)
                    prior_val = abs(prior) if prior else Decimal("0")
                    total_ncl += val
                    total_ncl_prior += prior_val
                    ft.add_line(name, val, prior_val, indent=1)

        if other_ncl_items:
            for code, name, balance, prior in other_ncl_items:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_ncl += val
                total_ncl_prior += prior_val
                ft.add_line(name, val, prior_val, indent=1)

        ft.add_subtotal("Total Non-Current Liabilities", total_ncl, total_ncl_prior, bold=True)

    # Total Liabilities
    total_liabilities = total_cl + total_ncl
    total_liabilities_prior = total_cl_prior + total_ncl_prior
    ft.add_spacer()
    ft.add_total("Total Liabilities", total_liabilities, total_liabilities_prior)

    # Net Assets + Equity (same table as liabilities for continuous flow)
    # Net Assets — grand total with double underline
    net_assets = total_assets - total_liabilities
    net_assets_prior = total_assets_prior - total_liabilities_prior
    ft.add_spacer()
    ft.add_total("Net Assets (Liabilities)", net_assets, net_assets_prior, is_grand_total=True)

    # ---- Equity (for non-sole-trader) ----
    if entity_type != "sole_trader":
        ft.add_section_heading("Equity", keep_with_next=True)

        total_equity = Decimal("0")
        total_equity_prior = Decimal("0")

        retained_note = nr.get("retained_profits") if nr else ""

        if sections["equity"]:
            equity_items = list(sections["equity"])
            for i, (code, name, balance, prior) in enumerate(equity_items):
                val = abs(balance) if balance < 0 else balance
                prior_val = abs(prior) if prior and prior < 0 else (prior or Decimal("0"))
                total_equity += val
                total_equity_prior += prior_val

                display_name = name
                if entity_type == "trust" and "retained" in name.lower():
                    display_name = "Undistributed income"

                # Add note ref to retained profits / undistributed income line
                line_note = ""
                name_lower = name.lower()
                if ("retained" in name_lower or "accumulated" in name_lower or
                    "undistributed" in name_lower):
                    line_note = retained_note

                # Keep all equity items together with Total Equity
                ft.add_line(display_name, val, prior_val,
                            keep_with_next=True, note_ref=line_note)
        else:
            if entity_type == "trust":
                label = "Undistributed income"
            else:
                label = "Retained profits / (accumulated losses)"
            ft.add_line(label, net_assets, net_assets_prior,
                        keep_with_next=True, note_ref=retained_note)
            total_equity = net_assets
            total_equity_prior = net_assets_prior

        ft.add_total("Total Equity", total_equity, total_equity_prior, is_grand_total=True)


# =============================================================================
# Summary P&L (Companies only)
# =============================================================================

def _add_summary_pnl(doc, entity, fy, sections, show_cents=False,
                     net_profit=Decimal("0"), net_profit_prior=Decimal("0"),
                     note_registry=None):
    """Add the Summary Profit and Loss Statement (companies only)."""
    nr = note_registry
    has_prior = _has_prior_year(fy)
    year = str(fy.end_date.year)
    prior_year_str = str(fy.end_date.year - 1) if has_prior else None

    _start_report_section(doc, entity,
                          f"Profit and Loss Statement\n{_get_period_text(fy)}",
                          footer_type="statement",
                          year=year, prior_year=prior_year_str,
                          has_prior=has_prior, show_column_headers=True,
                          include_note=True, show_cents=show_cents)

    ft = FinancialTable(doc, has_prior=has_prior, include_note=True, show_cents=show_cents)

    # Operating profit
    ft.add_line("Operating profit before income tax", net_profit, net_profit_prior)

    # Income tax (check for tax accounts in equity or expenses)
    tax_amount = Decimal("0")
    tax_amount_prior = Decimal("0")
    for code, name, balance, prior in sections["expenses"]:
        if "tax" in name.lower() and "income" in name.lower():
            tax_amount = abs(balance)
            tax_amount_prior = abs(prior) if prior else Decimal("0")

    if tax_amount > 0 or tax_amount_prior > 0:
        ft.add_line("Income tax attributable to operating profit (loss)",
                    -tax_amount, -tax_amount_prior)

    profit_after_tax = net_profit - tax_amount
    profit_after_tax_prior = net_profit_prior - tax_amount_prior

    ft.add_total("Operating profit after income tax", profit_after_tax,
                 profit_after_tax_prior)

    ft.add_spacer()

    # Retained profits
    opening_retained = Decimal("0")
    opening_retained_prior = Decimal("0")
    dividends = Decimal("0")
    dividends_prior = Decimal("0")

    for code, name, balance, prior in sections["equity"]:
        name_lower = name.lower()
        if "retained" in name_lower or "accumulated" in name_lower:
            opening_retained = abs(balance) if balance < 0 else balance
            opening_retained_prior = abs(prior) if prior and prior < 0 else (prior or Decimal("0"))
        elif "dividend" in name_lower:
            dividends = abs(balance) if balance else Decimal("0")
            dividends_prior = abs(prior) if prior else Decimal("0")

    ft.add_line("Retained profits at beginning of year",
                opening_retained - profit_after_tax,
                opening_retained_prior - profit_after_tax_prior)

    total_available = opening_retained
    total_available_prior = opening_retained_prior

    ft.add_subtotal("Total available for appropriation",
                    total_available, total_available_prior, bold=True)

    if dividends > 0 or dividends_prior > 0:
        ft.add_line("Dividends provided for or paid",
                    -dividends, -dividends_prior)

    closing_retained = total_available - dividends
    closing_retained_prior = total_available_prior - dividends_prior

    retained_note = nr.get("retained_profits") if nr else ""
    ft.add_total("Retained profits at end of year",
                 closing_retained, closing_retained_prior,
                 is_grand_total=True, note_ref=retained_note)


# =============================================================================
# Notes to Financial Statements
# =============================================================================

def _add_notes(doc, entity, fy, sections, show_cents=False, note_registry=None):
    """Add notes matching the real PDF format."""
    nr = note_registry
    _start_report_section(doc, entity,
                          f"Notes to the Financial Statements\n{_get_period_text(fy)}",
                          footer_type="notes",
                          show_column_headers=False)

    entity_type = entity.entity_type
    entity_ref_str = _entity_ref(entity_type)

    # ---- Note 1: Summary of Significant Accounting Policies ----
    note1_num = nr.get("accounting_policies") if nr else "1"
    _add_paragraph(doc, f"Note {note1_num}:  Summary of Significant Accounting Policies",
                   size=Pt(14), bold=True, space_after=12)

    # Basis of Preparation
    _add_paragraph(doc, "Basis of Preparation", size=FONT_SIZE_BODY, bold=True, space_after=6)

    if entity_type == "company":
        signatories = entity.officers.filter(is_signatory=True, date_ceased__isnull=True)
        singular = signatories.count() <= 1
        director_word = "director" if singular else "directors"
        has_have = "has" if singular else "have"

        _add_paragraph(
            doc,
            f"The {director_word} {has_have} prepared the financial statements on the basis that "
            f"the company is a non-reporting entity because there are no users dependent on general "
            f"purpose financial statements. The financial statements are therefore special purpose "
            f"financial statements that have been prepared in order to meet the needs of members.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    elif entity_type == "trust":
        _add_paragraph(
            doc,
            f"The trustee has prepared the financial statements of the trust on the basis that "
            f"the trust is a non-reporting entity because there are no users dependent on general "
            f"purpose financial statements. The financial statements are therefore special purpose "
            f"financial statements that have been prepared in order to meet the needs of the "
            f"trust deed and the directors of the trustee company.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    elif entity_type == "partnership":
        _add_paragraph(
            doc,
            f"The partners have prepared the financial statements on the basis that the partnership "
            f"is a non-reporting entity. The financial statements are therefore special purpose "
            f"financial statements that have been prepared in order to meet the needs of the partners.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    else:  # sole_trader
        _add_paragraph(
            doc,
            f"The owner has prepared the financial statements on the basis that the business "
            f"is a non-reporting entity because there are no users dependent on general purpose "
            f"financial statements. The financial statements are therefore special purpose "
            f"financial statements that have been prepared in order to meet the needs of the "
            f"owner and their bank.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    # Second paragraph
    responsible = _entity_label(entity_type)
    if entity_type == "company":
        has_have2 = "has" if singular else "have"
        is_are = "is" if singular else "are"
        _add_paragraph(
            doc,
            f"The financial statements have been prepared in accordance with the significant "
            f"accounting policies disclosed below, which {responsible} "
            f"{has_have2} determined {is_are} appropriate to meet "
            f"the needs of members. Such accounting policies are consistent with the previous period "
            f"unless stated otherwise.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    elif entity_type == "trust":
        _add_paragraph(
            doc,
            f"The financial statements have been prepared in accordance with the significant "
            f"accounting policies disclosed below, which the trustee has determined are appropriate "
            f"to meet the needs of the trust deed, the beneficiaries and the directors of the trustee "
            f"company. Such accounting policies are consistent with the previous period unless stated "
            f"otherwise.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    elif entity_type == "sole_trader":
        _add_paragraph(
            doc,
            f"The financial statements have been prepared in accordance with the significant "
            f"accounting policies disclosed below, which the owner has determined are appropriate "
            f"to meet the needs of the owner and their bank. Such accounting policies are consistent "
            f"with the previous period unless stated otherwise.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    else:
        _add_paragraph(
            doc,
            f"The financial statements have been prepared in accordance with the significant "
            f"accounting policies disclosed below, which {responsible} have determined are appropriate. "
            f"Such accounting policies are consistent with the previous period unless stated otherwise.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    _add_paragraph(
        doc,
        "The financial statements have been prepared on an accrual basis and are based on "
        "historical costs unless otherwise stated in the notes. The accounting policies that "
        "have been adopted in the preparation of the statements are as follows:",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    # Conditional accounting policies
    policy_letter = ord("a")

    # (a) Property, Plant and Equipment
    has_ppe = len(sections["noncurrent_assets"]) > 0
    if has_ppe:
        _add_paragraph(doc, f"({chr(policy_letter)})   Property, Plant and Equipment (PPE)",
                       size=FONT_SIZE_BODY, bold=True, space_after=6)
        _add_paragraph(
            doc,
            "All property, plant and equipment except for freehold land and buildings are initially "
            "measured at cost and are depreciated over their useful lives on a straight-line basis. "
            "Depreciation commences from the time the asset is available for its intended use. "
            "Leasehold improvements are depreciated over the shorter of either the unexpired period "
            "of the lease or the estimated useful lives of the improvements.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
            first_line_indent=Cm(1.5))
        _add_paragraph(
            doc,
            f"The carrying amount of plant and equipment is reviewed annually by {responsible} to "
            f"ensure it is not in excess of the recoverable amount from these assets. The recoverable "
            f"amount is assessed on the basis of the expected net cash flows that will be received "
            f"from the asset's employment and subsequent disposal. The expected net cash flows have "
            f"not been discounted in determining recoverable amounts.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
            first_line_indent=Cm(1.5))
        _add_paragraph(
            doc,
            f"Subsequent costs are included in the asset's carrying amount or recognised as a "
            f"separate asset, as appropriate, only when it is probable that future economic benefits "
            f"associated with the item will flow to {entity_ref_str} and the cost of the item can be "
            f"measured reliably. All other repairs and maintenance are recognised as expenses in "
            f"profit or loss during the financial period in which they are incurred.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
            first_line_indent=Cm(1.5))
        policy_letter += 1

    # Impairment of Assets
    if has_ppe:
        _add_paragraph(doc, f"({chr(policy_letter)})   Impairment of Assets",
                       size=FONT_SIZE_BODY, bold=True, space_after=6)
        _add_paragraph(
            doc,
            "At the end of each reporting period, property, plant and equipment, intangible assets "
            "and investments are reviewed to determine whether there is any indication that those "
            "assets have suffered an impairment loss. If there is an indication of possible "
            "impairment, the recoverable amount of any affected asset (or group of related assets) "
            "is estimated and compared with its carrying amount. The recoverable amount is the "
            "higher of the asset's fair value less costs of disposal and the present value of the "
            "asset's future cash flows discounted at the expected rate of return. If the estimated "
            "recoverable amount is lower, the carrying amount is reduced to the estimated "
            "recoverable amount and an impairment loss is recognised immediately in profit or loss.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
            first_line_indent=Cm(1.5))
        policy_letter += 1

    # Trade and Other Receivables (if receivables exist)
    has_receivables = any("debtor" in n.lower() or "receivable" in n.lower()
                          for _, n, _, _ in sections["current_assets"])
    if has_receivables:
        _add_paragraph(doc, f"({chr(policy_letter)})   Trade and Other Receivables",
                       size=FONT_SIZE_BODY, bold=True, space_after=6)
        _add_paragraph(
            doc,
            "Trade receivables are initially recognised at fair value and subsequently measured at "
            "amortised cost using the effective interest method, less any allowance for expected "
            "credit losses. Trade receivables are generally due for settlement within 30 days.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
            first_line_indent=Cm(1.5))
        policy_letter += 1

    # Cash and Cash Equivalents
    has_cash = any("cash" in n.lower() or "bank" in n.lower()
                   for _, n, _, _ in sections["current_assets"])
    if has_cash:
        _add_paragraph(doc, f"({chr(policy_letter)})   Cash and Cash Equivalents",
                       size=FONT_SIZE_BODY, bold=True, space_after=6)
        _add_paragraph(
            doc,
            "Cash and cash equivalents include cash on hand, deposits held at call with banks, "
            "other short-term highly liquid investments with original maturities of three months "
            "or less, and bank overdrafts. Bank overdrafts are shown within borrowings in current "
            "liabilities on the balance sheet.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
            first_line_indent=Cm(1.5))
        policy_letter += 1

    # Trade and Other Payables (if payables exist)
    has_payables = any("creditor" in n.lower() or "payable" in n.lower()
                       for _, n, _, _ in sections["current_liabilities"])
    if has_payables:
        _add_paragraph(doc, f"({chr(policy_letter)})   Trade and Other Payables",
                       size=FONT_SIZE_BODY, bold=True, space_after=6)
        _add_paragraph(
            doc,
            "Trade and other payables represent the liabilities for goods and services received "
            "by the entity that remain unpaid at the end of the reporting period. The balance is "
            "recognised as a current liability with the amounts normally paid within 30 days of "
            "recognition of the liability.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
            first_line_indent=Cm(1.5))
        policy_letter += 1

    # Revenue and Other Income
    _add_paragraph(doc, f"({chr(policy_letter)})   Revenue and Other Income",
                   size=FONT_SIZE_BODY, bold=True, space_after=6)
    _add_paragraph(
        doc,
        "Revenue is measured at the value of the consideration received or receivable after "
        "taking into account any trade discounts and volume rebates allowed. For this purpose, "
        "deferred consideration is not discounted to present values when recognising revenue.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
        first_line_indent=Cm(1.5))
    _add_paragraph(
        doc,
        "Interest revenue is recognised using the effective interest rate method, which, for "
        "floating rate financial assets, is the rate inherent in the instrument. Dividend revenue "
        "is recognised when the right to receive a dividend has been established.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
        first_line_indent=Cm(1.5))
    _add_paragraph(
        doc,
        "Revenue recognised related to the provision of services is determined with reference to "
        "the stage of completion of the transaction at the end of the reporting period and where "
        "outcome of the contract can be estimated reliably. Stage of completion is determined with "
        "reference to the services performed to date as a percentage of total anticipated services "
        "to be performed. Where the outcome cannot be estimated reliably, revenue is recognised "
        "only to the extent that related expenditure is recoverable.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
        first_line_indent=Cm(1.5))
    _add_paragraph(
        doc,
        "All revenue is stated net of the amount of goods and services tax (GST).",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
        first_line_indent=Cm(1.5))
    policy_letter += 1

    # Leases
    _add_paragraph(doc, f"({chr(policy_letter)})   Leases",
                   size=FONT_SIZE_BODY, bold=True, space_after=6)
    entity_name_ref = entity_ref_str.replace("the ", "")
    _add_paragraph(
        doc, f"The {entity_name_ref} as lessee",
        size=FONT_SIZE_BODY, bold=True, space_after=4, first_line_indent=Cm(1.5))
    _add_paragraph(
        doc,
        f"At inception of a contract, {entity_ref_str} assesses if the contract contains or is a lease "
        f"under AASB 16 Leases. Where a lease exists, a right-of-use asset and a corresponding "
        f"lease liability are recognised by {entity_ref_str} where {entity_ref_str} is a lessee. However, "
        f"all contracts that are classified as short-term leases (i.e. lease with remaining lease "
        f"term of 12 months or less) and leases of low value assets will be recognised as an "
        f"operating expense on a straight-line basis over the term of the lease.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
        first_line_indent=Cm(1.5))
    _add_paragraph(
        doc,
        f"{entity_ref_str.capitalize()} does not act as a lessor in relation to lease contracts.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
        first_line_indent=Cm(1.5))
    policy_letter += 1

    # Goods and Services Tax (GST)
    _add_paragraph(doc, f"({chr(policy_letter)})   Goods and Services Tax (GST)",
                   size=FONT_SIZE_BODY, bold=True, space_after=6)
    _add_paragraph(
        doc,
        "Revenues, expenses and assets are recognised net of the amount of GST, except where the "
        "amount of GST incurred is not recoverable from the Australian Taxation Office (ATO). In "
        "these circumstances, the GST is recognised as part of the cost of acquisition of the "
        "asset or as part of an item of the expense. Receivables and payables in the balance sheet "
        "are shown inclusive of GST.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
        first_line_indent=Cm(1.5))
    _add_paragraph(
        doc,
        "Cash flows are presented in the cash flow statement on a gross basis, except for the GST "
        "components of investing and financing activities, which are disclosed as operating cash flows.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10,
        first_line_indent=Cm(1.5))
    policy_letter += 1

    # ---- Note: Revenue ----
    note2_num = nr.get("revenue") if nr else "2"
    has_prior = _has_prior_year(fy)
    year = str(fy.end_date.year)
    prior_year_str = str(fy.end_date.year - 1) if has_prior else None
    has_trading = _has_cogs(sections)

    if not nr or nr.has("revenue"):
        _start_report_section(doc, entity,
                              f"Notes to the Financial Statements\n{_get_period_text(fy)}",
                              footer_type="notes",
                              show_column_headers=False)

        _add_paragraph(doc, f"Note {note2_num}:  Revenue", size=Pt(14), bold=True, space_after=8)
        _add_paragraph(doc, "Operating Activities:", size=FONT_SIZE_BODY, bold=True, space_after=4)

        ft_note2 = FinancialTable(doc, has_prior=has_prior, include_note=False, show_cents=show_cents)

        if has_trading:
            ft_note2.add_sub_heading("Sales revenue:", bold=False, space_before=2)
            total_revenue = Decimal("0")
            total_revenue_prior = Decimal("0")
            for code, name, balance, prior in sections["trading_income"]:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_revenue += val
                total_revenue_prior += prior_val
            ft_note2.add_line("Non-primary production trading revenue",
                              total_revenue, total_revenue_prior, indent=1)
        else:
            ft_note2.add_sub_heading("Other operating revenue:", bold=False, space_before=2)
            total_revenue = Decimal("0")
            total_revenue_prior = Decimal("0")
            for code, name, balance, prior in sections["trading_income"]:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_revenue += val
                total_revenue_prior += prior_val
                ft_note2.add_line(name, val, prior_val)

        # Other income
        if sections["income"]:
            ft_note2.add_spacer()
            ft_note2.add_sub_heading("Other revenue:", bold=False, space_before=2)
            total_other = Decimal("0")
            total_other_prior = Decimal("0")
            for code, name, balance, prior in sections["income"]:
                val = abs(balance)
                prior_val = abs(prior) if prior else Decimal("0")
                total_other += val
                total_other_prior += prior_val
                total_revenue += val
                total_revenue_prior += prior_val
                ft_note2.add_line(name, val, prior_val, indent=1)

        ft_note2.add_spacer()
        ft_note2.add_total("", total_revenue, total_revenue_prior, is_grand_total=True)

    # ---- Note: Profit from Ordinary Activities ----
    note3_num = nr.get("profit_ordinary") if nr else "3"
    if not nr or nr.has("profit_ordinary"):
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        _add_paragraph(doc, f"Note {note3_num}:  Profit from Ordinary Activities",
                       size=Pt(14), bold=True, space_after=6)
        _add_paragraph(
            doc,
            "Profit (loss) from ordinary activities before income tax has been determined after:",
            size=FONT_SIZE_BODY, space_after=8)

        _add_paragraph(doc, "Charging as Expense:", size=FONT_SIZE_BODY, bold=True, space_after=4)

        ft_note3 = FinancialTable(doc, has_prior=has_prior, include_note=False, show_cents=show_cents)

        # Check for borrowing costs
        borrowing_total = Decimal("0")
        borrowing_total_prior = Decimal("0")
        for code, name, balance, prior in sections["expenses"]:
            name_lower = name.lower()
            if "interest" in name_lower and ("loan" in name_lower or "australia" in name_lower or "mortgage" in name_lower):
                borrowing_total += abs(balance)
                borrowing_total_prior += abs(prior) if prior else Decimal("0")

        if borrowing_total > 0 or borrowing_total_prior > 0:
            ft_note3.add_sub_heading("Borrowing costs:", bold=False, space_before=2)
            ft_note3.add_line(" - Interest expense", borrowing_total, borrowing_total_prior, indent=1)

        # COGS
        if has_trading:
            total_cogs = Decimal("0")
            total_cogs_prior = Decimal("0")
            for code, name, balance, prior in sections["cogs"]:
                name_lower = name.lower()
                if "closing" not in name_lower:
                    total_cogs += abs(balance) if balance else Decimal("0")
                    total_cogs_prior += abs(prior) if prior else Decimal("0")
                else:
                    total_cogs -= abs(balance) if balance else Decimal("0")
                    total_cogs_prior -= abs(prior) if prior else Decimal("0")

            ft_note3.add_line("Cost of non-primary production goods traded",
                              total_cogs, total_cogs_prior)

        # Depreciation/amortisation
        depreciation_total = Decimal("0")
        depreciation_total_prior = Decimal("0")
        amortisation_total = Decimal("0")
        amortisation_total_prior = Decimal("0")

        for code, name, balance, prior in sections["expenses"]:
            name_lower = name.lower()
            val = abs(balance)
            prior_val = abs(prior) if prior else Decimal("0")
            if "depreciation" in name_lower:
                if "building" in name_lower:
                    ft_note3.add_sub_heading("Depreciation of non-current assets:", bold=False, space_before=2)
                    ft_note3.add_line(" - Buildings", val, prior_val, indent=1)
                    depreciation_total += val
                    depreciation_total_prior += prior_val
                else:
                    depreciation_total += val
                    depreciation_total_prior += prior_val
            if "amortisation" in name_lower or "amortization" in name_lower:
                amortisation_total += val
                amortisation_total_prior += prior_val

        if amortisation_total > 0 or amortisation_total_prior > 0:
            ft_note3.add_sub_heading("Amortisation of non-current assets:", bold=False, space_before=2)
            ft_note3.add_line(" - Leased assets", amortisation_total, amortisation_total_prior, indent=1)
            ft_note3.add_subtotal("Total amortisation expenses", amortisation_total,
                                  amortisation_total_prior)

        if depreciation_total > 0 or depreciation_total_prior > 0:
            has_building_dep = any("building" in n.lower() and "depreciation" in n.lower()
                                   for _, n, _, _ in sections["expenses"])
            if not has_building_dep:
                ft_note3.add_sub_heading("Depreciation of non-current assets:", bold=False, space_before=2)
            ft_note3.add_line(" - Other", depreciation_total, depreciation_total_prior, indent=1)
            ft_note3.add_subtotal("Total depreciation expenses", depreciation_total,
                                  depreciation_total_prior)

        # Bad debts
        bad_debts = Decimal("0")
        bad_debts_prior = Decimal("0")
        for code, name, balance, prior in sections["expenses"]:
            if "bad" in name.lower() and "debt" in name.lower():
                bad_debts += abs(balance)
                bad_debts_prior += abs(prior) if prior else Decimal("0")

        if bad_debts > 0 or bad_debts_prior > 0:
            ft_note3.add_line("Bad and doubtful debts", bad_debts, bad_debts_prior)

    # ---- Note: Retained Profits / Undistributed Income ----
    note4_num = nr.get("retained_profits") if nr else None
    if note4_num and nr and nr.has("retained_profits"):
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        if entity_type == "trust":
            note_title = f"Note {note4_num}:  Undistributed Income"
        else:
            note_title = f"Note {note4_num}:  Retained Profits"

        _add_paragraph(doc, note_title, size=Pt(14), bold=True, space_after=8)

        ft_note4 = FinancialTable(doc, has_prior=has_prior, include_note=False, show_cents=show_cents)

        # Calculate retained profits movement
        opening_retained = Decimal("0")
        opening_retained_prior = Decimal("0")
        dividends = Decimal("0")
        dividends_prior = Decimal("0")

        # Get net profit (need to recalculate from sections)
        total_income = Decimal("0")
        total_income_prior = Decimal("0")
        for code, name, balance, prior in sections["trading_income"]:
            total_income += abs(balance)
            total_income_prior += abs(prior) if prior else Decimal("0")
        for code, name, balance, prior in sections["income"]:
            total_income += abs(balance)
            total_income_prior += abs(prior) if prior else Decimal("0")

        total_expenses = Decimal("0")
        total_expenses_prior = Decimal("0")
        for code, name, balance, prior in sections["expenses"]:
            total_expenses += abs(balance)
            total_expenses_prior += abs(prior) if prior else Decimal("0")

        total_cogs_note = Decimal("0")
        total_cogs_note_prior = Decimal("0")
        if has_trading:
            for code, name, balance, prior in sections["cogs"]:
                name_lower = name.lower()
                if "closing" not in name_lower:
                    total_cogs_note += abs(balance) if balance else Decimal("0")
                    total_cogs_note_prior += abs(prior) if prior else Decimal("0")
                else:
                    total_cogs_note -= abs(balance) if balance else Decimal("0")
                    total_cogs_note_prior -= abs(prior) if prior else Decimal("0")

        net_profit_note = total_income - total_expenses - total_cogs_note
        net_profit_note_prior = total_income_prior - total_expenses_prior - total_cogs_note_prior

        # Get equity data
        for code, name, balance, prior in sections["equity"]:
            name_lower = name.lower()
            if "retained" in name_lower or "accumulated" in name_lower or "undistributed" in name_lower:
                opening_retained = abs(balance) if balance < 0 else balance
                opening_retained_prior = abs(prior) if prior and prior < 0 else (prior or Decimal("0"))
            elif "dividend" in name_lower:
                dividends = abs(balance) if balance else Decimal("0")
                dividends_prior = abs(prior) if prior else Decimal("0")

        # Opening balance = closing - profit + dividends
        opening_balance = opening_retained - net_profit_note
        opening_balance_prior = opening_retained_prior - net_profit_note_prior

        if entity_type == "trust":
            ft_note4.add_line("Undistributed income at beginning of year",
                              opening_balance, opening_balance_prior)
            ft_note4.add_line("Net profit / (loss) attributable to the trust",
                              net_profit_note, net_profit_note_prior)
        else:
            ft_note4.add_line("Retained profits at beginning of year",
                              opening_balance, opening_balance_prior)
            ft_note4.add_line("Net profit / (loss) attributable to members",
                              net_profit_note, net_profit_note_prior)

        if dividends > 0 or dividends_prior > 0:
            ft_note4.add_line("Dividends provided for or paid",
                              -dividends, -dividends_prior)

        if entity_type == "trust":
            ft_note4.add_total("Undistributed income at end of year",
                               opening_retained, opening_retained_prior, is_grand_total=True)
        else:
            ft_note4.add_total("Retained profits at end of year",
                               opening_retained, opening_retained_prior, is_grand_total=True)


# =============================================================================
# Depreciation Schedule
# =============================================================================

def _add_depreciation_schedule(doc, entity, fy, show_cents=False):
    """
    Add the depreciation schedule in LANDSCAPE orientation.
    Assets are grouped by category with subtotals per category.
    Columns: Asset | Total | Priv% | OWDV | Disposal(Date,Consid) | Addition(Date,Cost) |
             Value | T | Rate | Deprec | Priv | CWDV | Profit(Upto+,Above) | Loss(Total-,Priv)
    """
    assets = DepreciationAsset.objects.filter(
        financial_year=fy
    ).order_by("category", "display_order", "asset_name")

    if not assets.exists():
        return

    # Group assets by category
    categories = OrderedDict()
    for asset in assets:
        if asset.category not in categories:
            categories[asset.category] = []
        categories[asset.category].append(asset)

    def _dep_fmt(val):
        """Format a decimal value for the depreciation schedule."""
        if val is None or val == 0:
            return ""
        if show_cents:
            return f"{val:,.2f}"
        return f"{val:,.0f}"

    def _fmt_rate(val):
        if val is None or val == 0:
            return "0.00"
        return f"{val:.2f}"

    def _fmt_date(d):
        if d is None:
            return ""
        return d.strftime("%d/%m/%y")

    # Column widths in cm for the 14-column depreciation table
    # Asset | Total | Priv% | OWDV | Date | Consid | Date | Cost | Value | T | Rate | Deprec | Priv | CWDV
    _DEP_COL_WIDTHS_CM = [5.0, 1.7, 0.9, 1.7, 1.4, 1.4, 1.4, 1.4, 1.4, 0.5, 1.1, 1.7, 1.1, 1.7]

    for cat_name, cat_assets in categories.items():
        # New landscape section for each depreciation category
        _start_report_section(doc, entity,
                              f"Depreciation Schedule\n{_get_period_text(fy)}",
                              footer_type="statement",
                              show_column_headers=False, landscape=True)

        # Reduce margins for landscape depreciation schedule to give more room
        current_section = doc.sections[-1]
        current_section.left_margin = Cm(1.0)
        current_section.right_margin = Cm(1.0)

        _add_paragraph(doc, cat_name, size=FONT_SIZE_BODY, bold=True,
                       underline=True, space_after=6)

        # Create table with headers
        col_headers = [
            "Asset", "Total", "Priv\n%", "OWDV",
            "Date", "Consid",  # Disposal
            "Date", "Cost",    # Addition
            "Value", "T", "Rate", "Deprec", "Priv", "CWDV",
        ]

        num_cols = len(col_headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Force fixed table layout so Word respects our column widths exactly
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl_layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tbl_pr.append(tbl_layout)

        # Set explicit column widths to prevent text wrapping
        for row in table.rows:
            for i, width_cm in enumerate(_DEP_COL_WIDTHS_CM):
                row.cells[i].width = Cm(width_cm)

        # Style header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(col_headers):
            cell = hdr_cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.font.size = Pt(7)
            run.font.name = FONT_NAME
            run.font.bold = True
            # Shade header
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Add group header rows (DISPOSAL / ADDITION)
        # We'll add a second header row for the group labels

        # Add asset rows
        cat_total_cost = Decimal("0")
        cat_owdv = Decimal("0")
        cat_deprec = Decimal("0")
        cat_priv_dep = Decimal("0")
        cat_cwdv = Decimal("0")
        cat_add_cost = Decimal("0")
        cat_disp_consid = Decimal("0")

        for asset in cat_assets:
            row_cells = table.add_row().cells
            values = [
                asset.asset_name,
                _dep_fmt(asset.total_cost),
                f"{asset.private_use_pct:.2f}" if asset.private_use_pct else "",
                _dep_fmt(asset.opening_wdv),
                _fmt_date(asset.disposal_date),
                _dep_fmt(asset.disposal_consideration),
                _fmt_date(asset.addition_date),
                _dep_fmt(asset.addition_cost),
                _dep_fmt(asset.depreciable_value),
                asset.get_method_display()[0] if asset.method else "",
                _fmt_rate(asset.rate),
                _dep_fmt(asset.depreciation_amount),
                _dep_fmt(asset.private_depreciation),
                _dep_fmt(asset.closing_wdv),
            ]

            for i, val in enumerate(values):
                cell = row_cells[i]
                cell.width = Cm(_DEP_COL_WIDTHS_CM[i])
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(val))
                run.font.size = Pt(7)
                run.font.name = FONT_NAME

            # Accumulate category totals
            cat_total_cost += asset.total_cost or Decimal("0")
            cat_owdv += asset.opening_wdv or Decimal("0")
            cat_deprec += asset.depreciation_amount or Decimal("0")
            cat_priv_dep += asset.private_depreciation or Decimal("0")
            cat_cwdv += asset.closing_wdv or Decimal("0")
            cat_add_cost += asset.addition_cost or Decimal("0")
            cat_disp_consid += asset.disposal_consideration or Decimal("0")

        # Subtotals row
        sub_row = table.add_row().cells
        sub_values = [
            "Subtotals",
            _dep_fmt(cat_total_cost), "", _dep_fmt(cat_owdv),
            "", _dep_fmt(cat_disp_consid),
            "", _dep_fmt(cat_add_cost),
            "", "", "",
            _dep_fmt(cat_deprec), _dep_fmt(cat_priv_dep), _dep_fmt(cat_cwdv),
        ]
        for i, val in enumerate(sub_values):
            cell = sub_row[i]
            cell.width = Cm(_DEP_COL_WIDTHS_CM[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(7)
            run.font.name = FONT_NAME
            run.font.bold = True

        # Net depreciation line
        net_dep = cat_deprec - cat_priv_dep
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        _add_paragraph(doc, f"Deduct Private Portion: {_dep_fmt(cat_priv_dep)}",
                       size=Pt(8), space_after=2)
        p = doc.add_paragraph()
        run = p.add_run(f"Net Depreciation: {_dep_fmt(net_dep)}")
        run.font.size = Pt(8)
        run.font.name = FONT_NAME
        run.font.bold = True
        run.font.underline = True
        p.paragraph_format.space_after = Pt(6)

    # Note: No need to return to portrait here — the next report's
    # _start_report_section() call will create a new portrait section automatically.


# =============================================================================
# Partners' Profit Distribution Summary
# =============================================================================

def _add_partners_distribution(doc, entity, fy, sections, show_cents=False,
                               net_profit=Decimal("0"), net_profit_prior=Decimal("0")):
    """Add the partners' profit distribution summary (partnership only)."""
    has_prior = _has_prior_year(fy)
    year = str(fy.end_date.year)

    _start_report_section(doc, entity,
                          f"Partners' Profit Distribution Summary\n{_get_period_text(fy)}",
                          footer_type="statement",
                          year=year, has_prior=False,
                          show_column_headers=True,
                          include_note=False, show_cents=show_cents)

    _add_paragraph(doc, "Partners' Share of Profit", size=FONT_SIZE_BODY, bold=True, space_after=6)

    partners = entity.officers.filter(
        role=EntityOfficer.OfficerRole.PARTNER,
        date_ceased__isnull=True,
    ).order_by("display_order")

    for partner in partners:
        share_pct = partner.profit_share_percentage or Decimal("0")
        share_amount = net_profit * share_pct / Decimal("100") if share_pct else Decimal("0")
        _add_amount_line(doc, f"{partner.full_name} ({share_pct}%)",
                         share_amount, has_prior=False, indent=1, show_cents=show_cents)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _add_amount_line(doc, "Total Profit Distributed", net_profit, has_prior=False, bold=True,
                     show_cents=show_cents)


# =============================================================================
# Declaration
# =============================================================================

def _add_declaration(doc, entity, fy):
    """Add the declaration page — always starts on a new page for signing."""
    entity_type = entity.entity_type
    signatories = entity.officers.filter(
        is_signatory=True,
        date_ceased__isnull=True,
    ).order_by("display_order")

    num_signatories = signatories.count()
    singular = num_signatories <= 1

    if entity_type == "company":
        title = "Director's Declaration" if singular else "Directors' Declaration"
        _start_report_section(doc, entity, title,
                              footer_type="none", show_column_headers=False)

        director_word = "director" if singular else "directors"
        has_have = "has" if singular else "have"

        _add_paragraph(
            doc,
            f"The {director_word} {has_have} determined that the company is not a reporting entity "
            f"and that this special purpose financial report should be prepared in accordance with "
            f"the accounting policies prescribed in Note 1 to the financial statements.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

        _add_paragraph(
            doc,
            f"The {director_word} of the company declare{'s' if singular else ''} that:",
            size=FONT_SIZE_BODY, space_after=6)

        _add_paragraph(
            doc,
            f"1.  the financial statements and notes, present fairly the company's financial "
            f"position as at {fy.end_date.strftime('%-d %B %Y')} and its performance for the {_get_period_label(fy)} "
            f"ended on that date in accordance with the accounting policies described in Note 1 "
            f"to the financial statements;",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

        _add_paragraph(
            doc,
            f"2.  in the {director_word}'s opinion, there are reasonable grounds to believe that "
            f"the company will be able to pay its debts as and when they become due and payable.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

        _add_paragraph(
            doc,
            f"This declaration is made in accordance with a resolution of the {director_word}.",
            size=FONT_SIZE_BODY, space_after=20)

        # Signature blocks
        for officer in signatories:
            doc.add_paragraph().paragraph_format.space_after = Pt(20)
            _add_paragraph(doc, "_" * 50, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, officer.full_name, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, "Director", size=FONT_SIZE_BODY, space_after=6)

    elif entity_type == "trust":
        _start_report_section(doc, entity, "Trustee's Declaration",
                              footer_type="none", show_column_headers=False)

        _add_paragraph(
            doc,
            f"The trustee declares that the trust is not a reporting entity and that this special "
            f"purpose financial report should be prepared in accordance with the accounting policies "
            f"prescribed in Note 1 to the financial statements.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

        _add_paragraph(
            doc,
            "The directors of the trustee company declare that:",
            size=FONT_SIZE_BODY, space_after=6)

        _add_paragraph(
            doc,
            f"(i)  the financial statements and notes present fairly the trust's financial "
            f"position as at {fy.end_date.strftime('%-d %B %Y')} and its performance for the {_get_period_label(fy)} "
            f"ended on that date in accordance with the accounting policies described in Note 1 "
            f"to the financial statements;",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

        _add_paragraph(
            doc,
            f"(ii)  in the directors' opinion, there are reasonable grounds to believe that the "
            f"trust will be able to pay its debts as and when they become due and payable.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

        _add_paragraph(
            doc,
            "Signed in accordance with a resolution of the trustee by:",
            size=FONT_SIZE_BODY, space_after=20)

        for officer in signatories:
            doc.add_paragraph().paragraph_format.space_after = Pt(20)
            _add_paragraph(doc, "_" * 50, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, f"{officer.full_name}", size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, "(Trustee)", size=FONT_SIZE_BODY, space_after=6)

    elif entity_type == "partnership":
        _start_report_section(doc, entity, "Partner Declaration",
                              footer_type="none", show_column_headers=False)

        _add_paragraph(
            doc,
            "The partners have determined that the partnership is not a reporting entity and that "
            "this special purpose financial report should be prepared in accordance with the "
            "accounting policies described in the financial statements.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

        _add_paragraph(doc, "The partners declare that:", size=FONT_SIZE_BODY, space_after=6)

        _add_paragraph(
            doc,
            f"(a) the financial statements comply with the accounting policies described therein; and",
            size=FONT_SIZE_BODY, space_after=6)

        _add_paragraph(
            doc,
            f"(b) the financial statements present fairly the partnership's financial position as at "
            f"{fy.end_date.strftime('%-d %B %Y')} and its performance for the {_get_period_label(fy)} ended on that date.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

        _add_paragraph(
            doc,
            "In the partners' opinion, there are reasonable grounds to believe that the partnership "
            "will be able to pay its debts as and when they become due and payable.",
            size=FONT_SIZE_BODY, space_after=20)

        for officer in signatories:
            doc.add_paragraph().paragraph_format.space_after = Pt(20)
            _add_paragraph(doc, "_" * 50, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, officer.full_name, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, "Partner", size=FONT_SIZE_BODY, space_after=6)

    else:  # sole_trader
        _start_report_section(doc, entity, "Proprietor Declaration",
                              footer_type="none", show_column_headers=False)

        _add_paragraph(
            doc,
            "The proprietor has determined that the entity is not a reporting entity and that "
            "this special purpose financial statement should be prepared in accordance with the "
            "accounting policies described in the financial statements.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

        _add_paragraph(
            doc,
            "The proprietor declares that:",
            size=FONT_SIZE_BODY, space_after=6)

        _add_paragraph(
            doc,
            f"1.  the financial statements present fairly the business's financial position as at "
            f"{fy.end_date.strftime('%-d %B %Y')} and its performance for the {_get_period_label(fy)} ended on that date "
            f"in accordance with the accounting policies described in the financial statements;",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

        _add_paragraph(
            doc,
            f"2.  in the proprietor's opinion, there are reasonable grounds to believe that the "
            f"business will be able to pay its debts as and when they become due and payable.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=20)

        for officer in signatories:
            doc.add_paragraph().paragraph_format.space_after = Pt(20)
            _add_paragraph(doc, "_" * 50, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, officer.full_name, size=FONT_SIZE_BODY, space_after=0)
            _add_paragraph(doc, "Proprietor", size=FONT_SIZE_BODY, space_after=6)

    _add_paragraph(doc, "Dated:", size=FONT_SIZE_BODY, space_after=2)


# =============================================================================
# Compilation Report
# =============================================================================

def _add_compilation_report(doc, entity, fy):
    """Add the compilation report (APES 315)."""
    _start_report_section(doc, entity,
                          f"Compilation Report to {entity.entity_name}",
                          footer_type="none", show_column_headers=False)

    entity_type = entity.entity_type
    end_date_str = fy.end_date.strftime('%-d %B %Y')

    # Opening paragraph
    _add_paragraph(
        doc,
        f"We have compiled the accompanying special purpose financial statements of "
        f"{entity.entity_name}, which comprise the balance sheet as at {end_date_str}, "
        f"the Statement of Profit and Loss for the {_get_period_label(fy)} then ended, a summary of significant "
        f"accounting policies and other explanatory notes. The specific purpose for which the "
        f"special purpose financial statements have been prepared is set out in Note 1 to the "
        f"financial statements.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    # The Responsibility section
    if entity_type == "company":
        signatories = entity.officers.filter(is_signatory=True, date_ceased__isnull=True)
        singular = signatories.count() <= 1
        director_word = "Director" if singular else "Directors"
        director_lower = "director" if singular else "directors"

        _add_paragraph(doc, f"The Responsibility of the {director_word}",
                       size=FONT_SIZE_BODY, italic=True, space_after=4)
        _add_paragraph(
            doc,
            f"The {director_lower} of {entity.entity_name} is solely responsible for the information "
            f"contained in the special purpose financial statements, the reliability, accuracy "
            f"and completeness of the information and for the determination that the significant "
            f"accounting policies used are appropriate to meet the needs and for the purpose that "
            f"the financial statements were prepared.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    elif entity_type == "trust":
        _add_paragraph(doc, "The Responsibility of the Trustee",
                       size=FONT_SIZE_BODY, italic=True, space_after=4)
        _add_paragraph(
            doc,
            f"The directors of the trustee company are solely responsible for the information "
            f"contained in the special purpose financial statements, the reliability, accuracy "
            f"and completeness of the information and for the determination that the significant "
            f"accounting policies used are appropriate to meet the needs of the trust deed and "
            f"the directors of the trustee company.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    elif entity_type == "partnership":
        _add_paragraph(doc, "The Responsibility of the Partners",
                       size=FONT_SIZE_BODY, italic=True, space_after=4)
        _add_paragraph(
            doc,
            f"The partners are solely responsible for the information contained in the special "
            f"purpose financial statements, the reliability, accuracy and completeness of the "
            f"information and for the determination that the significant accounting policies used "
            f"are appropriate to meet the needs and for the purpose that the financial statements "
            f"were prepared.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    else:  # sole_trader
        _add_paragraph(doc, "The Responsibility of the Owner",
                       size=FONT_SIZE_BODY, italic=True, space_after=4)
        _add_paragraph(
            doc,
            f"The owner of {entity.entity_name} is solely responsible for the information "
            f"contained in the special purpose financial statements, the reliability, accuracy "
            f"and completeness of the information and for the determination that the significant "
            f"accounting policies used are appropriate to meet the needs of the owner and their bank.",
            size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    # Our Responsibility
    responsible = _entity_label(entity_type)
    _add_paragraph(doc, "Our Responsibility",
                   size=FONT_SIZE_BODY, italic=True, space_after=4)

    _add_paragraph(
        doc,
        f"On the basis of information provided by {responsible}, we have compiled the "
        f"accompanying special purpose financial statements in accordance with the significant "
        f"accounting policies as described in Note 1 to the financial statements and APES 315 "
        f"Compilation of Financial Information.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    _add_paragraph(
        doc,
        "We have applied our expertise in accounting and financial reporting to compile these "
        "financial statements in accordance with the significant accounting policies described "
        "in Note 1 to the financial statements. We have complied with the relevant ethical "
        "requirements of APES 110 Code of Ethics for Professional Accountants (including "
        "Independence Standards).",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

    # Assurance Disclaimer
    _add_paragraph(doc, "Assurance Disclaimer",
                   size=FONT_SIZE_BODY, italic=True, space_after=4)

    _add_paragraph(
        doc,
        "Since a compilation engagement is not an assurance engagement, we are not required to "
        "verify the reliability, accuracy or completeness of the information provided to us by "
        "management to compile these financial statements. Accordingly, we do not express an "
        "audit opinion or a review conclusion on these financial statements.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    _add_paragraph(
        doc,
        f"The special purpose financial statements were compiled exclusively for the benefit of "
        f"{responsible} who is responsible for the reliability, accuracy and completeness of the "
        f"information used to compile them. Accordingly, these special purpose financial statements "
        f"may not be suitable for other purposes. We do not accept responsibility for the contents "
        f"of the special purpose financial statements.",
        size=FONT_SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=20)

    # Signature block
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    _add_paragraph(doc, "_" * 40, size=FONT_SIZE_BODY, space_after=0)
    _add_paragraph(doc, FIRM_NAME, size=FONT_SIZE_BODY, space_after=0)
    _add_paragraph(doc, FIRM_ADDRESS_1, size=FONT_SIZE_BODY, space_after=0)
    _add_paragraph(doc, FIRM_ADDRESS_2, size=FONT_SIZE_BODY, space_after=6)
    _add_paragraph(doc, f"{date.today().strftime('%-d %B, %Y')}", size=FONT_SIZE_BODY, space_after=2)


# =============================================================================
# Main Generation Function
# =============================================================================

def generate_financial_statements(financial_year_id) -> io.BytesIO:
    """
    Generate a complete set of financial statements for a financial year.
    Returns a BytesIO object containing the Word document.
    """
    fy = FinancialYear.objects.select_related(
        "entity", "entity__client", "prior_year"
    ).get(pk=financial_year_id)

    entity = fy.entity
    entity_type = entity.entity_type
    show_cents = entity.show_cents

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Extract trial balance data
    sections = _get_tb_sections(fy)
    has_trading = _has_cogs(sections)

    # Build note registry — automatically assigns note numbers based on data
    note_registry = _build_note_registry(entity, sections)

    # =========================================================================
    # Build document in entity-type-specific order
    # Each entity type has a different section ordering based on the real
    # Access Ledger PDF output.
    # =========================================================================

    # Common: Cover + Contents
    _add_cover_page(doc, entity, fy)
    _add_contents_page(doc, entity, fy, sections)

    # Common: Trading Account (if COGS exist)
    gross_profit = None
    gross_profit_prior = None
    if has_trading:
        # Company with trading: Compilation Report comes FIRST
        if entity_type == "company":
            _add_compilation_report(doc, entity, fy)
        gross_profit, gross_profit_prior = _add_trading_account(
            doc, entity, fy, sections, show_cents=show_cents)

    # Common: Detailed P&L
    net_profit, net_profit_prior = _add_detailed_pnl(
        doc, entity, fy, sections, show_cents=show_cents,
        gross_profit=gross_profit, gross_profit_prior=gross_profit_prior,
        note_registry=note_registry)

    # Common: Detailed Balance Sheet
    _add_detailed_balance_sheet(doc, entity, fy, sections, show_cents=show_cents,
                                net_profit=net_profit, net_profit_prior=net_profit_prior,
                                note_registry=note_registry)

    # ---- Entity-type-specific ordering from here ----

    if entity_type == "company":
        # Company order: Summary P&L > Depreciation > Notes > Declaration > [Compilation if simple]
        if has_trading:
            _add_summary_pnl(doc, entity, fy, sections, show_cents=show_cents,
                             net_profit=net_profit, net_profit_prior=net_profit_prior,
                             note_registry=note_registry)
        _add_depreciation_schedule(doc, entity, fy, show_cents=show_cents)
        _add_notes(doc, entity, fy, sections, show_cents=show_cents,
                   note_registry=note_registry)
        _add_declaration(doc, entity, fy)
        if not has_trading:
            # Simple company: compilation report LAST
            _add_compilation_report(doc, entity, fy)

    elif entity_type == "trust":
        # Trust order: Notes > Depreciation > Trustee's Declaration > Compilation Report
        _add_notes(doc, entity, fy, sections, show_cents=show_cents,
                   note_registry=note_registry)
        _add_depreciation_schedule(doc, entity, fy, show_cents=show_cents)
        _add_declaration(doc, entity, fy)
        _add_compilation_report(doc, entity, fy)

    elif entity_type == "partnership":
        # Partnership order: Distribution > Depreciation > Notes > Declaration > Compilation
        _add_partners_distribution(doc, entity, fy, sections, show_cents=show_cents,
                                   net_profit=net_profit, net_profit_prior=net_profit_prior)
        _add_depreciation_schedule(doc, entity, fy, show_cents=show_cents)
        _add_notes(doc, entity, fy, sections, show_cents=show_cents,
                   note_registry=note_registry)
        _add_declaration(doc, entity, fy)
        _add_compilation_report(doc, entity, fy)

    else:  # sole_trader
        # Sole trader order: Notes > Depreciation > Compilation > Declaration
        _add_notes(doc, entity, fy, sections, show_cents=show_cents,
                   note_registry=note_registry)
        _add_depreciation_schedule(doc, entity, fy, show_cents=show_cents)
        _add_compilation_report(doc, entity, fy)
        _add_declaration(doc, entity, fy)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

