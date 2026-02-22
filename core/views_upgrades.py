"""
MCS Platform - Upgrade Views
Implements views for:
  1. Prior Year Comparatives Engine
  2. Document Version Control & Regeneration
  3. Trust Distribution Workflow
  4. Partnership Profit Allocation Module
  5. Working Paper Notes per Account
  6. Bulk Entity Import / Onboarding Tool
"""
import io
import csv
import json
import logging
import openpyxl
from collections import OrderedDict
from decimal import Decimal, InvalidOperation
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.exceptions import PermissionDenied
from django.core.files.base import ContentFile
from django.db import transaction
from django.db.models import Sum, Q
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.views.decorators.http import require_POST

from .models import (
    Entity, FinancialYear, TrialBalanceLine, AccountMapping,
    GeneratedDocument, AuditLog, EntityOfficer,
    TrustDistribution, BeneficiaryAllocation,
    PartnershipAllocation, PartnerShare, PartnerCapitalAccount,
    WorkpaperNote, EntityImportJob,
)
from config.authorization import get_financial_year_for_user

logger = logging.getLogger(__name__)


def _log_action(request, action, description, obj=None):
    AuditLog.objects.create(
        user=request.user,
        action=action,
        description=description,
        affected_object_type=type(obj).__name__ if obj else "",
        affected_object_id=str(obj.pk) if obj else "",
        ip_address=request.META.get("REMOTE_ADDR"),
    )


# ============================================================================
# 1. PRIOR YEAR COMPARATIVES ENGINE
# ============================================================================

@login_required
@require_POST
def populate_comparatives(request, pk):
    """
    Auto-populate prior year balances from the previous finalised financial year.
    Matches by account_code. Detects reclassifications where mapping changed.
    """
    fy = get_financial_year_for_user(request, pk)

    if fy.is_locked:
        messages.error(request, "Cannot modify comparatives on a finalised year.")
        return redirect("core:financial_year_detail", pk=pk)

    prior_fy = fy.prior_year
    if not prior_fy:
        # Try to auto-detect prior year
        prior_fy = (
            FinancialYear.objects
            .filter(entity=fy.entity, end_date__lt=fy.start_date)
            .order_by("-end_date")
            .first()
        )
        if prior_fy:
            fy.prior_year = prior_fy
            fy.save(update_fields=["prior_year"])

    if not prior_fy:
        messages.warning(
            request,
            "No prior year found. Use manual override to enter prior year comparatives."
        )
        return redirect("core:financial_year_detail", pk=pk)

    # Build lookup from prior year TB lines
    prior_lines = {
        line.account_code: line
        for line in prior_fy.trial_balance_lines.select_related("mapped_line_item").all()
    }

    updated = 0
    reclassified = 0
    current_lines = fy.trial_balance_lines.select_related("mapped_line_item").all()

    with transaction.atomic():
        for line in current_lines:
            if line.comparatives_locked:
                continue
            prior = prior_lines.get(line.account_code)
            if prior:
                line.prior_debit = prior.debit
                line.prior_credit = prior.credit
                line.prior_closing_balance = prior.closing_balance
                line.prior_balance_override = False

                # Check for reclassification
                if (prior.mapped_line_item and line.mapped_line_item and
                        prior.mapped_line_item_id != line.mapped_line_item_id):
                    line.reclassified = True
                    line.prior_mapped_line_item = prior.mapped_line_item
                    reclassified += 1
                else:
                    line.reclassified = False
                    line.prior_mapped_line_item = prior.mapped_line_item

                line.save(update_fields=[
                    "prior_debit", "prior_credit", "prior_closing_balance",
                    "prior_balance_override", "reclassified", "prior_mapped_line_item",
                ])
                updated += 1

    _log_action(
        request, "import",
        f"Populated comparatives from {prior_fy.year_label}: {updated} lines updated, {reclassified} reclassified",
        fy,
    )
    msg = f"Comparatives populated from {prior_fy.year_label}: {updated} lines updated."
    if reclassified:
        msg += f" {reclassified} accounts reclassified (note disclosure triggered)."
    messages.success(request, msg)
    return redirect("core:financial_year_detail", pk=pk)


@login_required
@require_POST
def override_comparative(request, pk):
    """
    Manually override a prior year comparative for a specific TB line.
    Used for Year 1 onboarding when prior year was prepared outside StatementHub.
    """
    line = get_object_or_404(TrialBalanceLine, pk=pk)
    fy = line.financial_year

    # Authorization: verify user has access to this financial year's entity
    get_financial_year_for_user(request, fy.pk)

    if fy.is_locked or line.comparatives_locked:
        return JsonResponse({"error": "Comparatives are locked."}, status=403)

    try:
        prior_debit = Decimal(request.POST.get("prior_debit", "0"))
        prior_credit = Decimal(request.POST.get("prior_credit", "0"))
    except (InvalidOperation, TypeError):
        return JsonResponse({"error": "Invalid amounts."}, status=400)

    line.prior_debit = prior_debit
    line.prior_credit = prior_credit
    line.prior_closing_balance = prior_debit - prior_credit
    line.prior_balance_override = True
    line.save(update_fields=[
        "prior_debit", "prior_credit", "prior_closing_balance", "prior_balance_override",
    ])

    _log_action(
        request, "adjustment",
        f"Manual comparative override: {line.account_code} Dr={prior_debit} Cr={prior_credit}",
        fy,
    )
    return JsonResponse({
        "success": True,
        "prior_debit": str(line.prior_debit),
        "prior_credit": str(line.prior_credit),
        "variance_amount": str(line.variance_amount),
        "variance_pct": str(line.variance_percentage) if line.variance_percentage is not None else "N/A",
    })


@login_required
@require_POST
def lock_comparatives(request, pk):
    """Lock all comparatives for a financial year (called on finalisation)."""
    fy = get_financial_year_for_user(request, pk)
    fy.trial_balance_lines.update(comparatives_locked=True)
    messages.success(request, "Prior year comparatives locked.")
    return redirect("core:financial_year_detail", pk=pk)


# ============================================================================
# 2. DOCUMENT VERSION CONTROL & REGENERATION
# ============================================================================

@login_required
@require_POST
def regenerate_document(request, pk):
    """
    Regenerate a document, creating a new version and preserving the old one.
    """
    fy = get_financial_year_for_user(request, pk)
    doc_type = request.POST.get("type", "financial_statements")
    fmt = request.POST.get("format", "docx").lower()

    if not fy.trial_balance_lines.exists() and doc_type == "financial_statements":
        messages.error(request, "Cannot generate statements: no trial balance data loaded.")
        return redirect("core:financial_year_detail", pk=pk)

    # Find the latest version of this document type
    latest = (
        GeneratedDocument.objects
        .filter(financial_year=fy, document_type=doc_type)
        .order_by("-version")
        .first()
    )
    new_version = (latest.version + 1) if latest else 1

    # Generate the document based on type
    if doc_type == "financial_statements":
        from .docgen import generate_financial_statements
        has_open_risks = fy.risk_flags.filter(status="open").exists()
        try:
            buffer = generate_financial_statements(fy.pk, has_open_risks=has_open_risks)
        except Exception as e:
            messages.error(request, f"Document generation failed: {e}")
            return redirect("core:financial_year_detail", pk=pk)
        file_content = buffer.getvalue()
        entity_name = fy.entity.entity_name.replace(" ", "_")
        filename = f"{entity_name}_Financial_Statements_{fy.year_label}_v{new_version}.{fmt}"
    elif doc_type == "distribution_minutes":
        from .distmin_gen import generate_distribution_minutes as gen_distmin
        try:
            buffer = gen_distmin(fy.pk)
        except Exception as e:
            messages.error(request, f"Distribution minutes generation failed: {e}")
            return redirect("core:financial_year_detail", pk=pk)
        file_content = buffer.getvalue()
        entity_name = fy.entity.entity_name.replace(" ", "_")
        filename = f"{entity_name}_Distribution_Minutes_{fy.year_label}_v{new_version}.docx"
        fmt = "docx"
    else:
        messages.error(request, f"Unknown document type: {doc_type}")
        return redirect("core:financial_year_detail", pk=pk)

    # Build change summary
    change_summary = ""
    if latest:
        # Mark old version as superseded
        change_summary = f"Regenerated from v{latest.version}."
        if request.POST.get("change_notes"):
            change_summary += f" Notes: {request.POST['change_notes']}"

    # Create new document record
    doc = GeneratedDocument(
        financial_year=fy,
        file_format=fmt,
        document_type=doc_type,
        version=new_version,
        status=GeneratedDocument.DocumentStatus.DRAFT,
        change_summary=change_summary,
        generated_by=request.user,
    )
    doc.file.save(filename, ContentFile(file_content), save=True)

    # Update superseded_by on the old version
    if latest:
        latest.superseded_by = doc
        latest.save(update_fields=["superseded_by"])

    _log_action(
        request, "generate",
        f"Generated {doc.get_document_type_display()} v{new_version} ({fmt.upper()}) for {fy}",
        fy,
    )
    messages.success(request, f"{doc.get_document_type_display()} v{new_version} generated.")

    # Return the file as download
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if fmt == "docx" else "application/pdf"
    )
    response = HttpResponse(file_content, content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
@require_POST
def bulk_regenerate(request, pk):
    """Regenerate all documents for a financial year in one action."""
    fy = get_financial_year_for_user(request, pk)

    if not fy.trial_balance_lines.exists():
        messages.error(request, "Cannot regenerate: no trial balance data loaded.")
        return redirect("core:financial_year_detail", pk=pk)

    # Get unique document types that have been generated
    doc_types = (
        GeneratedDocument.objects
        .filter(financial_year=fy)
        .values_list("document_type", flat=True)
        .distinct()
    )

    regenerated = 0
    failed = []
    for doc_type in doc_types:
        latest = (
            GeneratedDocument.objects
            .filter(financial_year=fy, document_type=doc_type)
            .order_by("-version")
            .first()
        )
        new_version = (latest.version + 1) if latest else 1

        try:
            if doc_type == "financial_statements":
                from .docgen import generate_financial_statements
                has_open_risks = fy.risk_flags.filter(status="open").exists()
                buffer = generate_financial_statements(fy.pk, has_open_risks=has_open_risks)
            elif doc_type == "distribution_minutes":
                from .distmin_gen import generate_distribution_minutes as gen_distmin
                buffer = gen_distmin(fy.pk)
            else:
                continue

            file_content = buffer.getvalue()
            entity_name = fy.entity.entity_name.replace(" ", "_")
            filename = f"{entity_name}_{doc_type}_v{new_version}.docx"

            doc = GeneratedDocument(
                financial_year=fy,
                file_format="docx",
                document_type=doc_type,
                version=new_version,
                status=GeneratedDocument.DocumentStatus.DRAFT,
                change_summary=f"Bulk regeneration from v{latest.version}." if latest else "Initial generation.",
                generated_by=request.user,
            )
            doc.file.save(filename, ContentFile(file_content), save=True)

            if latest:
                latest.superseded_by = doc
                latest.save(update_fields=["superseded_by"])

            regenerated += 1
        except Exception as e:
            logger.exception(f"Bulk regenerate failed for doc_type={doc_type} on {fy}: {e}")
            failed.append(doc_type)
            continue

    _log_action(request, "generate", f"Bulk regenerated {regenerated} documents for {fy}", fy)
    msg = f"Bulk regeneration complete: {regenerated} documents updated."
    if failed:
        msg += f" {len(failed)} failed: {', '.join(failed)}."
    messages.success(request, msg)
    return redirect("core:financial_year_detail", pk=pk)


@login_required
@require_POST
def mark_document_final(request, doc_pk):
    """Mark a document as Final."""
    doc = get_object_or_404(GeneratedDocument, pk=doc_pk)
    doc.status = GeneratedDocument.DocumentStatus.FINAL
    doc.save(update_fields=["status"])
    messages.success(request, f"{doc.get_document_type_display()} v{doc.version} marked as Final.")
    return redirect("core:financial_year_detail", pk=doc.financial_year_id)


# ============================================================================
# 4. TRUST DISTRIBUTION WORKFLOW
# ============================================================================

@login_required
def trust_distribution(request, pk):
    """Trust distribution workspace — view and manage distribution allocations."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    if entity.entity_type != "trust":
        messages.error(request, "Distribution workspace is only available for trusts.")
        return redirect("core:financial_year_detail", pk=pk)

    # Get or create distribution record
    dist, created = TrustDistribution.objects.get_or_create(financial_year=fy)

    if created:
        # Auto-populate from trial balance
        tb_lines = fy.trial_balance_lines.all()
        total_credit = tb_lines.aggregate(t=Sum("credit"))["t"] or Decimal("0")
        total_debit = tb_lines.aggregate(t=Sum("debit"))["t"] or Decimal("0")
        dist.accounting_profit = total_credit - total_debit
        dist.distributable_income = dist.accounting_profit
        dist.other_income = dist.accounting_profit
        dist.save()

    # Get beneficiaries from entity officers
    beneficiaries = EntityOfficer.objects.filter(
        entity=entity
    ).filter(
        Q(role="beneficiary") | Q(roles__contains=["beneficiary"])
    )

    # Get existing allocations
    allocations = dist.allocations.select_related("beneficiary").all()
    allocation_map = {a.beneficiary_id: a for a in allocations}

    # Calculate totals
    total_allocated_pct = allocations.aggregate(t=Sum("percentage"))["t"] or Decimal("0")
    total_allocated_amt = allocations.aggregate(t=Sum("total_distribution"))["t"] or Decimal("0")

    if request.method == "POST":
        action = request.POST.get("action", "")

        if action == "update_income":
            try:
                dist.accounting_profit = Decimal(request.POST.get("accounting_profit", "0"))
                dist.taxable_income = Decimal(request.POST.get("taxable_income", "0"))
                dist.distributable_income = Decimal(request.POST.get("distributable_income", "0"))
                dist.capital_gains = Decimal(request.POST.get("capital_gains", "0"))
                dist.franked_dividends = Decimal(request.POST.get("franked_dividends", "0"))
                dist.foreign_income = Decimal(request.POST.get("foreign_income", "0"))
                dist.other_income = Decimal(request.POST.get("other_income", "0"))
                dist.corpus = Decimal(request.POST.get("corpus", "0"))
                dist.notes = request.POST.get("notes", "")
                dist.save()
                messages.success(request, "Distribution income updated.")
            except (InvalidOperation, TypeError) as e:
                messages.error(request, f"Invalid amount: {e}")

        elif action == "save_allocations":
            try:
                with transaction.atomic():
                    for ben in beneficiaries:
                        pct_key = f"pct_{ben.pk}"
                        flag_key = f"s100a_{ben.pk}"
                        notes_key = f"s100a_notes_{ben.pk}"

                        pct = Decimal(request.POST.get(pct_key, "0") or "0")

                        alloc, _ = BeneficiaryAllocation.objects.update_or_create(
                            distribution=dist,
                            beneficiary=ben,
                            defaults={
                                "percentage": pct,
                                "section_100a_flag": flag_key in request.POST,
                                "section_100a_notes": request.POST.get(notes_key, ""),
                            },
                        )
                        alloc.calculate_allocation()
                        alloc.save()

                    # Check if fully allocated
                    total_pct = dist.allocations.aggregate(t=Sum("percentage"))["t"] or Decimal("0")
                    dist.is_fully_allocated = (total_pct == Decimal("100"))
                    dist.save(update_fields=["is_fully_allocated"])

                messages.success(request, "Beneficiary allocations saved.")
            except (InvalidOperation, TypeError) as e:
                messages.error(request, f"Invalid allocation value: {e}")
            # Refresh data
            allocations = dist.allocations.select_related("beneficiary").all()
            allocation_map = {a.beneficiary_id: a for a in allocations}
            total_allocated_pct = dist.allocations.aggregate(t=Sum("percentage"))["t"] or Decimal("0")
            total_allocated_amt = dist.allocations.aggregate(t=Sum("total_distribution"))["t"] or Decimal("0")

        elif action == "push_to_minutes":
            return redirect("core:generate_distribution_minutes", pk=pk)

        return redirect("core:trust_distribution", pk=pk)

    context = {
        "fy": fy,
        "entity": entity,
        "dist": dist,
        "beneficiaries": beneficiaries,
        "allocations": allocations,
        "allocation_map": allocation_map,
        "total_allocated_pct": total_allocated_pct,
        "total_allocated_amt": total_allocated_amt,
    }
    return render(request, "core/trust_distribution.html", context)


@login_required
def generate_beneficiary_statement(request, pk, officer_pk):
    """Generate a beneficiary statement for a specific beneficiary."""
    fy = get_financial_year_for_user(request, pk)
    beneficiary = get_object_or_404(EntityOfficer, pk=officer_pk)

    try:
        dist = TrustDistribution.objects.get(financial_year=fy)
        alloc = BeneficiaryAllocation.objects.get(
            distribution=dist, beneficiary=beneficiary
        )
    except (TrustDistribution.DoesNotExist, BeneficiaryAllocation.DoesNotExist):
        messages.error(request, "No distribution allocation found for this beneficiary.")
        return redirect("core:trust_distribution", pk=pk)

    # Generate a simple DOCX beneficiary statement
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Title
    title = doc.add_heading("Beneficiary Distribution Statement", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Trust: {fy.entity.entity_name}")
    doc.add_paragraph(f"Financial Year: {fy.year_label} ({fy.start_date:%d/%m/%Y} to {fy.end_date:%d/%m/%Y})")
    doc.add_paragraph(f"Beneficiary: {beneficiary.full_name}")
    doc.add_paragraph(f"Distribution Percentage: {alloc.percentage}%")
    doc.add_paragraph("")

    # Distribution breakdown table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Income Stream"
    hdr[1].text = "Amount ($)"

    rows_data = [
        ("Capital Gains", alloc.allocated_capital_gains),
        ("Franked Dividends", alloc.allocated_franked_dividends),
        ("Foreign Income", alloc.allocated_foreign_income),
        ("Other Income", alloc.allocated_other_income),
    ]
    for label, amount in rows_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"${amount:,.2f}"

    # Total
    total_row = table.add_row().cells
    total_row[0].text = "Total Distribution"
    total_row[1].text = f"${alloc.total_distribution:,.2f}"
    for cell in total_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    entity_name = fy.entity.entity_name.replace(" ", "_")
    ben_name = beneficiary.full_name.replace(" ", "_")
    filename = f"{entity_name}_Beneficiary_Statement_{ben_name}_{fy.year_label}.docx"

    # Determine next version number
    latest_ben_doc = (
        GeneratedDocument.objects
        .filter(
            financial_year=fy,
            document_type=GeneratedDocument.DocumentType.BENEFICIARY_STATEMENT,
            change_summary__contains=beneficiary.full_name,
        )
        .order_by("-version")
        .first()
    )
    ben_version = (latest_ben_doc.version + 1) if latest_ben_doc else 1

    # Save as GeneratedDocument
    gen_doc = GeneratedDocument(
        financial_year=fy,
        file_format="docx",
        document_type=GeneratedDocument.DocumentType.BENEFICIARY_STATEMENT,
        version=ben_version,
        change_summary=f"Beneficiary: {beneficiary.full_name}",
        generated_by=request.user,
    )
    gen_doc.file.save(filename, ContentFile(buffer.getvalue()), save=True)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ============================================================================
# 5. PARTNERSHIP PROFIT ALLOCATION MODULE
# ============================================================================

@login_required
def partnership_allocation(request, pk):
    """Partnership profit allocation workspace."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    if entity.entity_type != "partnership":
        messages.error(request, "Partnership allocation is only available for partnerships.")
        return redirect("core:financial_year_detail", pk=pk)

    # Get or create allocation record
    alloc, created = PartnershipAllocation.objects.get_or_create(financial_year=fy)

    if created:
        tb_lines = fy.trial_balance_lines.all()
        total_credit = tb_lines.aggregate(t=Sum("credit"))["t"] or Decimal("0")
        total_debit = tb_lines.aggregate(t=Sum("debit"))["t"] or Decimal("0")
        alloc.net_profit = total_credit - total_debit
        alloc.residual_profit = alloc.net_profit
        alloc.save()

    # Get partners from entity officers
    partners = EntityOfficer.objects.filter(
        entity=entity
    ).filter(
        Q(role="partner") | Q(roles__contains=["partner"])
    )

    # Get existing shares
    shares = alloc.partner_shares.select_related("partner").all()
    share_map = {s.partner_id: s for s in shares}

    # Capital accounts
    capital_accounts = PartnerCapitalAccount.objects.filter(
        financial_year=fy
    ).select_related("partner")
    capital_map = {ca.partner_id: ca for ca in capital_accounts}

    if request.method == "POST":
        action = request.POST.get("action", "")

        if action == "update_profit":
            try:
                alloc.net_profit = Decimal(request.POST.get("net_profit", "0"))
                alloc.total_salary_allowances = Decimal(request.POST.get("total_salary_allowances", "0"))
                alloc.total_interest_on_capital = Decimal(request.POST.get("total_interest_on_capital", "0"))
                alloc.residual_profit = (
                    alloc.net_profit - alloc.total_salary_allowances - alloc.total_interest_on_capital
                )
                alloc.notes = request.POST.get("notes", "")
                alloc.save()
                messages.success(request, "Partnership profit updated.")
            except (InvalidOperation, TypeError) as e:
                messages.error(request, f"Invalid amount: {e}")

        elif action == "save_shares":
            try:
                with transaction.atomic():
                    for partner in partners:
                        salary_key = f"salary_{partner.pk}"
                        interest_key = f"interest_{partner.pk}"
                        residual_key = f"residual_pct_{partner.pk}"

                        share, _ = PartnerShare.objects.update_or_create(
                            allocation=alloc,
                            partner=partner,
                            defaults={
                                "salary_allowance": Decimal(request.POST.get(salary_key, "0") or "0"),
                                "interest_on_capital": Decimal(request.POST.get(interest_key, "0") or "0"),
                                "residual_share_pct": Decimal(request.POST.get(residual_key, "0") or "0"),
                            },
                        )
                        share.calculate_share()
                        share.save()

                        # Update capital account
                        ca, _ = PartnerCapitalAccount.objects.update_or_create(
                            financial_year=fy,
                            partner=partner,
                            defaults={
                                "profit_share": share.total_share,
                                "drawings": Decimal(
                                    request.POST.get(f"drawings_{partner.pk}", "0") or "0"
                                ),
                                "capital_contributions": Decimal(
                                    request.POST.get(f"contributions_{partner.pk}", "0") or "0"
                                ),
                            },
                        )
                        # Auto-populate opening balance from prior year if available
                        if ca.opening_balance == 0 and fy.prior_year:
                            prior_ca = PartnerCapitalAccount.objects.filter(
                                financial_year=fy.prior_year, partner=partner
                            ).first()
                            if prior_ca:
                                ca.opening_balance = prior_ca.closing_balance
                        ca.calculate_closing()
                        ca.save()

                messages.success(request, "Partner shares and capital accounts updated.")
            except (InvalidOperation, TypeError) as e:
                messages.error(request, f"Invalid share value: {e}")
            # Refresh
            shares = alloc.partner_shares.select_related("partner").all()
            share_map = {s.partner_id: s for s in shares}
            capital_accounts = PartnerCapitalAccount.objects.filter(
                financial_year=fy
            ).select_related("partner")
            capital_map = {ca.partner_id: ca for ca in capital_accounts}

        return redirect("core:partnership_allocation", pk=pk)

    context = {
        "fy": fy,
        "entity": entity,
        "alloc": alloc,
        "partners": partners,
        "shares": shares,
        "share_map": share_map,
        "capital_accounts": capital_accounts,
        "capital_map": capital_map,
    }
    return render(request, "core/partnership_allocation.html", context)


@login_required
def generate_partner_statement(request, pk, officer_pk):
    """Generate a partner statement for a specific partner."""
    fy = get_financial_year_for_user(request, pk)
    partner = get_object_or_404(EntityOfficer, pk=officer_pk)

    try:
        alloc = PartnershipAllocation.objects.get(financial_year=fy)
        share = PartnerShare.objects.get(allocation=alloc, partner=partner)
    except (PartnershipAllocation.DoesNotExist, PartnerShare.DoesNotExist):
        messages.error(request, "No allocation found for this partner.")
        return redirect("core:partnership_allocation", pk=pk)

    capital = PartnerCapitalAccount.objects.filter(
        financial_year=fy, partner=partner
    ).first()

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    title = doc.add_heading("Partner Statement", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Partnership: {fy.entity.entity_name}")
    doc.add_paragraph(f"Financial Year: {fy.year_label}")
    doc.add_paragraph(f"Partner: {partner.full_name}")
    doc.add_paragraph("")

    # Profit share
    doc.add_heading("Profit Share", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Amount ($)"

    for label, amount in [
        ("Salary Allowance", share.salary_allowance),
        ("Interest on Capital", share.interest_on_capital),
        (f"Residual Share ({share.residual_share_pct}%)", share.residual_share_amount),
        ("Total Profit Share", share.total_share),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"${amount:,.2f}"

    # Capital account
    if capital:
        doc.add_paragraph("")
        doc.add_heading("Capital Account", level=2)
        cap_table = doc.add_table(rows=1, cols=2)
        cap_table.style = "Table Grid"
        hdr = cap_table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Amount ($)"

        for label, amount in [
            ("Opening Balance", capital.opening_balance),
            ("Capital Contributions", capital.capital_contributions),
            ("Profit Share", capital.profit_share),
            ("Less: Drawings", capital.drawings),
            ("Closing Balance", capital.closing_balance),
        ]:
            row = cap_table.add_row().cells
            row[0].text = label
            row[1].text = f"${amount:,.2f}"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    entity_name = fy.entity.entity_name.replace(" ", "_")
    partner_name = partner.full_name.replace(" ", "_")
    filename = f"{entity_name}_Partner_Statement_{partner_name}_{fy.year_label}.docx"

    # Determine next version number
    latest_partner_doc = (
        GeneratedDocument.objects
        .filter(
            financial_year=fy,
            document_type=GeneratedDocument.DocumentType.PARTNER_STATEMENT,
            change_summary__contains=partner.full_name,
        )
        .order_by("-version")
        .first()
    )
    partner_version = (latest_partner_doc.version + 1) if latest_partner_doc else 1

    gen_doc = GeneratedDocument(
        financial_year=fy,
        file_format="docx",
        document_type=GeneratedDocument.DocumentType.PARTNER_STATEMENT,
        version=partner_version,
        change_summary=f"Partner: {partner.full_name}",
        generated_by=request.user,
    )
    gen_doc.file.save(filename, ContentFile(buffer.getvalue()), save=True)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ============================================================================
# 6. WORKING PAPER NOTES PER ACCOUNT
# ============================================================================

@login_required
def workpaper_notes_api(request, pk):
    """
    GET: Return all workpaper notes for a financial year as JSON.
    POST: Create or update a workpaper note.
    """
    fy = get_financial_year_for_user(request, pk)

    if request.method == "GET":
        notes = WorkpaperNote.objects.filter(financial_year=fy)
        # Group by account_code: merge preparer + reviewer into one object per account
        accounts = {}
        for note in notes:
            code = note.account_code
            if code not in accounts:
                accounts[code] = {
                    "account_code": code,
                    "account_name": note.account_name,
                    "status": "blank",
                    "preparer_note": "",
                    "reviewer_note": "",
                }
            if note.note_type == "preparer":
                accounts[code]["preparer_note"] = note.content
                accounts[code]["status"] = note.status
            elif note.note_type == "reviewer":
                accounts[code]["reviewer_note"] = note.content

        return JsonResponse({"notes": list(accounts.values())})

    elif request.method == "POST":
        body = json.loads(request.body)
        account_code = body.get("account_code", "")
        preparer_note = body.get("preparer_note", "")
        reviewer_note = body.get("reviewer_note", "")
        status = body.get("status", "blank")

        if not account_code:
            return JsonResponse({"error": "account_code required"}, status=400)

        # Get account name from TB line
        tb_line = fy.trial_balance_lines.filter(account_code=account_code).first()
        account_name = tb_line.account_name if tb_line else ""

        # Save preparer note
        WorkpaperNote.objects.update_or_create(
            financial_year=fy,
            account_code=account_code,
            note_type="preparer",
            defaults={
                "account_name": account_name,
                "content": preparer_note,
                "status": status if preparer_note else "blank",
                "author": request.user,
            },
        )

        # Save reviewer note
        if reviewer_note:
            WorkpaperNote.objects.update_or_create(
                financial_year=fy,
                account_code=account_code,
                note_type="reviewer",
                defaults={
                    "account_name": account_name,
                    "content": reviewer_note,
                    "status": "in_progress",
                    "author": request.user,
                },
            )

        return JsonResponse({"ok": True})


@login_required
@require_POST
def carry_forward_notes(request, pk):
    """Carry forward working paper notes from the prior year."""
    fy = get_financial_year_for_user(request, pk)
    prior_fy = fy.prior_year

    if not prior_fy:
        messages.warning(request, "No prior year linked. Cannot carry forward notes.")
        return redirect("core:financial_year_detail", pk=pk)

    prior_notes = WorkpaperNote.objects.filter(
        financial_year=prior_fy,
        note_type="preparer",
    ).exclude(content="")

    carried = 0
    for pn in prior_notes:
        # Only carry forward if no note exists yet for this account
        exists = WorkpaperNote.objects.filter(
            financial_year=fy,
            account_code=pn.account_code,
            note_type="preparer",
        ).exists()
        if not exists:
            WorkpaperNote.objects.create(
                financial_year=fy,
                account_code=pn.account_code,
                account_name=pn.account_name,
                note_type="preparer",
                status=WorkpaperNote.NoteStatus.BLANK,
                content=f"[PY {prior_fy.year_label}] {pn.content}",
                is_carried_forward=True,
                source_year_label=prior_fy.year_label,
                author=request.user,
            )
            carried += 1

    _log_action(request, "import", f"Carried forward {carried} workpaper notes from {prior_fy.year_label}", fy)
    messages.success(request, f"Carried forward {carried} notes from {prior_fy.year_label}.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def export_workpaper_notes(request, pk):
    """Export all working paper notes as a single DOCX document."""
    fy = get_financial_year_for_user(request, pk)
    notes = WorkpaperNote.objects.filter(financial_year=fy).order_by("account_code", "note_type")

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    title = doc.add_heading("Working Paper Notes", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Entity: {fy.entity.entity_name}")
    doc.add_paragraph(f"Financial Year: {fy.year_label}")
    doc.add_paragraph(f"Exported: {timezone.now():%d/%m/%Y %H:%M}")
    doc.add_paragraph("")

    # Summary
    total = notes.count()
    completed = notes.filter(status="completed").count()
    doc.add_paragraph(f"Total notes: {total} | Completed: {completed}")
    doc.add_paragraph("")

    current_code = None
    for note in notes:
        if note.account_code != current_code:
            current_code = note.account_code
            doc.add_heading(f"{note.account_code} - {note.account_name}", level=2)

        prefix = "REVIEWER" if note.note_type == "reviewer" else "PREPARER"
        cf_tag = " [Carried Forward]" if note.is_carried_forward else ""
        doc.add_paragraph(
            f"[{prefix}] Status: {note.get_status_display()}{cf_tag}"
        )
        if note.content:
            doc.add_paragraph(note.content)
        if note.author:
            doc.add_paragraph(
                f"By: {note.author.get_full_name()} | Updated: {note.updated_at:%d/%m/%Y %H:%M}",
            ).style = doc.styles["Normal"]
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    entity_name = fy.entity.entity_name.replace(" ", "_")
    filename = f"{entity_name}_Workpaper_Notes_{fy.year_label}.docx"

    # Save as GeneratedDocument
    gen_doc = GeneratedDocument(
        financial_year=fy,
        file_format="docx",
        document_type=GeneratedDocument.DocumentType.WORKPAPER_NOTES,
        version=1,
        generated_by=request.user,
    )
    gen_doc.file.save(filename, ContentFile(buffer.getvalue()), save=True)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ============================================================================
# 7. BULK ENTITY IMPORT / ONBOARDING TOOL
# ============================================================================

def _require_senior(request):
    """Raise PermissionDenied if user is not admin or senior accountant."""
    if not (request.user.is_senior or request.user.is_admin):
        raise PermissionDenied("Only administrators and senior accountants can perform bulk imports.")


@login_required
def bulk_import_start(request):
    """Start a bulk entity import — upload CSV/Excel file."""
    _require_senior(request)

    if request.method == "POST":
        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            messages.error(request, "Please select a file to upload.")
            return redirect("core:bulk_import_start")

        ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
        if ext not in ("csv", "xlsx"):
            messages.error(request, "Only CSV and Excel (.xlsx) files are supported.")
            return redirect("core:bulk_import_start")

        # Limit upload size to 10 MB
        if uploaded_file.size > 10 * 1024 * 1024:
            messages.error(request, "File too large. Maximum upload size is 10 MB.")
            return redirect("core:bulk_import_start")

        job = EntityImportJob(
            original_filename=uploaded_file.name,
            uploaded_by=request.user,
            status=EntityImportJob.ImportStatus.PENDING,
        )
        job.file.save(uploaded_file.name, uploaded_file, save=True)

        return redirect("core:bulk_import_map", pk=job.pk)

    # GET: show upload form
    return render(request, "core/bulk_import_start.html")


@login_required
def bulk_import_template(request):
    """Download a blank template spreadsheet for bulk entity import."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Entity Import Template"

    headers = [
        "Entity Name", "Entity Type", "ABN", "ACN", "TFN",
        "Trading As", "Registration Date", "Financial Year End",
        "Contact Email", "Contact Phone",
        "Address Line 1", "Address Line 2", "Suburb", "State", "Postcode",
    ]
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)

    # Example row
    example = [
        "Example Pty Ltd", "company", "12345678901", "123456789", "",
        "Example Trading", "2020-01-15", "06-30",
        "info@example.com", "03 9555 1234",
        "123 Main Street", "Suite 4", "Melbourne", "VIC", "3000",
    ]
    for col, val in enumerate(example, 1):
        ws.cell(row=2, column=col, value=val)

    # Entity type reference
    ws2 = wb.create_sheet("Reference")
    ws2.cell(row=1, column=1, value="Valid Entity Types")
    for i, et in enumerate(["company", "trust", "partnership", "sole_trader", "smsf"], 2):
        ws2.cell(row=i, column=1, value=et)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = 'attachment; filename="StatementHub_Entity_Import_Template.xlsx"'
    return response


@login_required
def bulk_import_map(request, pk):
    """Column mapping screen for bulk import."""
    _require_senior(request)
    job = get_object_or_404(EntityImportJob, pk=pk)

    # Parse the file to get headers
    ext = job.original_filename.rsplit(".", 1)[-1].lower()
    headers = []

    try:
        if ext == "csv":
            content = job.file.read().decode("utf-8-sig")
            reader = csv.reader(io.StringIO(content))
            headers = next(reader, [])
        else:
            wb = openpyxl.load_workbook(job.file, read_only=True)
            ws = wb.active
            headers = [cell.value or "" for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            wb.close()
    except Exception as e:
        messages.error(request, f"Error reading file: {e}")
        return redirect("core:bulk_import_start")

    # StatementHub fields
    target_fields = [
        ("entity_name", "Entity Name *"),
        ("entity_type", "Entity Type *"),
        ("abn", "ABN"),
        ("acn", "ACN"),
        ("tfn", "TFN"),
        ("trading_as", "Trading As"),
        ("registration_date", "Registration Date"),
        ("financial_year_end", "Financial Year End"),
        ("contact_email", "Contact Email"),
        ("contact_phone", "Contact Phone"),
        ("address_line_1", "Address Line 1"),
        ("address_line_2", "Address Line 2"),
        ("suburb", "Suburb"),
        ("state", "State"),
        ("postcode", "Postcode"),
    ]

    if request.method == "POST":
        mapping = {}
        for field_key, _ in target_fields:
            col_idx = request.POST.get(f"map_{field_key}", "")
            if col_idx:
                mapping[field_key] = int(col_idx)

        job.column_mapping = mapping
        job.status = EntityImportJob.ImportStatus.VALIDATING
        job.save()

        return redirect("core:bulk_import_validate", pk=job.pk)

    context = {
        "job": job,
        "headers": list(enumerate(headers)),
        "target_fields": target_fields,
    }
    return render(request, "core/bulk_import_map.html", context)


@login_required
def bulk_import_validate(request, pk):
    """Validate the import data and show errors before importing."""
    _require_senior(request)
    job = get_object_or_404(EntityImportJob, pk=pk)
    mapping = job.column_mapping

    # Parse all rows
    ext = job.original_filename.rsplit(".", 1)[-1].lower()
    rows = []

    try:
        job.file.seek(0)
        if ext == "csv":
            content = job.file.read().decode("utf-8-sig")
            reader = csv.reader(io.StringIO(content))
            next(reader)  # skip header
            for row in reader:
                rows.append(row)
        else:
            wb = openpyxl.load_workbook(job.file, read_only=True)
            ws = wb.active
            first = True
            for row in ws.iter_rows():
                if first:
                    first = False
                    continue
                rows.append([cell.value for cell in row])
            wb.close()
    except Exception as e:
        messages.error(request, f"Error reading file: {e}")
        return redirect("core:bulk_import_start")

    job.total_rows = len(rows)
    errors = []
    valid_types = {"company", "trust", "partnership", "sole_trader", "smsf"}
    existing_abns = set(
        Entity.objects.exclude(abn="").values_list("abn", flat=True)
    )

    seen_abns = set()
    preview_rows = []
    for i, row in enumerate(rows, 2):  # Row 2 onwards (1 = header)
        row_data = {}
        for field_key, col_idx in mapping.items():
            val = row[col_idx] if col_idx < len(row) else ""
            row_data[field_key] = str(val).strip() if val else ""

        # Validate required fields
        if not row_data.get("entity_name"):
            errors.append({"row": i, "field": "entity_name", "message": "Entity name is required"})
        if not row_data.get("entity_type"):
            errors.append({"row": i, "field": "entity_type", "message": "Entity type is required"})
        elif row_data["entity_type"].lower() not in valid_types:
            errors.append({
                "row": i, "field": "entity_type",
                "message": f"Invalid entity type: {row_data['entity_type']}. Must be one of: {', '.join(valid_types)}"
            })

        # Check duplicate ABN (against database and within file)
        abn = row_data.get("abn", "")
        if abn:
            if abn in existing_abns:
                errors.append({"row": i, "field": "abn", "message": f"Duplicate ABN: {abn} already exists in database"})
            elif abn in seen_abns:
                errors.append({"row": i, "field": "abn", "message": f"Duplicate ABN: {abn} appears multiple times in this file"})
            seen_abns.add(abn)

        preview_rows.append({"row_num": i, "data": row_data})

    job.validation_errors = errors
    job.status = EntityImportJob.ImportStatus.VALIDATED
    job.save()

    if request.method == "POST" and not errors:
        return redirect("core:bulk_import_execute", pk=job.pk)

    # Count rows with errors (a row may have multiple errors)
    error_rows = len(set(e["row"] for e in errors))
    context = {
        "job": job,
        "errors": errors,
        "preview_rows": preview_rows[:50],  # Show first 50 for preview
        "total_rows": len(rows),
        "error_count": len(errors),
        "valid_row_count": len(rows) - error_rows,
    }
    return render(request, "core/bulk_import_validate.html", context)


@login_required
@require_POST
def bulk_import_execute(request, pk):
    """Execute the validated import — create entities in the database."""
    _require_senior(request)
    job = get_object_or_404(EntityImportJob, pk=pk)

    if job.validation_errors:
        messages.error(request, "Cannot import: there are validation errors. Fix them first.")
        return redirect("core:bulk_import_validate", pk=pk)

    mapping = job.column_mapping
    ext = job.original_filename.rsplit(".", 1)[-1].lower()
    rows = []

    try:
        job.file.seek(0)
        if ext == "csv":
            content = job.file.read().decode("utf-8-sig")
            reader = csv.reader(io.StringIO(content))
            next(reader)
            for row in reader:
                rows.append(row)
        else:
            wb = openpyxl.load_workbook(job.file, read_only=True)
            ws = wb.active
            first = True
            for row in ws.iter_rows():
                if first:
                    first = False
                    continue
                rows.append([cell.value for cell in row])
            wb.close()
    except Exception as e:
        job.status = EntityImportJob.ImportStatus.FAILED
        job.save()
        messages.error(request, f"Error reading file: {e}")
        return redirect("core:bulk_import_start")

    job.status = EntityImportJob.ImportStatus.IMPORTING
    job.save()

    created = 0
    skipped = 0
    errors = 0

    with transaction.atomic():
        for row in rows:
            row_data = {}
            for field_key, col_idx in mapping.items():
                val = row[col_idx] if col_idx < len(row) else ""
                row_data[field_key] = str(val).strip() if val else ""

            entity_name = row_data.get("entity_name", "")
            entity_type = row_data.get("entity_type", "").lower()

            if not entity_name or not entity_type:
                skipped += 1
                continue

            try:
                entity = Entity(
                    entity_name=entity_name,
                    entity_type=entity_type,
                    abn=row_data.get("abn", ""),
                    acn=row_data.get("acn", ""),
                    tfn=row_data.get("tfn", ""),
                    trading_as=row_data.get("trading_as", ""),
                    financial_year_end=row_data.get("financial_year_end", "06-30") or "06-30",
                    contact_email=row_data.get("contact_email", ""),
                    contact_phone=row_data.get("contact_phone", ""),
                    address_line_1=row_data.get("address_line_1", ""),
                    address_line_2=row_data.get("address_line_2", ""),
                    suburb=row_data.get("suburb", ""),
                    state=row_data.get("state", ""),
                    postcode=row_data.get("postcode", ""),
                )
                if row_data.get("registration_date"):
                    from datetime import datetime
                    try:
                        entity.registration_date = datetime.strptime(
                            row_data["registration_date"], "%Y-%m-%d"
                        ).date()
                    except ValueError:
                        pass
                entity.save()
                created += 1
            except Exception as e:
                logger.exception(f"Bulk import failed for row entity_name={entity_name}: {e}")
                errors += 1

    job.created_count = created
    job.skipped_count = skipped
    job.error_count = errors
    job.status = EntityImportJob.ImportStatus.COMPLETED
    job.completed_at = timezone.now()
    job.save()

    _log_action(
        request, "import",
        f"Bulk entity import: {created} created, {skipped} skipped, {errors} errors",
    )
    messages.success(
        request,
        f"Import complete: {created} entities created, {skipped} skipped, {errors} errors."
    )
    return redirect("core:entity_list")


# ---------------------------------------------------------------------------
# Entity Assignment Management
# ---------------------------------------------------------------------------
@login_required
def entity_assignments(request):
    """
    View and manage primary_accountant and reviewer assignments for all entities.
    Supports filtering by entity type, assignment status, and staff member.
    """
    if not request.user.is_senior:
        messages.error(request, "Only senior accountants and administrators can manage entity assignments.")
        return redirect("core:entity_list")

    from django.contrib.auth import get_user_model
    User = get_user_model()

    entities = Entity.objects.filter(is_archived=False).select_related(
        "primary_accountant", "reviewer"
    ).order_by("entity_name")

    # Filters
    entity_type_filter = request.GET.get("entity_type", "")
    assignment_filter = request.GET.get("assignment", "")
    staff_filter = request.GET.get("staff", "")

    if entity_type_filter:
        entities = entities.filter(entity_type=entity_type_filter)
    if assignment_filter == "unassigned":
        entities = entities.filter(primary_accountant__isnull=True)
    elif assignment_filter == "no_reviewer":
        entities = entities.filter(reviewer__isnull=True)
    elif assignment_filter == "fully_assigned":
        entities = entities.filter(
            primary_accountant__isnull=False, reviewer__isnull=False
        )
    if staff_filter:
        entities = entities.filter(
            Q(primary_accountant_id=staff_filter) | Q(reviewer_id=staff_filter)
        )

    # Staff list for dropdowns
    staff_members = User.objects.filter(is_active=True).order_by("first_name", "last_name")

    # Summary stats
    total_entities = Entity.objects.filter(is_archived=False).count()
    unassigned_count = Entity.objects.filter(
        is_archived=False, primary_accountant__isnull=True
    ).count()
    no_reviewer_count = Entity.objects.filter(
        is_archived=False, reviewer__isnull=True
    ).count()

    # Staff workload summary
    workload = []
    for staff in staff_members:
        primary_count = Entity.objects.filter(
            is_archived=False, primary_accountant=staff
        ).count()
        review_count = Entity.objects.filter(
            is_archived=False, reviewer=staff
        ).count()
        if primary_count > 0 or review_count > 0:
            workload.append({
                "user": staff,
                "primary_count": primary_count,
                "review_count": review_count,
                "total": primary_count + review_count,
            })
    workload.sort(key=lambda x: x["total"], reverse=True)

    context = {
        "entities": entities,
        "staff_members": staff_members,
        "total_entities": total_entities,
        "unassigned_count": unassigned_count,
        "no_reviewer_count": no_reviewer_count,
        "workload": workload,
        "entity_type_filter": entity_type_filter,
        "assignment_filter": assignment_filter,
        "staff_filter": staff_filter,
        "entity_type_choices": Entity.EntityType.choices,
    }
    return render(request, "core/entity_assignments.html", context)


@login_required
@require_POST
def bulk_assign_entities(request):
    """
    Bulk assign primary_accountant and/or reviewer to selected entities.
    Accepts a list of entity IDs and the staff member to assign.
    """
    if not request.user.is_senior:
        messages.error(request, "Only senior accountants and administrators can manage entity assignments.")
        return redirect("core:entity_assignments")

    from django.contrib.auth import get_user_model
    User = get_user_model()

    entity_ids = request.POST.getlist("entity_ids")
    primary_accountant_id = request.POST.get("primary_accountant_id", "")
    reviewer_id = request.POST.get("reviewer_id", "")
    action = request.POST.get("assign_action", "set")  # set or clear

    if not entity_ids:
        messages.warning(request, "No entities selected.")
        return redirect("core:entity_assignments")

    entities = Entity.objects.filter(pk__in=entity_ids, is_archived=False)
    count = entities.count()

    if action == "clear":
        clear_primary = request.POST.get("clear_primary") == "1"
        clear_reviewer = request.POST.get("clear_reviewer") == "1"
        update_fields = {}
        if clear_primary:
            update_fields["primary_accountant"] = None
        if clear_reviewer:
            update_fields["reviewer"] = None
        if update_fields:
            entities.update(**update_fields)
            cleared = ", ".join(k.replace("_", " ") for k in update_fields.keys())
            messages.success(request, f"Cleared {cleared} on {count} entities.")
            _log_action(
                request, "bulk_assign",
                f"Cleared {cleared} on {count} entities",
            )
    else:
        updates = {}
        if primary_accountant_id:
            try:
                pa = User.objects.get(pk=primary_accountant_id, is_active=True)
                updates["primary_accountant"] = pa
            except User.DoesNotExist:
                messages.error(request, "Selected primary accountant not found.")
                return redirect("core:entity_assignments")

        if reviewer_id:
            try:
                rev = User.objects.get(pk=reviewer_id, is_active=True)
                updates["reviewer"] = rev
            except User.DoesNotExist:
                messages.error(request, "Selected reviewer not found.")
                return redirect("core:entity_assignments")

        if not updates:
            messages.warning(request, "No staff member selected for assignment.")
            return redirect("core:entity_assignments")

        entities.update(**updates)
        assigned_roles = []
        if "primary_accountant" in updates:
            assigned_roles.append(f"Primary: {updates['primary_accountant'].get_full_name()}")
        if "reviewer" in updates:
            assigned_roles.append(f"Reviewer: {updates['reviewer'].get_full_name()}")

        messages.success(
            request,
            f"Assigned {', '.join(assigned_roles)} to {count} entities."
        )
        _log_action(
            request, "bulk_assign",
            f"Bulk assigned {', '.join(assigned_roles)} to {count} entities",
        )

    return redirect("core:entity_assignments")


@login_required
@require_POST
def update_entity_assignment(request, pk):
    """
    Update the primary_accountant and/or reviewer for a single entity (inline edit).
    """
    if not request.user.is_senior:
        messages.error(request, "Only senior accountants and administrators can manage entity assignments.")
        return redirect("core:entity_assignments")

    from django.contrib.auth import get_user_model
    User = get_user_model()

    entity = get_object_or_404(Entity, pk=pk)
    primary_id = request.POST.get("primary_accountant_id", "")
    reviewer_id = request.POST.get("reviewer_id", "")

    changed = []
    if primary_id:
        try:
            entity.primary_accountant = User.objects.get(pk=primary_id, is_active=True)
            changed.append("primary accountant")
        except User.DoesNotExist:
            pass
    elif "primary_accountant_id" in request.POST:
        entity.primary_accountant = None
        changed.append("primary accountant (cleared)")

    if reviewer_id:
        try:
            entity.reviewer = User.objects.get(pk=reviewer_id, is_active=True)
            changed.append("reviewer")
        except User.DoesNotExist:
            pass
    elif "reviewer_id" in request.POST:
        entity.reviewer = None
        changed.append("reviewer (cleared)")

    if changed:
        entity.save()
        messages.success(
            request,
            f"Updated {', '.join(changed)} for {entity.entity_name}."
        )

    # Return to the referring page or assignments page
    from django.utils.http import url_has_allowed_host_and_scheme
    next_url = request.POST.get("next", "")
    if next_url and url_has_allowed_host_and_scheme(next_url, allowed_hosts={request.get_host()}):
        return redirect(next_url)
    return redirect("core:entity_assignments")
