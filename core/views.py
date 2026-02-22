"""MCS Platform - Core Views"""
import openpyxl
from decimal import Decimal, InvalidOperation
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django_ratelimit.decorators import ratelimit
from django.db.models import Q, Count, Sum
from django.http import HttpResponse, JsonResponse
from django.views.decorators.http import require_POST
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone

from .models import (
    Client, Entity, FinancialYear, TrialBalanceLine,
    AccountMapping, ChartOfAccount, ClientAccountMapping, AdjustingJournal,
    JournalLine, GeneratedDocument, AuditLog, EntityOfficer,
    ClientAssociate, AccountingSoftware, MeetingNote,
    DepreciationAsset, RiskFlag, StockItem, ActivityLog,
)
from .forms import (
    ClientForm, EntityForm, FinancialYearForm,
    TrialBalanceUploadForm, AccountMappingForm,
    AdjustingJournalForm, JournalLineFormSet,
    EntityOfficerForm, ClientAssociateForm, AccountingSoftwareForm,
    MeetingNoteForm,
)
from config.authorization import get_entity_for_user, get_financial_year_for_user


def _log_action(request, action, description, obj=None):
    """Create an audit log entry."""
    AuditLog.objects.create(
        user=request.user,
        action=action,
        description=description,
        affected_object_type=type(obj).__name__ if obj else "",
        affected_object_id=str(obj.pk) if obj else "",
        ip_address=request.META.get("REMOTE_ADDR"),
    )


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------
@login_required
def dashboard(request):
    user = request.user
    if user.is_admin:
        clients = Client.objects.filter(is_active=True)
    else:
        clients = Client.objects.filter(
            Q(assigned_accountant=user) & Q(is_active=True)
        )

    # Time-based greeting
    import datetime
    now = timezone.localtime(timezone.now())
    hour = now.hour
    if hour < 12:
        greeting = "Good morning"
    elif hour < 17:
        greeting = "Good afternoon"
    else:
        greeting = "Good evening"

    # Unfinalised financial years (Draft, In Review, Reviewed) — grouped by client
    unfinalised_years = (
        FinancialYear.objects.filter(
            entity__client__in=clients,
            status__in=["draft", "in_review", "reviewed"],
        )
        .select_related("entity", "entity__client")
        .order_by("entity__client__name", "entity__entity_name", "-end_date")
    )

    # Open audit risk flags across all clients I'm working on
    open_risk_flags = (
        RiskFlag.objects.filter(
            financial_year__entity__client__in=clients,
            status__in=["open", "reviewed"],
        )
        .select_related("financial_year", "financial_year__entity", "financial_year__entity__client")
        .order_by("-severity", "-created_at")[:20]
    )

    # Group risk flags by client/entity for display
    risk_summary = {}
    for flag in open_risk_flags:
        key = flag.financial_year.entity.entity_name
        if key not in risk_summary:
            risk_summary[key] = {
                "entity_name": flag.financial_year.entity.entity_name,
                "fy_pk": flag.financial_year.pk,
                "flags": [],
            }
        risk_summary[key]["flags"].append(flag)

    # Recent activity log
    recent_activities = (
        ActivityLog.objects.all()
        .order_by("-created_at")[:30]
    )

    # Unread notification count for the bell
    unread_count = ActivityLog.objects.filter(is_read=False).count()

    context = {
        "greeting": greeting,
        "now": now,
        "unfinalised_years": unfinalised_years,
        "risk_summary": risk_summary,
        "open_risk_count": open_risk_flags.count(),
        "recent_activities": recent_activities,
        "unread_count": unread_count,
    }
    return render(request, "core/dashboard.html", context)


# ---------------------------------------------------------------------------
# Entities (top-level — replaces Clients)
# ---------------------------------------------------------------------------
@login_required
def entity_list(request):
    """Main list page showing all entities (replaces client_list)."""
    query = request.GET.get("q", "")
    entity_type = request.GET.get("entity_type", "")
    status_filter = request.GET.get("status", "")
    show_archived = request.GET.get("show_archived", "") == "1"

    if show_archived:
        if request.user.is_admin:
            entities = Entity.objects.filter(is_archived=True)
        else:
            entities = Entity.objects.filter(
                assigned_accountant=request.user, is_archived=True
            )
    else:
        if request.user.is_admin:
            entities = Entity.objects.filter(is_archived=False)
        else:
            entities = Entity.objects.filter(
                assigned_accountant=request.user, is_archived=False
            )

    if query:
        entities = entities.filter(
            Q(entity_name__icontains=query)
            | Q(abn__icontains=query)
            | Q(trading_as__icontains=query)
        ).distinct()

    if entity_type:
        entities = entities.filter(entity_type=entity_type)

    entities = entities.annotate(num_fys=Count("financial_years"))

    context = {
        "entities": entities,
        "query": query,
        "entity_type": entity_type,
        "status_filter": status_filter,
        "show_archived": show_archived,
        "entity_types": Entity.EntityType.choices,
    }

    if request.htmx:
        return render(request, "partials/entity_list_rows.html", context)
    return render(request, "core/entity_list.html", context)


# Keep client_list as alias for backward compatibility
client_list = entity_list


@login_required
def client_create(request):
    """Redirect to entity_create — clients are no longer created directly."""
    return redirect("core:entity_create")


@login_required
def client_detail(request, pk):
    """Legacy redirect — try to find entity or client."""
    # Try as entity first
    entity = Entity.objects.filter(pk=pk).first()
    if entity:
        return redirect("core:entity_detail", pk=pk)
    # Try as client — redirect to first entity
    client = Client.objects.filter(pk=pk).first()
    if client:
        first_entity = client.entities.first()
        if first_entity:
            return redirect("core:entity_detail", pk=first_entity.pk)
    return redirect("core:entity_list")


@login_required
def client_edit(request, pk):
    """Legacy redirect."""
    return redirect("core:entity_list")


# ---------------------------------------------------------------------------
# Entities
# ---------------------------------------------------------------------------
@login_required
def entity_create(request, client_pk=None):
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to create entities.")
        return redirect("core:entity_list")

    if request.method == "POST":
        form = EntityForm(request.POST, user=request.user)
        if form.is_valid():
            entity = form.save()
            _log_action(request, "import", f"Created entity: {entity.entity_name}", entity)
            messages.success(request, f"Entity '{entity.entity_name}' created.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = EntityForm(user=request.user)
    return render(request, "core/entity_form.html", {
        "form": form, "title": "Create Entity"
    })


@login_required
def entity_detail(request, pk):
    entity = get_entity_for_user(request, pk)
    # Audit log: data access
    _log_action(request, "view", f"Viewed entity: {entity.entity_name}", entity)
    financial_years = entity.financial_years.all()
    officers = entity.officers.all().order_by('date_ceased', 'full_name')
    unfinalised_count = financial_years.exclude(status="finalised").count()
    associates = entity.associates.filter(is_active=True)
    family_associates = [a for a in associates if a.is_family]
    business_associates = [a for a in associates if not a.is_family]
    software_configs = entity.software_configs.all()
    meeting_notes = entity.meeting_notes.all()[:20]
    pending_followups = entity.meeting_notes.filter(
        follow_up_completed=False, follow_up_date__isnull=False
    ).order_by("follow_up_date")
    context = {
        "entity": entity,
        "financial_years": financial_years,
        "officers": officers,
        "unfinalised_count": unfinalised_count,
        "family_associates": family_associates,
        "business_associates": business_associates,
        "software_configs": software_configs,
        "meeting_notes": meeting_notes,
        "pending_followups": pending_followups,
    }
    return render(request, "core/entity_detail.html", context)


@login_required
def entity_edit(request, pk):
    entity = get_entity_for_user(request, pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to edit entities.")
        return redirect("core:entity_detail", pk=pk)

    if request.method == "POST":
        form = EntityForm(request.POST, instance=entity, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, f"Entity '{entity.entity_name}' updated.")
            return redirect("core:entity_detail", pk=pk)
    else:
        form = EntityForm(instance=entity, user=request.user)
    return render(request, "core/entity_form.html", {
        "form": form, "title": f"Edit: {entity.entity_name}"
    })


# ---------------------------------------------------------------------------
# Financial Years
# ---------------------------------------------------------------------------
@login_required
def financial_year_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = FinancialYearForm(request.POST)
        if form.is_valid():
            fy = form.save(commit=False)
            fy.entity = entity
            # Link to prior year if exists
            prior = entity.financial_years.order_by("-end_date").first()
            if prior:
                fy.prior_year = prior
            fy.save()
            _log_action(request, "import", f"Created financial year: {fy.year_label}", fy)
            messages.success(request, f"Financial year '{fy.year_label}' created.")
            return redirect("core:financial_year_detail", pk=fy.pk)
    else:
        form = FinancialYearForm()
    return render(request, "core/financial_year_form.html", {
        "form": form, "entity": entity, "title": "Create Financial Year"
    })


@login_required
def financial_year_detail(request, pk):
    fy = get_financial_year_for_user(request, pk)
    # Audit log: data access
    _log_action(request, "view", f"Viewed financial year: {fy.year_label} for {fy.entity.entity_name}", fy)
    tb_lines = fy.trial_balance_lines.select_related("mapped_line_item").all()
    adjustments = fy.adjusting_journals.all()
    unmapped_count = tb_lines.filter(mapped_line_item__isnull=True).count()
    documents = fy.generated_documents.all().order_by('-version', '-generated_at')

    # Calculate totals
    total_debit = tb_lines.aggregate(total=Sum("debit"))["total"] or Decimal("0")
    total_credit = tb_lines.aggregate(total=Sum("credit"))["total"] or Decimal("0")
    total_prior_debit = tb_lines.aggregate(total=Sum("prior_debit"))["total"] or Decimal("0")
    total_prior_credit = tb_lines.aggregate(total=Sum("prior_credit"))["total"] or Decimal("0")

    # Group lines by section for comparative display
    from collections import OrderedDict
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }
    sections = OrderedDict()
    for line in tb_lines:
        if line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = 'Unmapped'
        if display_section not in sections:
            sections[display_section] = []
        sections[display_section].append(line)
    ordered_sections = OrderedDict()
    section_keys_ordered = []
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in section_keys_ordered:
            section_keys_ordered.append(ds)
    for key in section_keys_ordered:
        if key in sections:
            ordered_sections[key] = sections[key]
    for key in sections:
        if key not in ordered_sections:
            ordered_sections[key] = sections[key]

    # Year labels
    current_year = str(fy.year_label)
    # Extract year number from label like "FY2026" or "2026"
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year = str(fy.prior_year.year_label)
    else:
        prior_year = 'Prior'

    # Audit Risk flags
    risk_flags = RiskFlag.objects.filter(financial_year=fy).order_by('-severity', '-created_at')
    open_risk_count = risk_flags.filter(status='open').count()

    # Build a lookup: account_code -> list of open risk flag titles/descriptions
    # affected_accounts is a JSON list of dicts: [{"account_code": "2992", "account_name": "Clearing", ...}]
    flagged_accounts = {}  # {account_code: [{"title": ..., "severity": ..., "description": ...}]}
    for flag in risk_flags.filter(status__in=['open', 'reviewed']):
        for acc in (flag.affected_accounts or []):
            code = acc.get('account_code', '') if isinstance(acc, dict) else str(acc)
            if code:
                if code not in flagged_accounts:
                    flagged_accounts[code] = []
                flagged_accounts[code].append({
                    'title': flag.title,
                    'severity': flag.severity,
                    'description': flag.description[:120],
                })

    # Annotate TB lines with risk flag info and variance calculations
    for line in tb_lines:
        line.risk_flags_list = flagged_accounts.get(line.account_code, [])
        # Variance calculations for Prior Year Comparatives Engine
        current_net = (line.debit or Decimal('0')) - (line.credit or Decimal('0'))
        prior_net = (line.prior_debit or Decimal('0')) - (line.prior_credit or Decimal('0'))
        line.variance_amount = current_net - prior_net
        if prior_net != 0:
            line.variance_percentage = round(((current_net - prior_net) / abs(prior_net)) * 100, 1)
        else:
            line.variance_percentage = None

    # Depreciation assets
    depreciation_assets = DepreciationAsset.objects.filter(financial_year=fy)
    dep_categories = OrderedDict()
    dep_total_opening = Decimal('0')
    dep_total_depreciation = Decimal('0')
    dep_total_closing = Decimal('0')
    for asset in depreciation_assets:
        if asset.category not in dep_categories:
            dep_categories[asset.category] = []
        dep_categories[asset.category].append(asset)
        dep_total_opening += asset.opening_wdv
        dep_total_depreciation += asset.depreciation_amount
        dep_total_closing += asset.closing_wdv

    # Stock items
    stock_items = StockItem.objects.filter(financial_year=fy)
    stock_total_opening = stock_items.aggregate(total=Sum('opening_value'))['total'] or Decimal('0')
    stock_total_closing = stock_items.aggregate(total=Sum('closing_value'))['total'] or Decimal('0')

    # Review items (pending transactions from bank statement uploads for this entity)
    from review.models import PendingTransaction, ReviewJob
    review_jobs = ReviewJob.objects.filter(entity=fy.entity)
    pending_review = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=False,
    ).select_related('job').order_by('date')
    confirmed_review = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=True,
    ).select_related('job').order_by('date')

    # Activity / Audit trail — all AuditLog entries for this financial year
    activity_logs = AuditLog.objects.filter(
        affected_object_id=str(fy.pk),
    ).select_related('user').order_by('-timestamp')
    # Also include journal-specific logs (where affected_object_id is a journal PK)
    journal_pks = list(fy.adjusting_journals.values_list('pk', flat=True))
    journal_logs = AuditLog.objects.filter(
        affected_object_id__in=[str(pk) for pk in journal_pks],
    ).select_related('user').order_by('-timestamp') if journal_pks else AuditLog.objects.none()
    # Merge and deduplicate, ordered by timestamp descending
    from itertools import chain
    all_log_pks = set(activity_logs.values_list('pk', flat=True)) | set(journal_logs.values_list('pk', flat=True))
    activity_logs = AuditLog.objects.filter(pk__in=all_log_pks).select_related('user').order_by('-timestamp')

    context = {
        "fy": fy,
        "entity": fy.entity,
        "tb_lines": tb_lines,
        "tb_sections": ordered_sections,
        "adjustments": adjustments,
        "unmapped_count": unmapped_count,
        "documents": documents,
        "total_debit": total_debit,
        "total_credit": total_credit,
        "total_prior_debit": total_prior_debit,
        "total_prior_credit": total_prior_credit,
        "current_year": current_year,
        "prior_year": prior_year,
        # Audit Risk
        "risk_flags": risk_flags,
        "open_risk_count": open_risk_count,
        "flagged_accounts": flagged_accounts,
        # Depreciation
        "depreciation_assets": depreciation_assets,
        "dep_categories": dep_categories,
        "dep_total_opening": dep_total_opening,
        "dep_total_depreciation": dep_total_depreciation,
        "dep_total_closing": dep_total_closing,
        # Stock
        "stock_items": stock_items,
        "stock_total_opening": stock_total_opening,
        "stock_total_closing": stock_total_closing,
        # Review
        "pending_review": pending_review,
        "confirmed_review": confirmed_review,
        # Activity
        "activity_logs": activity_logs,
    }
    return render(request, "core/financial_year_detail.html", context)


@login_required
def financial_year_status(request, pk):
    """Change the status of a financial year."""
    fy = get_financial_year_for_user(request, pk)
    new_status = request.POST.get("status")

    if not new_status or new_status not in dict(FinancialYear.Status.choices):
        messages.error(request, "Invalid status.")
        return redirect("core:financial_year_detail", pk=pk)

    # Permission checks
    if new_status == "finalised" and not request.user.can_finalise:
        messages.error(request, "Only senior accountants can finalise.")
        return redirect("core:financial_year_detail", pk=pk)

    if new_status == "reviewed" and not request.user.is_senior:
        messages.error(request, "Only senior accountants can mark as reviewed.")
        return redirect("core:financial_year_detail", pk=pk)

    # Finalisation Gate: enforce risk engine completion and flag resolution
    if new_status == "finalised":
        override_reason = request.POST.get("override_reason", "").strip()
        force_override = request.POST.get("force_override") == "1"

        # Gate 1: HARD BLOCK — Risk engine must have been run at least once
        total_flags = fy.risk_flags.count()
        if total_flags == 0:
            messages.error(
                request,
                "BLOCKED: The Risk Engine has not been run for this financial year. "
                "This is a mandatory step. Please run the Risk Engine from the Audit Risk tab."
            )
            return redirect("core:financial_year_detail", pk=pk)

        # Gate 2: HARD BLOCK — No CRITICAL flags can remain open (no override)
        critical_open = fy.risk_flags.filter(status="open", severity="CRITICAL").count()
        if critical_open > 0:
            messages.error(
                request,
                f"BLOCKED: {critical_open} CRITICAL risk flag(s) remain unresolved. "
                f"Critical flags cannot be overridden — they must be resolved before finalisation. "
                f"These include Division 7A breaches, solvency issues, and TB imbalances."
            )
            return redirect("core:financial_year_detail", pk=pk)

        # Gate 3: WARN — HIGH flags can be overridden with a reason logged to audit trail
        high_open = fy.risk_flags.filter(status="open", severity="HIGH").count()
        if high_open > 0:
            if not force_override or not override_reason:
                messages.warning(
                    request,
                    f"WARNING: {high_open} HIGH severity risk flag(s) remain unresolved. "
                    f"You may override this gate by providing a reason. "
                    f"This will be logged to the audit trail for partner review."
                )
                # Set a session flag so the template can show the override form
                request.session["finalisation_gate_warn"] = {
                    "high_open": high_open,
                    "fy_pk": pk,
                }
                return redirect("core:financial_year_detail", pk=pk)
            else:
                # Override accepted — log to audit trail
                _log_action(
                    request, "finalisation_override",
                    f"Finalisation gate overridden for {high_open} HIGH flags. "
                    f"Reason: {override_reason}",
                    fy
                )
                # Mark the HIGH flags as reviewed with override note
                fy.risk_flags.filter(status="open", severity="HIGH").update(
                    status="reviewed",
                    resolution_notes=f"Overridden at finalisation by {request.user.get_full_name()}. "
                                     f"Reason: {override_reason}",
                    resolved_at=timezone.now(),
                    resolved_by=request.user,
                )
                # Clear the session flag
                request.session.pop("finalisation_gate_warn", None)

        # Gate 4: WARN — MEDIUM/LOW flags can be overridden with a reason
        open_medium_low = fy.risk_flags.filter(
            status="open", severity__in=["MEDIUM", "LOW"]
        ).count()
        if open_medium_low > 0:
            if not force_override or not override_reason:
                messages.warning(
                    request,
                    f"WARNING: {open_medium_low} MEDIUM/LOW risk flag(s) remain open. "
                    f"You may override this gate by providing a reason."
                )
                request.session["finalisation_gate_warn"] = {
                    "medium_low_open": open_medium_low,
                    "fy_pk": pk,
                }
                return redirect("core:financial_year_detail", pk=pk)
            else:
                # Override accepted — log to audit trail
                _log_action(
                    request, "finalisation_override",
                    f"Finalisation gate overridden for {open_medium_low} MEDIUM/LOW flags. "
                    f"Reason: {override_reason}",
                    fy
                )
                fy.risk_flags.filter(
                    status="open", severity__in=["MEDIUM", "LOW"]
                ).update(
                    status="reviewed",
                    resolution_notes=f"Overridden at finalisation by {request.user.get_full_name()}. "
                                     f"Reason: {override_reason}",
                    resolved_at=timezone.now(),
                    resolved_by=request.user,
                )
                request.session.pop("finalisation_gate_warn", None)

    old_status = fy.status
    fy.status = new_status
    if new_status == "reviewed":
        fy.reviewed_by = request.user
    if new_status == "finalised":
        fy.finalised_at = timezone.now()
        # Lock comparatives and mark final documents
        fy.trial_balance_lines.update(comparatives_locked=True)
        # Lock and mark the latest version of each document type as final
        from django.db.models import Max
        for doc_type in fy.generated_documents.values_list('document_type', flat=True).distinct():
            latest = fy.generated_documents.filter(document_type=doc_type).order_by('-version').first()
            if latest:
                latest.status = 'final'
                latest.is_locked = True
                latest.save(update_fields=['status', 'is_locked'])
    fy.save()

    _log_action(
        request, "status_change",
        f"Changed {fy} from {old_status} to {new_status}", fy
    )
    messages.success(request, f"Status changed to {fy.get_status_display()}.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def roll_forward(request, pk):
    """Create a new financial year from the current one, carrying closing balances."""
    current_fy = get_financial_year_for_user(request, pk)
    entity = current_fy.entity

    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    # Cannot roll forward unless the current FY is finalised
    if current_fy.status != "finalised":
        messages.error(
            request,
            "Cannot roll forward: this financial year must be finalised before rolling forward."
        )
        return redirect("core:financial_year_detail", pk=pk)

    if request.method == "POST":
        # Calculate new dates (add 1 year)
        from dateutil.relativedelta import relativedelta
        new_start = current_fy.end_date + relativedelta(days=1)
        new_end = current_fy.end_date + relativedelta(years=1)
        new_label = f"FY{new_end.year}"

        # Check if already exists
        if entity.financial_years.filter(year_label=new_label).exists():
            messages.warning(request, f"{new_label} already exists for this entity.")
            return redirect("core:financial_year_detail", pk=pk)

        # Create new FY
        new_fy = FinancialYear.objects.create(
            entity=entity,
            year_label=new_label,
            start_date=new_start,
            end_date=new_end,
            prior_year=current_fy,
            status=FinancialYear.Status.DRAFT,
        )

        # Determine which sections are balance sheet
        BS_SECTIONS = {"assets", "liabilities", "equity", "capital_accounts"}
        BS_STATEMENTS = {"balance_sheet", "equity"}

        # Build lookup of account_code -> section from ChartOfAccount
        coa_sections = dict(
            ChartOfAccount.objects.filter(
                entity_type=entity.entity_type, is_active=True
            ).values_list("account_code", "section")
        )

        carried = 0
        for line in current_fy.trial_balance_lines.filter(is_adjustment=False):
            # Determine if this is a balance sheet account
            is_bs = False
            if line.mapped_line_item:
                is_bs = line.mapped_line_item.financial_statement in BS_STATEMENTS
            elif line.account_code in coa_sections:
                is_bs = coa_sections[line.account_code] in BS_SECTIONS
            else:
                # Fallback: check if account code starts with 2/3 (common for BS)
                # or if the section keyword is in the account name
                code_prefix = line.account_code.split(".")[0] if line.account_code else ""
                if code_prefix.isdigit():
                    is_bs = int(code_prefix) >= 2000

            if not is_bs:
                continue  # Skip P&L accounts (revenue, expenses, COGS)

            # Only carry forward if there's actually a balance
            if line.closing_balance == 0:
                continue

            TrialBalanceLine.objects.create(
                financial_year=new_fy,
                account_code=line.account_code,
                account_name=line.account_name,
                opening_balance=line.closing_balance,
                debit=Decimal("0"),
                credit=Decimal("0"),
                closing_balance=line.closing_balance,  # = opening, waiting for movement
                prior_debit=line.debit,
                prior_credit=line.credit,
                mapped_line_item=line.mapped_line_item,
                is_adjustment=False,
            )
            carried += 1

        _log_action(request, "import", f"Rolled forward to {new_label} with {carried} balance sheet items", new_fy)
        messages.success(request, f"Rolled forward to {new_label}. {carried} balance sheet items carried with opening balances and comparatives.")
        return redirect("core:financial_year_detail", pk=new_fy.pk)

    return render(request, "core/roll_forward_confirm.html", {"fy": current_fy})


# ---------------------------------------------------------------------------
# Trial Balance Import
# ---------------------------------------------------------------------------
@login_required
def trial_balance_import(request, pk):
    fy = get_financial_year_for_user(request, pk)

    if fy.is_locked:
        messages.error(request, "Cannot import into a finalised financial year.")
        return redirect("core:financial_year_detail", pk=pk)

    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    if request.method == "POST":
        form = TrialBalanceUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES["file"]

            # Validate file type and size
            if uploaded_file.size > 20 * 1024 * 1024:
                messages.error(request, "File too large. Maximum size is 20MB.")
                return redirect("core:financial_year_detail", pk=pk)

            import os as _os
            file_ext = _os.path.splitext(uploaded_file.name)[1].lower()
            if file_ext not in {".xlsx", ".xls"}:
                messages.error(request, f"Unsupported file type: {file_ext}. Only Excel files (.xlsx) are supported.")
                return redirect("core:financial_year_detail", pk=pk)

            file = uploaded_file
            try:
                result = _process_trial_balance_upload(fy, file)
                _log_action(request, "import", f"Imported trial balance: {result['imported']} lines", fy)
                messages.success(
                    request,
                    f"Imported {result['imported']} lines. "
                    f"{result['unmapped']} unmapped accounts need attention."
                )
                if result["errors"]:
                    for err in result["errors"][:5]:
                        messages.warning(request, err)
                return redirect("core:financial_year_detail", pk=pk)
            except Exception as e:
                messages.error(request, "Import failed. Please check the file format and try again.")
    else:
        form = TrialBalanceUploadForm()

    return render(request, "core/trial_balance_import.html", {
        "form": form, "fy": fy
    })


def _process_trial_balance_upload(fy, file):
    """Process an Excel trial balance upload."""
    wb = openpyxl.load_workbook(file, read_only=True, data_only=True)
    ws = wb.active

    # Clear existing non-adjustment lines
    fy.trial_balance_lines.filter(is_adjustment=False).delete()

    imported = 0
    unmapped = 0
    errors = []
    entity = fy.entity

    rows = list(ws.iter_rows(min_row=2, values_only=True))  # Skip header

    for i, row in enumerate(rows, start=2):
        if not row or not row[0]:
            continue

        try:
            account_code = str(row[0]).strip()
            account_name = str(row[1]).strip() if row[1] else ""
            opening_balance = Decimal(str(row[2] or 0))
            debit = Decimal(str(row[3] or 0))
            credit = Decimal(str(row[4] or 0))
            closing_balance = opening_balance + debit - credit
        except (InvalidOperation, ValueError, IndexError) as e:
            errors.append(f"Row {i}: {str(e)}")
            continue

        # Try to find existing mapping for this entity (learning from prior imports)
        mapping = ClientAccountMapping.objects.filter(
            entity=entity, client_account_code=account_code
        ).first()

        mapped_item = mapping.mapped_line_item if mapping else None

        # If no existing mapping, try to match against the standard chart of accounts
        # for this entity type. This provides automatic mapping for known account codes.
        coa_match = None
        if not mapped_item:
            coa_match = ChartOfAccount.objects.filter(
                entity_type=entity.entity_type,
                account_code=account_code,
                is_active=True,
            ).first()
            if coa_match and coa_match.maps_to:
                mapped_item = coa_match.maps_to

        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code=account_code,
            account_name=account_name,
            opening_balance=opening_balance,
            debit=debit,
            credit=credit,
            closing_balance=closing_balance,
            mapped_line_item=mapped_item,
            is_adjustment=False,
        )

        # Create or update client account mapping record
        ClientAccountMapping.objects.update_or_create(
            entity=entity,
            client_account_code=account_code,
            defaults={"client_account_name": account_name, "mapped_line_item": mapped_item},
        )

        imported += 1
        if not mapped_item:
            unmapped += 1

    wb.close()
    return {"imported": imported, "unmapped": unmapped, "errors": errors}


@login_required
def trial_balance_view(request, pk):
    """Preview trial balance with comparative columns, grouped by section."""
    from collections import OrderedDict
    fy = get_financial_year_for_user(request, pk)
    tb_lines = fy.trial_balance_lines.select_related("mapped_line_item").order_by("account_code")

    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }

    sections = OrderedDict()
    grand_total_dr = Decimal('0')
    grand_total_cr = Decimal('0')
    grand_total_prior_dr = Decimal('0')
    grand_total_prior_cr = Decimal('0')

    for line in tb_lines:
        if line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = 'Unmapped'
        if display_section not in sections:
            sections[display_section] = []
        sections[display_section].append(line)
        grand_total_dr += line.debit or Decimal('0')
        grand_total_cr += line.credit or Decimal('0')
        grand_total_prior_dr += line.prior_debit or Decimal('0')
        grand_total_prior_cr += line.prior_credit or Decimal('0')

    # Sort sections by defined order
    ordered_sections = OrderedDict()
    seen = set()
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in seen and ds in sections:
            ordered_sections[ds] = sections[ds]
            seen.add(ds)
    for key in sections:
        if key not in ordered_sections:
            ordered_sections[key] = sections[key]

    # Year labels
    current_year_label = str(fy.year_label)
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year_label = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year_label = str(fy.prior_year.year_label)
    else:
        prior_year_label = 'Prior'

    return render(request, "core/trial_balance_view.html", {
        "fy": fy,
        "sections": ordered_sections,
        "current_year_label": current_year_label,
        "prior_year_label": prior_year_label,
        "grand_total_dr": grand_total_dr,
        "grand_total_cr": grand_total_cr,
        "grand_total_prior_dr": grand_total_prior_dr,
        "grand_total_prior_cr": grand_total_prior_cr,
    })


# ---------------------------------------------------------------------------
# Account Mapping
# ---------------------------------------------------------------------------
@login_required
def account_mapping_list(request):
    mappings = AccountMapping.objects.all()
    return render(request, "core/account_mapping_list.html", {"mappings": mappings})


@login_required
def account_mapping_create(request):
    if not request.user.is_admin:
        messages.error(request, "Only administrators can manage standard mappings.")
        return redirect("core:account_mapping_list")

    if request.method == "POST":
        form = AccountMappingForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Account mapping created.")
            return redirect("core:account_mapping_list")
    else:
        form = AccountMappingForm()
    return render(request, "core/account_mapping_form.html", {"form": form, "title": "Create Standard Mapping"})


@login_required
def map_client_accounts(request, pk):
    """Map unmapped client accounts to standard line items for a financial year."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    unmapped = ClientAccountMapping.objects.filter(
        entity=entity, mapped_line_item__isnull=True
    )
    standard_mappings = AccountMapping.objects.all()

    if request.method == "POST":
        mapped_count = 0
        for cam in unmapped:
            mapping_id = request.POST.get(f"mapping_{cam.pk}")
            if mapping_id:
                try:
                    cam.mapped_line_item = AccountMapping.objects.get(pk=mapping_id)
                    cam.save()
                    # Update trial balance lines with this mapping
                    TrialBalanceLine.objects.filter(
                        financial_year=fy,
                        account_code=cam.client_account_code,
                    ).update(mapped_line_item=cam.mapped_line_item)
                    mapped_count += 1
                except AccountMapping.DoesNotExist:
                    pass

        _log_action(request, "mapping_change", f"Mapped {mapped_count} accounts", fy)
        messages.success(request, f"Mapped {mapped_count} accounts.")
        return redirect("core:financial_year_detail", pk=pk)

    context = {
        "fy": fy,
        "unmapped": unmapped,
        "standard_mappings": standard_mappings,
    }
    return render(request, "core/map_client_accounts.html", context)


# ---------------------------------------------------------------------------
# Journal Entries (Enhanced)
# ---------------------------------------------------------------------------
@login_required
def adjustment_list(request, pk):
    fy = get_financial_year_for_user(request, pk)
    adjustments = fy.adjusting_journals.prefetch_related("lines").all()
    return render(request, "core/adjustment_list.html", {"fy": fy, "adjustments": adjustments})


@login_required
def adjustment_create(request, pk):
    fy = get_financial_year_for_user(request, pk)

    if fy.is_locked:
        messages.error(request, "Cannot add journals to a finalised year.")
        return redirect("core:financial_year_detail", pk=pk)

    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    # Get available accounts for the account picker
    # Always use the master Chart of Accounts for the entity type
    # so journals can be created even before a trial balance is imported
    entity_type = fy.entity.entity_type
    master_accounts = list(
        ChartOfAccount.objects.filter(entity_type=entity_type, is_active=True)
        .order_by("account_code")
        .values("account_code", "account_name")
    )
    # Rename fields to match what the template JS expects
    accounts = [
        {"client_account_code": a["account_code"], "client_account_name": a["account_name"]}
        for a in master_accounts
    ]

    if request.method == "POST":
        form = AdjustingJournalForm(request.POST)
        formset = JournalLineFormSet(request.POST)
        if form.is_valid() and formset.is_valid():
            journal = form.save(commit=False)
            journal.financial_year = fy
            journal.created_by = request.user
            journal.save()  # This auto-generates reference_number

            formset.instance = journal
            lines = formset.save()

            # Validate debits = credits
            total_dr = sum(l.debit for l in lines)
            total_cr = sum(l.credit for l in lines)
            if total_dr != total_cr:
                journal.delete()
                messages.error(request, f"Journal does not balance: Dr ${total_dr:,.2f} \u2260 Cr ${total_cr:,.2f}")
                return render(request, "core/adjustment_form.html", {
                    "form": form, "formset": formset, "fy": fy, "accounts": accounts
                })

            # Update cached totals
            journal.total_debit = total_dr
            journal.total_credit = total_cr
            journal.save(update_fields=["total_debit", "total_credit"])

            _log_action(request, "adjustment", f"Created journal {journal.reference_number}: {journal.description}", journal)
            messages.success(request, f"Journal {journal.reference_number} created as Draft.")
            return redirect("core:financial_year_detail", pk=pk)
    else:
        form = AdjustingJournalForm(initial={"journal_date": fy.end_date})
        formset = JournalLineFormSet()

    return render(request, "core/adjustment_form.html", {
        "form": form, "formset": formset, "fy": fy, "accounts": accounts
    })


@login_required
def journal_detail(request, pk):
    """View a single journal entry with all its lines and audit info."""
    journal = get_object_or_404(
        AdjustingJournal.objects.select_related(
            "financial_year", "financial_year__entity",
            "created_by", "posted_by"
        ).prefetch_related("lines"),
        pk=pk,
    )
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check
    entity = fy.entity
    return render(request, "core/journal_detail.html", {
        "journal": journal, "fy": fy, "entity": entity,
    })


@login_required
def journal_post(request, pk):
    """Post a draft journal — creates adjustment TB lines."""
    journal = get_object_or_404(AdjustingJournal, pk=pk)
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if journal.status != AdjustingJournal.JournalStatus.DRAFT:
        messages.error(request, "Only draft journals can be posted.")
        return redirect("core:journal_detail", pk=pk)

    if not journal.is_balanced:
        messages.error(request, "Journal does not balance. Cannot post.")
        return redirect("core:journal_detail", pk=pk)

    if fy.is_locked:
        messages.error(request, "Cannot post to a finalised year.")
        return redirect("core:journal_detail", pk=pk)

    # Create adjustment trial balance lines
    for line in journal.lines.all():
        mapping = ClientAccountMapping.objects.filter(
            entity=fy.entity, client_account_code=line.account_code
        ).first()
        mapped_item = mapping.mapped_line_item if mapping else None

        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code=line.account_code,
            account_name=line.account_name,
            opening_balance=Decimal("0"),
            debit=line.debit,
            credit=line.credit,
            closing_balance=line.debit - line.credit,
            mapped_line_item=mapped_item,
            is_adjustment=True,
        )

    # Update journal status
    journal.status = AdjustingJournal.JournalStatus.POSTED
    journal.posted_by = request.user
    journal.posted_at = timezone.now()
    journal.save(update_fields=["status", "posted_by", "posted_at"])

    _log_action(request, "adjustment", f"Posted journal {journal.reference_number}", journal)
    messages.success(request, f"Journal {journal.reference_number} has been posted.")
    return redirect("core:financial_year_detail", pk=fy.pk)


@login_required
def journal_delete(request, pk):
    """Delete a journal entry. If posted, also removes its adjustment TB lines."""
    journal = get_object_or_404(AdjustingJournal, pk=pk)
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if fy.is_locked:
        messages.error(request, "Cannot delete journals in a finalised year.")
        return redirect("core:journal_detail", pk=pk)

    if not request.user.can_edit:
        messages.error(request, "You do not have permission to delete journals.")
        return redirect("core:journal_detail", pk=pk)

    ref = journal.reference_number
    status = journal.status

    # If the journal was posted, remove its adjustment TB lines
    if status == AdjustingJournal.JournalStatus.POSTED:
        for line in journal.lines.all():
            TrialBalanceLine.objects.filter(
                financial_year=fy,
                account_code=line.account_code,
                debit=line.debit,
                credit=line.credit,
                is_adjustment=True,
            ).first().delete() if TrialBalanceLine.objects.filter(
                financial_year=fy,
                account_code=line.account_code,
                debit=line.debit,
                credit=line.credit,
                is_adjustment=True,
            ).exists() else None

    journal.delete()
    _log_action(request, "adjustment", f"Deleted {status} journal {ref}")
    messages.success(request, f"Journal {ref} has been deleted.")
    return redirect("core:financial_year_detail", pk=fy.pk)


@login_required
def account_list_api(request, pk):
    """JSON API endpoint returning available accounts for a financial year's entity."""
    fy = get_financial_year_for_user(request, pk)
    accounts = list(
        ClientAccountMapping.objects.filter(entity=fy.entity)
        .order_by("client_account_code")
        .values("client_account_code", "client_account_name")
    )
    # Also include accounts from the trial balance that may not be mapped
    tb_accounts = list(
        fy.trial_balance_lines.filter(is_adjustment=False)
        .values("account_code", "account_name")
        .distinct()
        .order_by("account_code")
    )
    # Merge: use TB accounts as base, supplement with mapped accounts
    account_dict = {}
    for a in tb_accounts:
        account_dict[a["account_code"]] = a["account_name"]
    for a in accounts:
        if a["client_account_code"] not in account_dict:
            account_dict[a["client_account_code"]] = a["client_account_name"]

    result = [
        {"code": code, "name": name}
        for code, name in sorted(account_dict.items())
    ]
    return JsonResponse(result, safe=False)


# ---------------------------------------------------------------------------
# Financial Statements Preview
# ---------------------------------------------------------------------------
@login_required
def financial_statements_view(request, pk):
    """Render a preview of the financial statements on screen."""
    get_financial_year_for_user(request, pk)  # IDOR check
    fy = get_object_or_404(
        FinancialYear.objects.select_related("entity", "prior_year"),
        pk=pk,
    )

    # Aggregate trial balance by mapped line item
    from django.db.models import Sum as DSum
    current_data = (
        fy.trial_balance_lines
        .filter(mapped_line_item__isnull=False)
        .values(
            "mapped_line_item__standard_code",
            "mapped_line_item__line_item_label",
            "mapped_line_item__financial_statement",
            "mapped_line_item__statement_section",
            "mapped_line_item__display_order",
        )
        .annotate(total=DSum("closing_balance"))
        .order_by("mapped_line_item__display_order")
    )

    # Get prior year data if available
    prior_data = {}
    if fy.prior_year:
        prior_lines = (
            fy.prior_year.trial_balance_lines
            .filter(mapped_line_item__isnull=False)
            .values("mapped_line_item__standard_code")
            .annotate(total=DSum("closing_balance"))
        )
        prior_data = {item["mapped_line_item__standard_code"]: item["total"] for item in prior_lines}

    # Organise into statements
    income_statement = []
    balance_sheet = []
    for item in current_data:
        entry = {
            "code": item["mapped_line_item__standard_code"],
            "label": item["mapped_line_item__line_item_label"],
            "section": item["mapped_line_item__statement_section"],
            "current": item["total"],
            "prior": prior_data.get(item["mapped_line_item__standard_code"], Decimal("0")),
        }
        if item["mapped_line_item__financial_statement"] == "income_statement":
            income_statement.append(entry)
        elif item["mapped_line_item__financial_statement"] == "balance_sheet":
            balance_sheet.append(entry)

    context = {
        "fy": fy,
        "entity": fy.entity,
        "income_statement": income_statement,
        "balance_sheet": balance_sheet,
    }
    return render(request, "core/financial_statements_view.html", context)


# ---------------------------------------------------------------------------
# Document Generation (placeholder for Phase 2)
# ---------------------------------------------------------------------------
@login_required
def generate_document(request, pk):
    fy = get_financial_year_for_user(request, pk)

    # Check that there are trial balance lines
    if not fy.trial_balance_lines.exists():
        messages.error(request, "Cannot generate statements: no trial balance data loaded.")
        return redirect("core:financial_year_detail", pk=pk)

    # Determine requested format (default: docx)
    fmt = request.GET.get("format", "docx").lower()
    if fmt not in ("docx", "pdf"):
        fmt = "docx"

    from .docgen import generate_financial_statements

    # Check for open audit risk flags — if any exist, watermark the document
    has_open_risks = fy.risk_flags.filter(status="open").exists()

    try:
        buffer = generate_financial_statements(fy.pk, has_open_risks=has_open_risks)
    except Exception as e:
        messages.error(request, f"Document generation failed: {e}")
        return redirect("core:financial_year_detail", pk=pk)

    # Build filename
    entity_name = fy.entity.entity_name.replace(" ", "_")
    base_filename = f"{entity_name}_Financial_Statements_{fy.year_label}"

    if fmt == "pdf":
        # Convert DOCX buffer to PDF via LibreOffice (soffice)
        import subprocess, tempfile, os
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, f"{base_filename}.docx")
                with open(docx_path, "wb") as f:
                    f.write(buffer.getvalue())

                # Try multiple LibreOffice binary names
                lo_bin = None
                for candidate in ["soffice", "libreoffice", "/usr/bin/soffice", "/usr/bin/libreoffice"]:
                    try:
                        subprocess.run([candidate, "--version"], capture_output=True, timeout=5)
                        lo_bin = candidate
                        break
                    except (FileNotFoundError, subprocess.TimeoutExpired):
                        continue

                if not lo_bin:
                    raise RuntimeError(
                        "LibreOffice is not installed. Install with: sudo apt-get install -y libreoffice-writer"
                    )

                result = subprocess.run(
                    [lo_bin, "--headless", "--norestore", "--convert-to", "pdf",
                     "--outdir", tmpdir, docx_path],
                    capture_output=True, timeout=120,
                    env={**os.environ, "HOME": tmpdir},
                )
                pdf_path = os.path.join(tmpdir, f"{base_filename}.pdf")
                if not os.path.exists(pdf_path):
                    stderr = result.stderr.decode('utf-8', errors='replace')
                    stdout = result.stdout.decode('utf-8', errors='replace')
                    raise RuntimeError(
                        f"PDF conversion failed (exit code {result.returncode}).\n"
                        f"stdout: {stdout[:500]}\nstderr: {stderr[:500]}"
                    )
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()

            filename = f"{base_filename}.pdf"
            response = HttpResponse(
                pdf_bytes,
                content_type="application/pdf",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = pdf_bytes
            file_format = "pdf"
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF conversion failed: {e}")
            messages.error(request, f"PDF conversion failed: {e}. Falling back to DOCX.")
            filename = f"{base_filename}.docx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = buffer.getvalue()
            file_format = "docx"
    else:
        filename = f"{base_filename}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        file_content = buffer.getvalue()
        file_format = "docx"

    # Log the generation
    _log_action(request, "generate", f"Generated financial statements ({file_format.upper()}) for {fy}", fy)

    # Save record
    from django.core.files.base import ContentFile
    doc = GeneratedDocument(
        financial_year=fy,
        file_format=file_format,
        generated_by=request.user,
    )
    doc.file.save(filename, ContentFile(file_content), save=True)

    return response


# ---------------------------------------------------------------------------
# Distribution Minutes Generation
# ---------------------------------------------------------------------------
@login_required
def generate_distribution_minutes(request, pk):
    """Generate distribution minutes document for a trust entity."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    # Only trusts have distribution minutes
    if entity.entity_type != 'trust':
        messages.error(request, "Distribution minutes are only applicable to trust entities.")
        return redirect("core:financial_year_detail", pk=pk)

    fmt = request.GET.get("format", "docx").lower()
    if fmt not in ("docx", "pdf"):
        fmt = "docx"

    from .distmin_gen import generate_distribution_minutes as gen_distmin

    try:
        buffer = gen_distmin(fy.pk)
    except (ValueError, FileNotFoundError) as e:
        messages.error(request, f"Distribution minutes generation failed: {e}")
        return redirect("core:financial_year_detail", pk=pk)
    except Exception as e:
        messages.error(request, f"Unexpected error generating distribution minutes: {e}")
        return redirect("core:financial_year_detail", pk=pk)

    entity_name = entity.entity_name.replace(" ", "_")
    base_filename = f"{entity_name}_Distribution_Minutes_{fy.year_label}"

    if fmt == "pdf":
        import subprocess, tempfile, os
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, f"{base_filename}.docx")
                with open(docx_path, "wb") as f:
                    f.write(buffer.getvalue())

                lo_bin = None
                for candidate in ["soffice", "libreoffice", "/usr/bin/soffice", "/usr/bin/libreoffice"]:
                    try:
                        subprocess.run([candidate, "--version"], capture_output=True, timeout=5)
                        lo_bin = candidate
                        break
                    except (FileNotFoundError, subprocess.TimeoutExpired):
                        continue

                if not lo_bin:
                    raise RuntimeError("LibreOffice is not installed.")

                result = subprocess.run(
                    [lo_bin, "--headless", "--norestore", "--convert-to", "pdf",
                     "--outdir", tmpdir, docx_path],
                    capture_output=True, timeout=120,
                    env={**os.environ, "HOME": tmpdir},
                )
                pdf_path = os.path.join(tmpdir, f"{base_filename}.pdf")
                if not os.path.exists(pdf_path):
                    raise RuntimeError("PDF conversion failed.")
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()

            filename = f"{base_filename}.pdf"
            response = HttpResponse(pdf_bytes, content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = pdf_bytes
            file_format = "pdf"
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF conversion failed: {e}")
            messages.error(request, f"PDF conversion failed: {e}. Falling back to DOCX.")
            filename = f"{base_filename}.docx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = buffer.getvalue()
            file_format = "docx"
    else:
        filename = f"{base_filename}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        file_content = buffer.getvalue()
        file_format = "docx"

    _log_action(request, "generate", f"Generated distribution minutes ({file_format.upper()}) for {fy}", fy)

    # Save record
    from django.core.files.base import ContentFile
    doc = GeneratedDocument(
        financial_year=fy,
        file_format=file_format,
        generated_by=request.user,
    )
    doc.file.save(filename, ContentFile(file_content), save=True)

    return response


# ---------------------------------------------------------------------------
# HTMX Partials
# ---------------------------------------------------------------------------
@login_required
def htmx_client_search(request):
    """HTMX search endpoint for entities (replaces client search)."""
    query = request.GET.get("q", "")
    if len(query) < 2:
        return HttpResponse("")

    if request.user.is_admin:
        entities = Entity.objects.filter(is_archived=False)
    else:
        entities = Entity.objects.filter(
            assigned_accountant=request.user, is_archived=False
        )

    entities = entities.filter(
        Q(entity_name__icontains=query)
        | Q(abn__icontains=query)
        | Q(trading_as__icontains=query)
    ).distinct().annotate(num_fys=Count("financial_years"))[:20]

    return render(request, "partials/entity_list_rows.html", {"entities": entities})


@login_required
def htmx_map_tb_line(request, pk):
    """HTMX endpoint to map a single trial balance line."""
    line = get_object_or_404(TrialBalanceLine, pk=pk)
    get_financial_year_for_user(request, line.financial_year.pk)  # IDOR check
    if not request.user.can_edit:
        return HttpResponse("Permission denied", status=403)
    mapping_id = request.POST.get("mapped_line_item")

    if mapping_id:
        try:
            mapping = AccountMapping.objects.get(pk=mapping_id)
            line.mapped_line_item = mapping
            line.save()

            # Also save to client account mapping for reuse
            ClientAccountMapping.objects.update_or_create(
                entity=line.financial_year.entity,
                client_account_code=line.account_code,
                defaults={
                    "client_account_name": line.account_name,
                    "mapped_line_item": mapping,
                },
            )
        except AccountMapping.DoesNotExist:
            pass

    mappings = AccountMapping.objects.all()
    return render(request, "partials/tb_line_row.html", {
        "line": line, "mappings": mappings
    })


# ---------------------------------------------------------------------------
# Entity Officers / Signatories
# ---------------------------------------------------------------------------
@login_required
def entity_officers(request, pk):
    """List all officers/signatories for an entity."""
    entity = get_entity_for_user(request, pk)
    officers = entity.officers.all()

    officer_label_map = {
        "company": "Director / Officer",
        "trust": "Trustee / Beneficiary",
        "partnership": "Partner",
        "sole_trader": "Proprietor",
        "smsf": "Trustee / Director",
    }
    officer_label = officer_label_map.get(entity.entity_type, "Officer")

    return render(request, "core/entity_officers.html", {
        "entity": entity,
        "officers": officers,
        "officer_label": officer_label,
    })


@login_required
def entity_officer_create(request, entity_pk):
    """Add a new officer/signatory to an entity."""
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_officers", pk=entity.pk)

    if request.method == "POST":
        form = EntityOfficerForm(request.POST, entity_type=entity.entity_type)
        if form.is_valid():
            officer = form.save(commit=False)
            officer.entity = entity
            officer.save()
            _log_action(request, "user_change",
                        f"Added officer {officer.full_name} to {entity.entity_name}",
                        officer)
            messages.success(request, f"Added {officer.full_name} as {officer.get_role_display()}.")
            return redirect("core:entity_officers", pk=entity.pk)
    else:
        form = EntityOfficerForm(entity_type=entity.entity_type)

    return render(request, "core/entity_officer_form.html", {
        "form": form,
        "entity": entity,
    })


@login_required
def entity_officer_edit(request, pk):
    """Edit an existing officer/signatory."""
    officer = get_object_or_404(EntityOfficer, pk=pk)
    entity = officer.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_officers", pk=entity.pk)

    if request.method == "POST":
        form = EntityOfficerForm(request.POST, instance=officer, entity_type=entity.entity_type)
        if form.is_valid():
            form.save()
            _log_action(request, "user_change",
                        f"Updated officer {officer.full_name} for {entity.entity_name}",
                        officer)
            messages.success(request, f"Updated {officer.full_name}.")
            return redirect("core:entity_officers", pk=entity.pk)
    else:
        form = EntityOfficerForm(instance=officer, entity_type=entity.entity_type)

    return render(request, "core/entity_officer_form.html", {
        "form": form,
        "entity": entity,
    })


@login_required
@require_POST
def entity_officer_delete(request, pk):
    """Delete an officer/signatory."""
    officer = get_object_or_404(EntityOfficer, pk=pk)
    entity = officer.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_officers", pk=entity.pk)
    name = officer.full_name
    _log_action(request, "user_change",
                f"Removed officer {name} from {entity.entity_name}",
                officer)
    officer.delete()
    messages.success(request, f"Removed {name}.")
    return redirect("core:entity_officers", pk=entity.pk)


# ---------------------------------------------------------------------------
# Access Ledger Import
# ---------------------------------------------------------------------------
@login_required
def access_ledger_import(request):
    """Import an Access Ledger ZIP export."""
    from .access_ledger_import import import_access_ledger_zip

    result = None

    if request.method == "POST":
        zip_file = request.FILES.get("zip_file")
        if not zip_file:
            messages.error(request, "Please select a ZIP file.")
        elif not zip_file.name.lower().endswith(".zip"):
            messages.error(request, "File must be a .zip file.")
        else:
            replace = request.POST.get("replace_existing") == "1"
            try:
                result = import_access_ledger_zip(
                    zip_file,
                    replace_existing=replace,
                )
                if result["errors"]:
                    messages.warning(
                        request,
                        f"Import completed with {len(result['errors'])} error(s)."
                    )
                else:
                    messages.success(
                        request,
                        f"Successfully imported {result['entity'].entity_name}: "
                        f"{result['years_imported']} years, "
                        f"{result['total_tb_lines']} TB lines, "
                        f"{result['total_dep_assets']} depreciation assets."
                    )
                _log_action(
                    request, "import",
                    f"Imported Access Ledger ZIP: {zip_file.name} "
                    f"({result['years_imported']} years)",
                    result.get("entity"),
                )
            except Exception as e:
                messages.error(request, "Import failed. Please check the file format and try again.")

    return render(request, "core/access_ledger_import.html", {
        "result": result,
    })


# ============================================================
# PDF Downloads: Trial Balance & Journals List
# ============================================================

@login_required
def trial_balance_pdf(request, pk):
    """Generate a comparative trial balance PDF matching professional accounting format."""
    from io import BytesIO
    from collections import OrderedDict
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Table, TableStyle,
        Paragraph, Spacer, Image, NextPageTemplate, PageBreak,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    import os
    from django.conf import settings

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    tb_lines = TrialBalanceLine.objects.filter(financial_year=fy).select_related('mapped_line_item').order_by('account_code')

    # Determine section ordering and grouping
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }

    # Group lines by section
    sections = OrderedDict()
    unmapped_lines = []
    for line in tb_lines:
        if line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = 'Unmapped'
        if display_section not in sections:
            sections[display_section] = []
        sections[display_section].append(line)

    # Sort sections by defined order
    ordered_sections = OrderedDict()
    section_keys_ordered = []
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in section_keys_ordered:
            section_keys_ordered.append(ds)
    for key in section_keys_ordered:
        if key in sections:
            ordered_sections[key] = sections[key]
    # Add any remaining sections
    for key in sections:
        if key not in ordered_sections:
            ordered_sections[key] = sections[key]

    # Year labels
    current_year = str(fy.year_label)
    # Extract year number from label like "FY2026" or "2026"
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year = str(fy.prior_year.year_label)
    else:
        prior_year = 'Prior'

    # ABN
    abn = entity.abn if hasattr(entity, 'abn') and entity.abn else ''
    abn_display = ''
    if abn:
        abn_clean = abn.replace(' ', '')
        if len(abn_clean) == 11:
            abn_display = f'ABN {abn_clean[:2]} {abn_clean[2:5]} {abn_clean[5:8]} {abn_clean[8:11]}'
        else:
            abn_display = f'ABN {abn}'

    buffer = BytesIO()

    # Custom page template with header and footer on every page
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'mcs_logo.png')

    class TBDocTemplate(BaseDocTemplate):
        def __init__(self, *args, **kwargs):
            self.entity_name = kwargs.pop('entity_name', '')
            self.abn_display = kwargs.pop('abn_display', '')
            self.end_date_str = kwargs.pop('end_date_str', '')
            self.current_year = kwargs.pop('current_year', '')
            self.prior_year = kwargs.pop('prior_year', '')
            super().__init__(*args, **kwargs)

        def afterPage(self):
            """Draw footer on every page."""
            canvas = self.canv
            canvas.saveState()
            # Footer line
            canvas.setStrokeColor(colors.HexColor('#333333'))
            canvas.setLineWidth(0.5)
            canvas.line(15*mm, 18*mm, A4[0] - 15*mm, 18*mm)
            # Footer text
            canvas.setFont('Helvetica-Bold', 7)
            canvas.setFillColor(colors.HexColor('#333333'))
            footer_text = (
                "These financial statements are unaudited. They must be read in conjunction with the attached "
                "Accountant's Compilation Report and Notes which form part of these financial statements."
            )
            # Center the text
            text_width = canvas.stringWidth(footer_text, 'Helvetica-Bold', 7)
            page_width = A4[0]
            if text_width > page_width - 30*mm:
                # Two-line footer
                line1 = "These financial statements are unaudited. They must be read in conjunction with the attached Accountant's"
                line2 = "Compilation Report and Notes which form part of these financial statements."
                w1 = canvas.stringWidth(line1, 'Helvetica-Bold', 7)
                w2 = canvas.stringWidth(line2, 'Helvetica-Bold', 7)
                canvas.drawString((page_width - w1) / 2, 13*mm, line1)
                canvas.drawString((page_width - w2) / 2, 10*mm, line2)
            else:
                canvas.drawString((page_width - text_width) / 2, 12*mm, footer_text)
            canvas.restoreState()

    doc = TBDocTemplate(
        buffer, pagesize=A4, topMargin=15*mm, bottomMargin=25*mm,
        leftMargin=15*mm, rightMargin=15*mm,
        entity_name=entity.entity_name.upper(),
        abn_display=abn_display,
        end_date_str=fy.end_date.strftime('%d %B %Y'),
        current_year=current_year,
        prior_year=prior_year,
    )

    frame = Frame(
        doc.leftMargin, doc.bottomMargin,
        doc.width, doc.height,
        id='main'
    )
    doc.addPageTemplates([PageTemplate(id='main', frames=[frame])])

    styles = getSampleStyleSheet()
    elements = []

    # Styles
    s_entity = ParagraphStyle('Entity', fontName='Helvetica-Bold', fontSize=14,
                               alignment=TA_CENTER, spaceAfter=2*mm)
    s_abn = ParagraphStyle('ABN', fontName='Helvetica', fontSize=10,
                            alignment=TA_CENTER, spaceAfter=2*mm)
    s_title = ParagraphStyle('TBTitle', fontName='Helvetica-Bold', fontSize=11,
                              alignment=TA_CENTER, spaceAfter=6*mm)
    s_section = ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=11,
                                spaceBefore=5*mm, spaceAfter=2*mm)
    s_cell = ParagraphStyle('Cell', fontName='Helvetica', fontSize=9)
    s_cell_bold = ParagraphStyle('CellBold', fontName='Helvetica-Bold', fontSize=9)
    s_num = ParagraphStyle('Num', fontName='Helvetica', fontSize=9, alignment=TA_RIGHT)
    s_num_bold = ParagraphStyle('NumBold', fontName='Helvetica-Bold', fontSize=9, alignment=TA_RIGHT)

    # Logo
    if os.path.exists(logo_path):
        logo = Image(logo_path, width=25*mm, height=25*mm)
        logo.hAlign = 'CENTER'
        elements.append(logo)
        elements.append(Spacer(1, 2*mm))

    # Header
    elements.append(Paragraph(entity.entity_name.upper(), s_entity))
    if abn_display:
        elements.append(Paragraph(abn_display, s_abn))
    elements.append(Paragraph(
        f"Comparative Trial Balance as at {fy.end_date.strftime('%d %B %Y')}", s_title
    ))

    # Column header table
    col_widths = [50, 165, 75, 75, 75, 75]
    header_data = [
        ['', '', current_year, current_year, prior_year, prior_year],
        ['', '', '$ Dr', '$ Cr', '$ Dr', '$ Cr'],
    ]
    header_table = Table(header_data, colWidths=col_widths)
    header_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
        ('LINEBELOW', (2, 1), (-1, 1), 0.8, colors.HexColor('#333333')),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    elements.append(header_table)

    def fmt(val):
        """Format a decimal value with commas, or return empty string if zero."""
        if val and val != Decimal('0'):
            return f"{val:,.2f}"
        return ''

    # Totals accumulators
    grand_total_dr = Decimal('0')
    grand_total_cr = Decimal('0')
    grand_total_prior_dr = Decimal('0')
    grand_total_prior_cr = Decimal('0')

    # Build section tables
    for section_name, lines in ordered_sections.items():
        elements.append(Paragraph(f"<b>{section_name}</b>", s_section))

        data = []
        for line in lines:
            dr = line.debit if line.debit else Decimal('0')
            cr = line.credit if line.credit else Decimal('0')
            prior_dr = line.prior_debit if line.prior_debit else Decimal('0')
            prior_cr = line.prior_credit if line.prior_credit else Decimal('0')

            grand_total_dr += dr
            grand_total_cr += cr
            grand_total_prior_dr += prior_dr
            grand_total_prior_cr += prior_cr

            row = [
                Paragraph(f"<b>{line.account_code}</b>", ParagraphStyle('Code', fontName='Helvetica-Bold', fontSize=9)),
                Paragraph(line.account_name, s_cell),
                fmt(dr),
                fmt(cr),
                fmt(prior_dr),
                fmt(prior_cr),
            ]
            data.append(row)

        if data:
            section_table = Table(data, colWidths=col_widths)
            style_cmds = [
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('LINEBELOW', (0, 0), (-1, -1), 0.3, colors.HexColor('#cccccc')),
            ]
            section_table.setStyle(TableStyle(style_cmds))
            elements.append(section_table)

    # Grand totals
    elements.append(Spacer(1, 3*mm))
    totals_data = [[
        '', '',
        f"{grand_total_dr:,.2f}",
        f"{grand_total_cr:,.2f}",
        f"{grand_total_prior_dr:,.2f}",
        f"{grand_total_prior_cr:,.2f}",
    ]]
    totals_table = Table(totals_data, colWidths=col_widths)
    totals_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
        ('LINEABOVE', (2, 0), (-1, 0), 1.0, colors.HexColor('#333333')),
        ('LINEBELOW', (2, 0), (-1, 0), 0.5, colors.HexColor('#333333')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(totals_table)

    # Net Profit
    net_profit_current = grand_total_cr - grand_total_dr
    net_profit_prior = grand_total_prior_cr - grand_total_prior_dr
    elements.append(Spacer(1, 3*mm))
    profit_data = [[
        '', Paragraph('<b>Net Profit</b>', s_cell_bold),
        '', f"{abs(net_profit_current):,.2f}",
        '', f"{abs(net_profit_prior):,.2f}",
    ]]
    profit_table = Table(profit_data, colWidths=col_widths)
    profit_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
        ('LINEBELOW', (2, 0), (-1, 0), 0.8, colors.HexColor('#333333')),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    elements.append(profit_table)

    doc.build(elements)
    buffer.seek(0)

    filename = f"Comparative_TB_{entity.entity_name.replace(' ', '_')}_{fy.year_label}.pdf"
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def journals_pdf(request, pk):
    """Generate a PDF listing all journal entries for a financial year."""
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, KeepTogether, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    import os

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    journals = AdjustingJournal.objects.filter(financial_year=fy).prefetch_related('lines').order_by('created_at')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm,
                            leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    elements = []

    # Logo
    from django.conf import settings
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'mcs_logo.png')
    if os.path.exists(logo_path):
        logo = Image(logo_path, width=40*mm, height=40*mm)
        logo.hAlign = 'LEFT'
        elements.append(logo)
        elements.append(Spacer(1, 3*mm))

    # Title
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=16, spaceAfter=4*mm)
    elements.append(Paragraph(f"{entity.entity_name}", title_style))

    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, textColor=colors.grey, spaceAfter=2*mm)
    elements.append(Paragraph(f"Journal Entries — {fy.year_label}", subtitle_style))
    elements.append(Paragraph(
        f"{fy.start_date.strftime('%d %b %Y')} to {fy.end_date.strftime('%d %b %Y')} · Status: {fy.get_status_display()}",
        ParagraphStyle('DateLine', parent=styles['Normal'], fontSize=9, textColor=colors.grey, spaceAfter=6*mm)
    ))

    if not journals.exists():
        elements.append(Paragraph("No journal entries recorded for this financial year.", styles['Normal']))
    else:
        # Summary table
        summary_data = [
            ['Total Journals', str(journals.count())],
            ['Posted', str(journals.filter(status='posted').count())],
            ['Draft', str(journals.filter(status='draft').count())],
        ]
        summary_table = Table(summary_data, colWidths=[100, 80])
        summary_table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 6*mm))

        # Each journal
        for journal in journals:
            journal_elements = []

            # Journal header
            ref = journal.reference_number or 'DRAFT'
            status_text = journal.get_status_display()
            jtype = journal.get_journal_type_display()
            header_text = (
                f"<b>{ref}</b> · {jtype} · {status_text} · "
                f"{journal.journal_date.strftime('%d %b %Y')}"
            )
            journal_elements.append(Paragraph(header_text, ParagraphStyle(
                'JournalHeader', fontSize=10, spaceAfter=1*mm,
                textColor=colors.HexColor('#212529')
            )))

            if journal.description:
                journal_elements.append(Paragraph(
                    journal.description,
                    ParagraphStyle('JDesc', fontSize=8, textColor=colors.grey, spaceAfter=2*mm)
                ))

            # Lines table
            lines = journal.lines.all()
            line_header = ['Account', 'Name', 'Description', 'Debit', 'Credit']
            line_data = [line_header]
            total_dr = Decimal('0')
            total_cr = Decimal('0')

            for line in lines:
                line_data.append([
                    line.account_code,
                    Paragraph(line.account_name, ParagraphStyle('LCell', fontSize=7)),
                    Paragraph(line.description or '', ParagraphStyle('LCell2', fontSize=7, textColor=colors.grey)),
                    f"{line.debit:,.2f}" if line.debit else '',
                    f"{line.credit:,.2f}" if line.credit else '',
                ])
                total_dr += line.debit or Decimal('0')
                total_cr += line.credit or Decimal('0')

            line_data.append(['', '', Paragraph('<b>Total</b>', ParagraphStyle('Bold', fontSize=7)),
                              f"{total_dr:,.2f}", f"{total_cr:,.2f}"])

            line_table = Table(line_data, colWidths=[50, 120, 130, 65, 65], repeatRows=1)
            line_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#495057')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ALIGN', (3, 0), (4, -1), 'RIGHT'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#f8f9fa')]),
                ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#e9ecef')),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            journal_elements.append(line_table)

            # Created by info
            created_by = journal.created_by.get_full_name() if journal.created_by else 'System'
            journal_elements.append(Paragraph(
                f"Created by {created_by} on {journal.created_at.strftime('%d/%m/%Y %H:%M')}",
                ParagraphStyle('CreatedBy', fontSize=6, textColor=colors.grey, spaceBefore=1*mm)
            ))
            journal_elements.append(Spacer(1, 5*mm))

            elements.append(KeepTogether(journal_elements))

    # Footer
    elements.append(Spacer(1, 4*mm))
    footer_style = ParagraphStyle('Footer', fontSize=7, textColor=colors.grey)
    elements.append(Paragraph(
        f"Generated by StatementHub on {timezone.now().strftime('%d %b %Y at %H:%M')} · MC & S Pty Ltd",
        footer_style
    ))

    doc.build(elements)
    buffer.seek(0)

    filename = f"Journals_{entity.entity_name.replace(' ', '_')}_{fy.year_label}.pdf"
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Client Associates (Family & Business)
# ---------------------------------------------------------------------------
@login_required
def associate_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to add associates.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = ClientAssociateForm(request.POST)
        if form.is_valid():
            assoc = form.save(commit=False)
            assoc.entity = entity
            assoc.save()
            _log_action(request, "import", f"Added associate: {assoc.name} ({assoc.get_relationship_type_display()})", assoc)
            messages.success(request, f"Associate '{assoc.name}' added.")
            return redirect("core:entity_detail", pk=entity_pk)
    else:
        form = ClientAssociateForm()
    return render(request, "core/associate_form.html", {
        "form": form, "entity": entity, "title": "Add Associate"
    })


@login_required
def associate_edit(request, pk):
    assoc = get_object_or_404(ClientAssociate, pk=pk)
    entity = assoc.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to edit associates.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        form = ClientAssociateForm(request.POST, instance=assoc)
        if form.is_valid():
            form.save()
            messages.success(request, f"Associate '{assoc.name}' updated.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = ClientAssociateForm(instance=assoc)
    return render(request, "core/associate_form.html", {
        "form": form, "entity": entity, "title": f"Edit: {assoc.name}"
    })


@login_required
def associate_delete(request, pk):
    assoc = get_object_or_404(ClientAssociate, pk=pk)
    entity = assoc.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to delete associates.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        name = assoc.name
        assoc.delete()
        messages.success(request, f"Associate '{name}' removed.")
    return redirect("core:entity_detail", pk=entity.pk)


# ---------------------------------------------------------------------------
# Accounting Software
# ---------------------------------------------------------------------------
@login_required
def software_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to add software configs.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = AccountingSoftwareForm(request.POST)
        if form.is_valid():
            sw = form.save(commit=False)
            sw.entity = entity
            sw.save()
            _log_action(request, "import", f"Added software: {sw.get_software_type_display()}", sw)
            messages.success(request, f"Software '{sw.get_software_type_display()}' added.")
            return redirect("core:entity_detail", pk=entity_pk)
    else:
        form = AccountingSoftwareForm()
    return render(request, "core/software_form.html", {
        "form": form, "entity": entity, "title": "Add Accounting Software", "is_new": True
    })


@login_required
def software_edit(request, pk):
    sw = get_object_or_404(AccountingSoftware, pk=pk)
    entity = sw.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to edit software configs.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        form = AccountingSoftwareForm(request.POST, instance=sw)
        if form.is_valid():
            form.save()
            messages.success(request, f"Software '{sw.get_software_type_display()}' updated.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = AccountingSoftwareForm(instance=sw)
    return render(request, "core/software_form.html", {
        "form": form, "entity": entity, "title": f"Edit: {sw.get_software_type_display()}", "is_new": False
    })


@login_required
def software_delete(request, pk):
    sw = get_object_or_404(AccountingSoftware, pk=pk)
    entity = sw.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to delete software configs.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        label = sw.get_software_type_display()
        sw.delete()
        messages.success(request, f"Software '{label}' removed.")
    return redirect("core:entity_detail", pk=entity.pk)


# ---------------------------------------------------------------------------
# Meeting Notes
# ---------------------------------------------------------------------------
@login_required
def meeting_note_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = MeetingNoteForm(request.POST)
        if form.is_valid():
            note = form.save(commit=False)
            note.entity = entity
            note.created_by = request.user
            note.save()
            _log_action(request, "import", f"Added meeting note: {note.title}", note)
            messages.success(request, f"Meeting note '{note.title}' created.")
            return redirect("core:entity_detail", pk=entity_pk)
    else:
        form = MeetingNoteForm(initial={"meeting_date": timezone.now().date()})
    return render(request, "core/meeting_note_form.html", {
        "form": form, "entity": entity, "title": "New Meeting Note"
    })


@login_required
def meeting_note_edit(request, pk):
    note = get_object_or_404(MeetingNote, pk=pk)
    entity = note.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        form = MeetingNoteForm(request.POST, instance=note)
        if form.is_valid():
            form.save()
            messages.success(request, f"Meeting note '{note.title}' updated.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = MeetingNoteForm(instance=note)
    return render(request, "core/meeting_note_form.html", {
        "form": form, "entity": entity, "title": f"Edit: {note.title}"
    })


@login_required
def meeting_note_detail(request, pk):
    note = get_object_or_404(MeetingNote, pk=pk)
    entity = note.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    return render(request, "core/meeting_note_detail.html", {
        "note": note, "entity": entity,
    })


@login_required
def meeting_note_delete(request, pk):
    note = get_object_or_404(MeetingNote, pk=pk)
    entity = note.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        title = note.title
        note.delete()
        messages.success(request, f"Meeting note '{title}' deleted.")
    return redirect("core:entity_detail", pk=entity.pk)


@login_required
def meeting_note_toggle_followup(request, pk):
    """HTMX endpoint to toggle follow-up completion."""
    note = get_object_or_404(MeetingNote, pk=pk)
    get_entity_for_user(request, note.entity.pk)  # IDOR check
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)
    note.follow_up_completed = not note.follow_up_completed
    note.save(update_fields=["follow_up_completed"])
    return JsonResponse({"completed": note.follow_up_completed})



# ---------------------------------------------------------------------------
# GST Activity Statement
# ---------------------------------------------------------------------------
@login_required
def gst_activity_statement(request, pk):
    """
    Generate a GST Activity Statement (BAS summary) for a financial year.
    Maps trial balance lines to BAS labels using ChartOfAccount tax codes.
    Supports both Simpler BAS (G1, 1A, 1B) and Full BAS (G1-G20).
    """
    fy = get_financial_year_for_user(request, pk)
    # Re-fetch with select_related for efficiency
    fy = get_object_or_404(
        FinancialYear.objects.select_related("entity", "entity__client"),
        pk=pk,
    )
    entity = fy.entity
    entity_type = entity.entity_type  # company, trust, partnership, sole_trader

    # Build a lookup: account_code -> ChartOfAccount for this entity type
    coa_lookup = {}
    for coa in ChartOfAccount.objects.filter(entity_type=entity_type, is_active=True):
        coa_lookup[coa.account_code] = coa

    # Get all trial balance lines for this financial year
    tb_lines = fy.trial_balance_lines.all()

    # Initialize BAS labels
    g1 = Decimal("0")   # Total sales (including GST)
    g2 = Decimal("0")   # Export sales
    g3 = Decimal("0")   # Other GST-free sales
    g4 = Decimal("0")   # Input taxed sales
    g10 = Decimal("0")  # Capital purchases (including GST)
    g11 = Decimal("0")  # Non-capital purchases (including GST)
    g13 = Decimal("0")  # Purchases for making input taxed sales
    g14 = Decimal("0")  # Purchases without GST (GST-free purchases)
    g15 = Decimal("0")  # Private use / not income tax deductible
    g7 = Decimal("0")   # Sales adjustments
    g18 = Decimal("0")  # Purchase adjustments

    # Detailed line items for the breakdown
    sales_lines = []
    purchase_lines = []
    capital_lines = []
    excluded_lines = []

    for line in tb_lines:
        coa = coa_lookup.get(line.account_code)
        if not coa:
            # Check if the line has its own tax_type from bank statement review
            line_tax = (getattr(line, 'tax_type', '') or '').strip()
            if line_tax:
                # GST accounts from bank statement review (9100 GST Collected, 9110 GST Paid)
                if line.account_code in ('9100', '9110'):
                    # These are GST clearing accounts - they contribute to 1A/1B directly
                    pass  # Handled via the G9/G20 calculation from revenue/expense lines
                excluded_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "amount": abs(line.closing_balance),
                    "reason": f"Not in chart of accounts (tax: {line_tax})",
                })
            else:
                excluded_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "amount": abs(line.closing_balance),
                    "reason": "Not in chart of accounts",
                })
            continue

        section = coa.section
        # Use line-level tax_type if ChartOfAccount has no tax code
        coa_tax = (coa.tax_code or "").upper().strip()
        line_tax = (getattr(line, 'tax_type', '') or '').strip()
        # Map line-level tax types to COA tax codes
        if not coa_tax and line_tax:
            tax_map = {
                'GST on Income': 'GST',
                'GST on Expenses': 'GST',
                'GST Free Income': 'FRE',
                'GST Free Expenses': 'FRE',
                'BAS Excluded': 'N-T',
                'N-T': 'N-T',
            }
            tax_code = tax_map.get(line_tax, coa_tax)
        else:
            tax_code = coa_tax
        # Use closing balance; revenue is typically credit (negative in TB)
        amount = abs(line.closing_balance)

        if section == "Revenue":
            # All revenue goes to G1
            g1 += amount
            bas_label = "G1"

            if tax_code == "GST":
                bas_label = "G1 (Taxable)"
            elif tax_code == "ITS":
                g4 += amount
                bas_label = "G4 (Input Taxed)"
            elif tax_code == "ADS":
                g7 += amount
                bas_label = "G7 (Adjustment)"
            elif tax_code in ("", "FRE", "N-T"):
                g3 += amount
                bas_label = "G3 (GST-Free)"

            sales_lines.append({
                "code": line.account_code,
                "name": line.account_name,
                "tax_code": tax_code or "N-T",
                "amount": amount,
                "bas_label": bas_label,
            })

        elif section == "Expenses":
            # Non-capital purchases go to G11
            if tax_code in ("INP", "GST"):
                g11 += amount
                bas_label = "G11 (Non-Capital)"
            elif tax_code in ("IOA",):
                g11 += amount
                g13 += amount
                bas_label = "G11/G13 (Input Taxed)"
            elif tax_code in ("FOA", "FRE"):
                g11 += amount
                g14 += amount
                bas_label = "G11/G14 (GST-Free)"
            elif tax_code == "ADS":
                g11 += amount
                g18 += amount
                bas_label = "G11/G18 (Adjustment)"
            else:
                # No tax code - still include in G11 but also G14
                g11 += amount
                g14 += amount
                bas_label = "G11/G14 (No GST)"

            purchase_lines.append({
                "code": line.account_code,
                "name": line.account_name,
                "tax_code": tax_code or "N-T",
                "amount": amount,
                "bas_label": bas_label,
            })

        elif section == "Assets":
            if tax_code == "CAP":
                g10 += amount
                bas_label = "G10 (Capital)"
                capital_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "tax_code": tax_code,
                    "amount": amount,
                    "bas_label": bas_label,
                })
            elif tax_code == "FCA":
                g10 += amount
                g14 += amount
                bas_label = "G10/G14 (GST-Free Capital)"
                capital_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "tax_code": tax_code,
                    "amount": amount,
                    "bas_label": bas_label,
                })
            # Other assets (no tax code) are balance sheet items, not on BAS

        # Liabilities, Equity, Capital Accounts - not on BAS

    # Calculated fields
    g5 = g2 + g3 + g4                    # Total non-taxable sales
    g6 = g1 - g5                          # Sales subject to GST
    g8 = g6 + g7                          # Sales subject to GST after adjustments
    g9 = (g8 / Decimal("11")).quantize(Decimal("0.01")) if g8 else Decimal("0")  # GST on sales

    g12 = g10 + g11                       # Total purchases
    g16 = g13 + g14 + g15                 # Non-creditable purchases
    g17 = g12 - g16                       # Purchases subject to GST
    g19 = g17 + g18                       # Purchases subject to GST after adjustments
    g20 = (g19 / Decimal("11")).quantize(Decimal("0.01")) if g19 else Decimal("0")  # GST on purchases

    label_1a = g9                         # GST on sales (payable)
    label_1b = g20                        # GST on purchases (credit)
    gst_payable = label_1a - label_1b     # Net GST payable (or refund if negative)

    # Build the context
    bas_data = {
        # Sales section
        "G1": g1, "G2": g2, "G3": g3, "G4": g4,
        "G5": g5, "G6": g6, "G7": g7, "G8": g8, "G9": g9,
        # Purchases section
        "G10": g10, "G11": g11, "G12": g12, "G13": g13,
        "G14": g14, "G15": g15, "G16": g16, "G17": g17,
        "G18": g18, "G19": g19, "G20": g20,
        # Summary
        "1A": label_1a, "1B": label_1b,
        "gst_payable": gst_payable,
    }

    context = {
        "fy": fy,
        "entity": entity,
        "bas_data": bas_data,
        "sales_lines": sales_lines,
        "purchase_lines": purchase_lines,
        "capital_lines": capital_lines,
        "excluded_lines": excluded_lines,
        "is_gst_registered": entity.is_gst_registered,
    }
    return render(request, "core/gst_activity_statement.html", context)


@login_required
def gst_activity_statement_download(request, pk):
    """Download GST Activity Statement as Excel or PDF."""
    import io

    fmt = request.GET.get("format", "excel").lower()

    fy = get_object_or_404(
        FinancialYear.objects.select_related("entity", "entity__client"),
        pk=pk,
    )
    entity = fy.entity
    entity_type = entity.entity_type

    # Re-run the calculation (same logic as the view)
    coa_lookup = {}
    for coa in ChartOfAccount.objects.filter(entity_type=entity_type, is_active=True):
        coa_lookup[coa.account_code] = coa

    tb_lines = fy.trial_balance_lines.all()

    g1 = g2 = g3 = g4 = g7 = Decimal("0")
    g10 = g11 = g13 = g14 = g15 = g18 = Decimal("0")
    detail_rows = []

    for line in tb_lines:
        coa = coa_lookup.get(line.account_code)
        if not coa:
            continue
        section = coa.section
        tax_code = (coa.tax_code or "").upper().strip()
        amount = abs(line.closing_balance)

        bas_label = ""
        if section == "Revenue":
            g1 += amount
            if tax_code == "GST":
                bas_label = "G1"
            elif tax_code == "ITS":
                g4 += amount
                bas_label = "G4"
            elif tax_code == "ADS":
                g7 += amount
                bas_label = "G7"
            else:
                g3 += amount
                bas_label = "G3"
        elif section == "Expenses":
            if tax_code in ("INP", "GST"):
                g11 += amount
                bas_label = "G11"
            elif tax_code == "IOA":
                g11 += amount
                g13 += amount
                bas_label = "G11/G13"
            elif tax_code in ("FOA", "FRE"):
                g11 += amount
                g14 += amount
                bas_label = "G11/G14"
            elif tax_code == "ADS":
                g11 += amount
                g18 += amount
                bas_label = "G11/G18"
            else:
                g11 += amount
                g14 += amount
                bas_label = "G11/G14"
        elif section == "Assets":
            if tax_code == "CAP":
                g10 += amount
                bas_label = "G10"
            elif tax_code == "FCA":
                g10 += amount
                g14 += amount
                bas_label = "G10/G14"

        if bas_label:
            detail_rows.append({
                "code": line.account_code,
                "name": line.account_name,
                "tax_code": tax_code or "N-T",
                "amount": float(amount),
                "bas_label": bas_label,
            })

    g5 = g2 + g3 + g4
    g6 = g1 - g5
    g8 = g6 + g7
    g9 = (g8 / Decimal("11")).quantize(Decimal("0.01")) if g8 else Decimal("0")
    g12 = g10 + g11
    g16 = g13 + g14 + g15
    g17 = g12 - g16
    g19 = g17 + g18
    g20 = (g19 / Decimal("11")).quantize(Decimal("0.01")) if g19 else Decimal("0")
    label_1a = g9
    label_1b = g20

    # PDF format
    if fmt == "pdf":
        return _gst_download_pdf(
            fy, entity, detail_rows,
            g1, g2, g3, g4, g5, g6, g7, g8, g9,
            g10, g11, g12, g13, g14, g15, g16, g17, g18, g19, g20,
            label_1a, label_1b,
        )

    # Create Excel workbook
    wb = openpyxl.Workbook()

    # Sheet 1: BAS Summary
    ws = wb.active
    ws.title = "GST Activity Statement"

    # Header
    ws.append([f"GST Activity Statement — {entity.entity_name}"])
    ws.append([f"Period: {fy.start_date.strftime('%d/%m/%Y')} to {fy.end_date.strftime('%d/%m/%Y')}"])
    ws.append([f"ABN: {entity.abn or 'N/A'}"])
    ws.append([])

    # GST on Sales
    ws.append(["GST ON SALES"])
    ws.append(["Label", "Description", "Amount"])
    ws.append(["G1", "Total sales (including any GST)", float(g1)])
    ws.append(["G2", "Export sales", float(g2)])
    ws.append(["G3", "Other GST-free sales", float(g3)])
    ws.append(["G4", "Input taxed sales", float(g4)])
    ws.append(["G5", "G2 + G3 + G4", float(g5)])
    ws.append(["G6", "Total sales subject to GST (G1 - G5)", float(g6)])
    ws.append(["G7", "Adjustments", float(g7)])
    ws.append(["G8", "Total sales subject to GST after adjustments (G6 + G7)", float(g8)])
    ws.append(["G9", "GST on sales (G8 ÷ 11)", float(g9)])
    ws.append([])

    # GST on Purchases
    ws.append(["GST ON PURCHASES"])
    ws.append(["Label", "Description", "Amount"])
    ws.append(["G10", "Capital purchases (including any GST)", float(g10)])
    ws.append(["G11", "Non-capital purchases (including any GST)", float(g11)])
    ws.append(["G12", "G10 + G11", float(g12)])
    ws.append(["G13", "Purchases for making input taxed sales", float(g13)])
    ws.append(["G14", "Purchases without GST in the price", float(g14)])
    ws.append(["G15", "Estimated purchases for private use", float(g15)])
    ws.append(["G16", "G13 + G14 + G15", float(g16)])
    ws.append(["G17", "Total purchases subject to GST (G12 - G16)", float(g17)])
    ws.append(["G18", "Adjustments", float(g18)])
    ws.append(["G19", "Total purchases subject to GST after adjustments (G17 + G18)", float(g19)])
    ws.append(["G20", "GST on purchases (G19 ÷ 11)", float(g20)])
    ws.append([])

    # Summary
    ws.append(["BAS SUMMARY"])
    ws.append(["Label", "Description", "Amount"])
    ws.append(["1A", "GST on sales", float(label_1a)])
    ws.append(["1B", "GST on purchases", float(label_1b)])
    ws.append(["", "Net GST payable / (refundable)", float(label_1a - label_1b)])

    # Format columns
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 55
    ws.column_dimensions["C"].width = 18

    # Bold headers
    from openpyxl.styles import Font, numbers
    bold = Font(bold=True)
    for row_idx in [1, 2, 3, 5, 6, 17, 18, 31, 32]:
        for cell in ws[row_idx]:
            cell.font = bold

    # Number format for amount column
    for row in ws.iter_rows(min_row=7, max_col=3, max_row=ws.max_row):
        cell = row[2]
        if isinstance(cell.value, (int, float)):
            cell.number_format = '#,##0.00'

    # Sheet 2: Detail Breakdown
    ws2 = wb.create_sheet("Detail Breakdown")
    ws2.append(["Account Code", "Account Name", "Tax Code", "Amount", "BAS Label"])
    for row in detail_rows:
        ws2.append([row["code"], row["name"], row["tax_code"], row["amount"], row["bas_label"]])

    ws2.column_dimensions["A"].width = 14
    ws2.column_dimensions["B"].width = 40
    ws2.column_dimensions["C"].width = 12
    ws2.column_dimensions["D"].width = 18
    ws2.column_dimensions["E"].width = 14

    for cell in ws2[1]:
        cell.font = bold
    for row in ws2.iter_rows(min_row=2, min_col=4, max_col=4, max_row=ws2.max_row):
        row[0].number_format = '#,##0.00'

    # Write to response
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    entity_name = entity.entity_name.replace(" ", "_")
    filename = f"GST_Activity_Statement_{entity_name}_{fy.start_date.strftime('%Y%m%d')}_{fy.end_date.strftime('%Y%m%d')}.xlsx"

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response



def _gst_download_pdf(fy, entity, detail_rows,
                      g1, g2, g3, g4, g5, g6, g7, g8, g9,
                      g10, g11, g12, g13, g14, g15, g16, g17, g18, g19, g20,
                      label_1a, label_1b):
    """Generate GST Activity Statement as PDF using WeasyPrint."""
    import io
    import weasyprint
    from django.utils.html import escape as html_escape

    abn = html_escape(entity.abn or 'N/A')
    period = f"{fy.start_date.strftime('%d/%m/%Y')} to {fy.end_date.strftime('%d/%m/%Y')}"
    net_gst = label_1a - label_1b
    net_label = "Net GST Payable to ATO" if net_gst > 0 else "Net GST Refundable from ATO"

    def fmt(v):
        return f"${v:,.2f}"

    # Build detail rows HTML (escaped to prevent HTML injection)
    detail_html = ""
    for row in detail_rows:
        detail_html += f"""<tr>
            <td>{html_escape(str(row['code']))}</td><td>{html_escape(str(row['name']))}</td>
            <td>{html_escape(str(row['tax_code']))}</td><td class="r">${row['amount']:,.2f}</td>
            <td>{html_escape(str(row['bas_label']))}</td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @page {{ size: A4; margin: 15mm; }}
    body {{ font-family: 'Times New Roman', serif; font-size: 10pt; color: #333; }}
    h1 {{ font-size: 14pt; text-align: center; margin-bottom: 2mm; }}
    h2 {{ font-size: 11pt; margin-top: 6mm; margin-bottom: 3mm; border-bottom: 1px solid #333; padding-bottom: 2mm; }}
    .sub {{ text-align: center; font-size: 10pt; color: #666; margin-bottom: 4mm; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 4mm; }}
    th, td {{ padding: 3px 6px; font-size: 9pt; border-bottom: 1px solid #ddd; }}
    th {{ background: #f5f5f5; text-align: left; font-weight: bold; }}
    .r {{ text-align: right; }}
    .bold {{ font-weight: bold; }}
    .highlight {{ background: #e8f4fd; }}
    .summary {{ background: #f8f9fa; }}
    .total-row {{ border-top: 2px solid #333; font-weight: bold; }}
</style></head><body>

<h1>GST Activity Statement</h1>
<p class="sub">{html_escape(entity.entity_name)} &mdash; ABN: {abn}</p>
<p class="sub">Period: {period}</p>

<h2>GST on Sales</h2>
<table>
    <tr><th>Label</th><th>Description</th><th class="r">Amount</th></tr>
    <tr><td class="bold">G1</td><td>Total sales (including any GST)</td><td class="r">{fmt(g1)}</td></tr>
    <tr><td>G2</td><td>Export sales</td><td class="r">{fmt(g2)}</td></tr>
    <tr><td>G3</td><td>Other GST-free sales</td><td class="r">{fmt(g3)}</td></tr>
    <tr><td>G4</td><td>Input taxed sales</td><td class="r">{fmt(g4)}</td></tr>
    <tr class="summary"><td class="bold">G5</td><td>G2 + G3 + G4</td><td class="r bold">{fmt(g5)}</td></tr>
    <tr><td class="bold">G6</td><td>Total sales subject to GST (G1 &minus; G5)</td><td class="r bold">{fmt(g6)}</td></tr>
    <tr><td>G7</td><td>Adjustments</td><td class="r">{fmt(g7)}</td></tr>
    <tr class="summary"><td class="bold">G8</td><td>Total sales subject to GST after adj. (G6 + G7)</td><td class="r bold">{fmt(g8)}</td></tr>
    <tr class="highlight"><td class="bold">G9</td><td>GST on sales (G8 &divide; 11)</td><td class="r bold">{fmt(g9)}</td></tr>
</table>

<h2>GST on Purchases</h2>
<table>
    <tr><th>Label</th><th>Description</th><th class="r">Amount</th></tr>
    <tr><td class="bold">G10</td><td>Capital purchases (including any GST)</td><td class="r">{fmt(g10)}</td></tr>
    <tr><td class="bold">G11</td><td>Non-capital purchases (including any GST)</td><td class="r">{fmt(g11)}</td></tr>
    <tr class="summary"><td class="bold">G12</td><td>G10 + G11</td><td class="r bold">{fmt(g12)}</td></tr>
    <tr><td>G13</td><td>Purchases for making input taxed sales</td><td class="r">{fmt(g13)}</td></tr>
    <tr><td>G14</td><td>Purchases without GST in the price</td><td class="r">{fmt(g14)}</td></tr>
    <tr><td>G15</td><td>Estimated purchases for private use</td><td class="r">{fmt(g15)}</td></tr>
    <tr class="summary"><td class="bold">G16</td><td>G13 + G14 + G15</td><td class="r bold">{fmt(g16)}</td></tr>
    <tr><td class="bold">G17</td><td>Total purchases subject to GST (G12 &minus; G16)</td><td class="r bold">{fmt(g17)}</td></tr>
    <tr><td>G18</td><td>Adjustments</td><td class="r">{fmt(g18)}</td></tr>
    <tr class="summary"><td class="bold">G19</td><td>Total purchases subject to GST after adj. (G17 + G18)</td><td class="r bold">{fmt(g19)}</td></tr>
    <tr class="highlight"><td class="bold">G20</td><td>GST on purchases (G19 &divide; 11)</td><td class="r bold">{fmt(g20)}</td></tr>
</table>

<h2>Activity Statement Summary</h2>
<table style="max-width: 400px;">
    <tr><td class="bold">1A</td><td>GST on sales</td><td class="r bold">{fmt(label_1a)}</td></tr>
    <tr><td class="bold">1B</td><td>GST on purchases (credit)</td><td class="r bold">{fmt(label_1b)}</td></tr>
    <tr class="total-row"><td colspan="2">{net_label}</td><td class="r">{fmt(abs(net_gst))}</td></tr>
</table>

<h2>Detail Breakdown</h2>
<table>
    <tr><th>Code</th><th>Account Name</th><th>Tax Code</th><th class="r">Amount</th><th>BAS Label</th></tr>
    {detail_html}
</table>

</body></html>"""

    pdf_bytes = weasyprint.HTML(string=html).write_pdf()

    entity_name = entity.entity_name.replace(' ', '_')
    filename = f"GST_Activity_Statement_{entity_name}_{fy.start_date.strftime('%Y%m%d')}_{fy.end_date.strftime('%Y%m%d')}.pdf"

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Depreciation Asset AJAX endpoints
# ---------------------------------------------------------------------------
@login_required
def depreciation_add(request, pk):
    """Add a new depreciation asset to a financial year."""
    fy = get_financial_year_for_user(request, pk)
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        asset = DepreciationAsset.objects.create(
            financial_year=fy,
            category=request.POST.get("category", "Other"),
            asset_name=request.POST.get("asset_name", ""),
            purchase_date=request.POST.get("purchase_date") or None,
            total_cost=Decimal(request.POST.get("total_cost", "0") or "0"),
            private_use_pct=Decimal(request.POST.get("private_use_pct", "0") or "0"),
            opening_wdv=Decimal(request.POST.get("opening_wdv", "0") or "0"),
            method=request.POST.get("method", "D"),
            rate=Decimal(request.POST.get("rate", "0") or "0"),
            addition_cost=Decimal(request.POST.get("addition_cost", "0") or "0"),
            addition_date=request.POST.get("addition_date") or None,
            disposal_date=request.POST.get("disposal_date") or None,
            disposal_consideration=Decimal(request.POST.get("disposal_consideration", "0") or "0"),
        )
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=pk)
    # Calculate depreciation
    _calc_depreciation(asset)
    asset.save()

    _log_action(request, "create", f"Added depreciation asset: {asset.asset_name}", asset)
    messages.success(request, f"Asset '{asset.asset_name}' added.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def depreciation_edit(request, pk):
    """Edit a depreciation asset."""
    asset = get_object_or_404(DepreciationAsset, pk=pk)
    fy_pk = asset.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        asset.category = request.POST.get("category", asset.category)
        asset.asset_name = request.POST.get("asset_name", asset.asset_name)
        asset.purchase_date = request.POST.get("purchase_date") or asset.purchase_date
        asset.total_cost = Decimal(request.POST.get("total_cost", "0") or "0")
        asset.private_use_pct = Decimal(request.POST.get("private_use_pct", "0") or "0")
        asset.opening_wdv = Decimal(request.POST.get("opening_wdv", "0") or "0")
        asset.method = request.POST.get("method", asset.method)
        asset.rate = Decimal(request.POST.get("rate", "0") or "0")
        asset.addition_cost = Decimal(request.POST.get("addition_cost", "0") or "0")
        asset.addition_date = request.POST.get("addition_date") or None
        asset.disposal_date = request.POST.get("disposal_date") or None
        asset.disposal_consideration = Decimal(request.POST.get("disposal_consideration", "0") or "0")
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    _calc_depreciation(asset)
    asset.save()

    _log_action(request, "update", f"Updated depreciation asset: {asset.asset_name}", asset)
    messages.success(request, f"Asset '{asset.asset_name}' updated.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
@require_POST
def depreciation_delete(request, pk):
    """Delete a depreciation asset."""
    asset = get_object_or_404(DepreciationAsset, pk=pk)
    fy_pk = asset.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    name = asset.asset_name
    asset.delete()
    _log_action(request, "delete", f"Deleted depreciation asset: {name}")
    messages.success(request, f"Asset '{name}' deleted.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
@require_POST
def depreciation_roll_forward(request, pk):
    """Roll forward depreciation schedule from prior year."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    if not fy.prior_year:
        messages.error(request, "No prior year linked. Cannot roll forward depreciation.")
        return redirect("core:financial_year_detail", pk=pk)

    prior_assets = DepreciationAsset.objects.filter(financial_year=fy.prior_year)
    if not prior_assets.exists():
        messages.warning(request, "No depreciation assets in the prior year to roll forward.")
        return redirect("core:financial_year_detail", pk=pk)

    # Clear existing assets if any
    DepreciationAsset.objects.filter(financial_year=fy).delete()

    count = 0
    for pa in prior_assets:
        if pa.closing_wdv <= 0 and not pa.disposal_date:
            continue  # Skip fully depreciated with no disposal
        DepreciationAsset.objects.create(
            financial_year=fy,
            category=pa.category,
            asset_name=pa.asset_name,
            purchase_date=pa.purchase_date,
            total_cost=pa.total_cost,
            private_use_pct=pa.private_use_pct,
            opening_wdv=pa.closing_wdv,  # Prior closing = new opening
            method=pa.method,
            rate=pa.rate,
            display_order=pa.display_order,
        )
        count += 1

    _log_action(request, "roll_forward", f"Rolled forward {count} depreciation assets to {fy}", fy)
    messages.success(request, f"Rolled forward {count} depreciation assets from prior year.")
    return redirect("core:financial_year_detail", pk=pk)


def _calc_depreciation(asset):
    """Calculate depreciation amount and closing WDV for an asset."""
    depreciable = asset.opening_wdv + asset.addition_cost
    if asset.disposal_date:
        # On disposal, depreciation is up to disposal date
        # Profit/loss = disposal consideration - WDV at disposal
        asset.depreciation_amount = Decimal("0")
        asset.closing_wdv = Decimal("0")
        wdv_at_disposal = depreciable
        asset.profit_on_disposal = max(Decimal("0"), asset.disposal_consideration - wdv_at_disposal)
        asset.loss_on_disposal = max(Decimal("0"), wdv_at_disposal - asset.disposal_consideration)
    elif asset.method == "W":
        # Written off entirely
        asset.depreciation_amount = depreciable
        asset.closing_wdv = Decimal("0")
    elif asset.method == "D":
        # Diminishing value
        dep = (depreciable * asset.rate / Decimal("100")).quantize(Decimal("0.01"))
        asset.depreciation_amount = dep
        asset.closing_wdv = (depreciable - dep).quantize(Decimal("0.01"))
    elif asset.method == "P":
        # Prime cost (straight line)
        dep = (asset.total_cost * asset.rate / Decimal("100")).quantize(Decimal("0.01"))
        asset.depreciation_amount = min(dep, depreciable)
        asset.closing_wdv = (depreciable - asset.depreciation_amount).quantize(Decimal("0.01"))
    else:
        asset.depreciation_amount = Decimal("0")
        asset.closing_wdv = depreciable

    # Private use adjustment
    if asset.private_use_pct > 0:
        asset.private_depreciation = (
            asset.depreciation_amount * asset.private_use_pct / Decimal("100")
        ).quantize(Decimal("0.01"))
    else:
        asset.private_depreciation = Decimal("0")

    asset.depreciable_value = depreciable


# ---------------------------------------------------------------------------
# Stock Item AJAX endpoints
# ---------------------------------------------------------------------------
@login_required
def stock_add(request, pk):
    """Add a stock item to a financial year."""
    fy = get_financial_year_for_user(request, pk)
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        StockItem.objects.create(
            financial_year=fy,
            item_name=request.POST.get("item_name", ""),
            opening_quantity=Decimal(request.POST.get("opening_quantity", "0") or "0"),
            opening_value=Decimal(request.POST.get("opening_value", "0") or "0"),
            closing_quantity=Decimal(request.POST.get("closing_quantity", "0") or "0"),
            closing_value=Decimal(request.POST.get("closing_value", "0") or "0"),
            notes=request.POST.get("notes", ""),
        )
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=pk)
    messages.success(request, "Stock item added.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def stock_edit(request, pk):
    """Edit a stock item."""
    item = get_object_or_404(StockItem, pk=pk)
    fy_pk = item.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        item.item_name = request.POST.get("item_name", item.item_name)
        item.opening_quantity = Decimal(request.POST.get("opening_quantity", "0") or "0")
        item.opening_value = Decimal(request.POST.get("opening_value", "0") or "0")
        item.closing_quantity = Decimal(request.POST.get("closing_quantity", "0") or "0")
        item.closing_value = Decimal(request.POST.get("closing_value", "0") or "0")
        item.notes = request.POST.get("notes", "")
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    item.save()
    messages.success(request, f"Stock item '{item.item_name}' updated.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
@require_POST
def stock_delete(request, pk):
    """Delete a stock item."""
    item = get_object_or_404(StockItem, pk=pk)
    fy_pk = item.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    name = item.item_name
    item.delete()
    messages.success(request, f"Stock item '{name}' deleted.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
def stock_push_to_tb(request, pk):
    """Push stock values to the trial balance."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    stock_items = StockItem.objects.filter(financial_year=fy)

    if not stock_items.exists():
        messages.warning(request, "No stock items to push.")
        return redirect("core:financial_year_detail", pk=pk)

    total_opening = sum(s.opening_value for s in stock_items)
    total_closing = sum(s.closing_value for s in stock_items)

    # Remove any existing stock TB lines
    TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_name__in=["Opening Stock", "Closing Stock"],
    ).delete()

    # Create Opening Stock (debit = cost of goods)
    if total_opening > 0:
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code="5100",
            account_name="Opening Stock",
            debit=total_opening,
            credit=Decimal("0"),
        )

    # Create Closing Stock (credit to P&L, debit to Balance Sheet)
    if total_closing > 0:
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code="5200",
            account_name="Closing Stock",
            debit=Decimal("0"),
            credit=total_closing,
        )
        # Balance sheet current asset
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code="1300",
            account_name="Stock on Hand",
            debit=total_closing,
            credit=Decimal("0"),
        )

    # Mark as pushed
    stock_items.update(pushed_to_tb=True)

    messages.success(request, f"Stock pushed to trial balance: Opening ${total_opening}, Closing ${total_closing}.")
    return redirect("core:financial_year_detail", pk=pk)


# ---------------------------------------------------------------------------
# Review → Trial Balance Push
# ---------------------------------------------------------------------------
@login_required
def review_push_to_tb(request, pk):
    """Push confirmed review transactions to the trial balance as journal entries."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    from review.models import PendingTransaction

    confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=True,
    )

    if not confirmed.exists():
        messages.warning(request, "No confirmed transactions to push.")
        return redirect("core:financial_year_detail", pk=pk)

    # Group by confirmed account code and aggregate
    from django.db.models import Sum as DSum
    aggregated = confirmed.values(
        "confirmed_code", "confirmed_name"
    ).annotate(
        total_amount=DSum("amount"),
    )

    count = 0
    for entry in aggregated:
        code = entry["confirmed_code"]
        name = entry["confirmed_name"]
        total = entry["total_amount"]
        if total is None or total == 0:
            continue

        # Check if TB line already exists for this code
        tb_line, created = TrialBalanceLine.objects.get_or_create(
            financial_year=fy,
            account_code=code,
            defaults={
                "account_name": name,
                "debit": max(Decimal("0"), total),
                "credit": abs(min(Decimal("0"), total)),
            },
        )
        if not created:
            # Add to existing line
            if total > 0:
                tb_line.debit += total
            else:
                tb_line.credit += abs(total)
            tb_line.save()
        count += 1

    messages.success(request, f"Pushed {count} account lines to trial balance from {confirmed.count()} transactions.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def review_approve_transaction(request, pk):
    """Approve a single pending transaction (AJAX).
    Also auto-pushes the transaction to the trial balance so TB and GST/BAS
    update live without needing a separate 'Push to TB' step.
    """
    from review.models import PendingTransaction
    txn = get_object_or_404(PendingTransaction, pk=pk)

    # IDOR check: verify user has access to the linked entity
    if txn.job and txn.job.entity:
        get_entity_for_user(request, txn.job.entity.pk)

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)

    txn.confirmed_code = request.POST.get("confirmed_code", txn.ai_suggested_code)
    txn.confirmed_name = request.POST.get("confirmed_name", txn.ai_suggested_name)
    txn.confirmed_tax_type = request.POST.get("confirmed_tax_type", txn.ai_suggested_tax_type)

    # Handle GST toggle from the review form
    has_gst = request.POST.get("has_gst", "0") == "1"
    if has_gst:
        abs_amount = abs(txn.amount)
        gst_amount = (abs_amount / Decimal("11")).quantize(Decimal("0.01"))
        net_amount = abs_amount - gst_amount
        txn.gst_amount = gst_amount
        txn.net_amount = net_amount
        txn.confirmed_gst_amount = gst_amount
        # Ensure tax type is set correctly for GST
        if not txn.confirmed_tax_type or 'Free' in (txn.confirmed_tax_type or '') or txn.confirmed_tax_type in ('BAS Excluded', 'N-T', ''):
            txn.confirmed_tax_type = 'GST on Expenses' if txn.amount < 0 else 'GST on Income'
    else:
        txn.gst_amount = Decimal("0.00")
        txn.net_amount = abs(txn.amount)
        txn.confirmed_gst_amount = Decimal("0.00")
        # Set GST-free tax type if currently GST
        if txn.confirmed_tax_type in ('GST on Income', 'GST on Expenses'):
            txn.confirmed_tax_type = 'GST Free Expenses' if txn.amount < 0 else 'GST Free Income'

    txn.is_confirmed = True
    txn.save()

    # Auto-push to trial balance immediately
    tb_updated = False
    if txn.job and txn.job.entity:
        # Find the financial year for this entity that covers this transaction
        from datetime import datetime as dt
        entity = txn.job.entity
        fys = FinancialYear.objects.filter(entity=entity, status__in=['draft', 'in_review', 'reviewed'])
        target_fy = None
        for fy_candidate in fys:
            # Try to parse the transaction date
            txn_date = None
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d %b %Y"):
                try:
                    txn_date = dt.strptime(txn.date.strip(), fmt).date()
                    break
                except (ValueError, AttributeError):
                    continue
            if txn_date and fy_candidate.start_date <= txn_date <= fy_candidate.end_date:
                target_fy = fy_candidate
                break
        if not target_fy and fys.exists():
            target_fy = fys.order_by('-end_date').first()

        if target_fy and txn.confirmed_code:
            amount = txn.amount
            code = txn.confirmed_code
            name = txn.confirmed_name
            tax_type = txn.confirmed_tax_type or ''

            # Push the net amount (ex-GST) to the expense/income account
            net_for_tb = txn.net_amount if has_gst else abs(amount)
            tb_line, created = TrialBalanceLine.objects.get_or_create(
                financial_year=target_fy,
                account_code=code,
                defaults={
                    "account_name": name,
                    "debit": net_for_tb if amount > 0 else Decimal("0"),
                    "credit": net_for_tb if amount < 0 else Decimal("0"),
                    "tax_type": tax_type,
                },
            )
            if not created:
                if amount > 0:
                    tb_line.debit += net_for_tb
                else:
                    tb_line.credit += net_for_tb
                if not tb_line.tax_type:
                    tb_line.tax_type = tax_type
                tb_line.save()
            tb_updated = True

            # If GST applies, also post the GST component to the GST Collected/Paid account
            if has_gst and txn.confirmed_gst_amount > 0:
                gst_amt = txn.confirmed_gst_amount
                if amount > 0:
                    # Income: GST Collected (liability) - code 9100
                    gst_code = '9100'
                    gst_name = 'GST Collected'
                    gst_line, gst_created = TrialBalanceLine.objects.get_or_create(
                        financial_year=target_fy,
                        account_code=gst_code,
                        defaults={
                            "account_name": gst_name,
                            "debit": Decimal("0"),
                            "credit": gst_amt,
                            "tax_type": "GST on Income",
                        },
                    )
                    if not gst_created:
                        gst_line.credit += gst_amt
                        gst_line.save()
                else:
                    # Expense: GST Paid (asset) - code 9110
                    gst_code = '9110'
                    gst_name = 'GST Paid'
                    gst_line, gst_created = TrialBalanceLine.objects.get_or_create(
                        financial_year=target_fy,
                        account_code=gst_code,
                        defaults={
                            "account_name": gst_name,
                            "debit": gst_amt,
                            "credit": Decimal("0"),
                            "tax_type": "GST on Expenses",
                        },
                    )
                    if not gst_created:
                        gst_line.debit += gst_amt
                        gst_line.save()

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        # Return rich response for live UI updates
        remaining_pending = 0
        remaining_confirmed = 0
        if txn.job and txn.job.entity:
            from review.models import PendingTransaction as PT
            remaining_pending = PT.objects.filter(
                job__entity=txn.job.entity, is_confirmed=False
            ).count()
            remaining_confirmed = PT.objects.filter(
                job__entity=txn.job.entity, is_confirmed=True
            ).count()
        return JsonResponse({
            "status": "success",
            "id": str(txn.pk),
            "tb_updated": tb_updated,
            "has_gst": has_gst,
            "gst_amount": str(txn.confirmed_gst_amount or Decimal("0.00")),
            "net_amount": str(txn.net_amount or abs(txn.amount)),
            "pending_count": remaining_pending,
            "confirmed_count": remaining_confirmed,
            "confirmed_code": txn.confirmed_code,
            "confirmed_name": txn.confirmed_name,
            "confirmed_tax_type": txn.confirmed_tax_type,
            "amount": str(txn.amount),
            "date": txn.date,
            "description": txn.description,
        })

    messages.success(request, f"Transaction approved: {txn.description[:50]}")
    return redirect(request.META.get("HTTP_REFERER", "/"))


@login_required
def review_approve_all(request, pk):
    """Approve all pending transactions for a financial year's entity.
    Also auto-pushes all approved transactions to the trial balance.
    """
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    from review.models import PendingTransaction

    pending = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=False,
    )
    count = 0
    tb_count = 0
    for txn in pending:
        if txn.ai_suggested_code:
            txn.confirmed_code = txn.ai_suggested_code
            txn.confirmed_name = txn.ai_suggested_name
            txn.confirmed_tax_type = txn.ai_suggested_tax_type
            txn.is_confirmed = True
            txn.save()
            count += 1

            # Auto-push to trial balance
            amount = txn.amount
            code = txn.confirmed_code
            name = txn.confirmed_name
            if code and amount != 0:
                tb_line, created = TrialBalanceLine.objects.get_or_create(
                    financial_year=fy,
                    account_code=code,
                    defaults={
                        "account_name": name,
                        "debit": max(Decimal("0"), amount),
                        "credit": abs(min(Decimal("0"), amount)),
                    },
                )
                if not created:
                    if amount > 0:
                        tb_line.debit += amount
                    else:
                        tb_line.credit += abs(amount)
                    tb_line.save()
                tb_count += 1

    messages.success(request, f"Approved {count} transactions with AI suggestions. {tb_count} lines pushed to trial balance.")
    return redirect("core:financial_year_detail", pk=pk)


# ---------------------------------------------------------------------------
# Activity Log / Notifications
# ---------------------------------------------------------------------------
@login_required
@require_POST
@ratelimit(key='user', rate='30/m', method='POST')
def mark_notification_read(request, pk):
    """Mark a single activity log entry as read."""
    activity = get_object_or_404(ActivityLog, pk=pk, user=request.user)
    activity.is_read = True
    activity.save()
    return JsonResponse({"status": "ok"})


@login_required
@require_POST
@ratelimit(key='user', rate='10/m', method='POST')
def mark_all_notifications_read(request):
    """Mark all unread notifications as read for the current user."""
    ActivityLog.objects.filter(is_read=False, user=request.user).update(is_read=True)
    return JsonResponse({"status": "ok"})


@login_required
@ratelimit(key='user', rate='30/m', method='GET')
def notifications_api(request):
    """Return recent unread notifications as JSON for polling (scoped to current user)."""
    activities = (
        ActivityLog.objects.filter(is_read=False, user=request.user)
        .order_by("-created_at")[:10]
    )
    data = []
    for a in activities:
        data.append({
            "id": str(a.pk),
            "event_type": a.event_type,
            "title": a.title,
            "description": a.description,
            "url": a.url,
            "created_at": a.created_at.isoformat(),
        })
    return JsonResponse({"unread_count": ActivityLog.objects.filter(is_read=False, user=request.user).count(), "items": data})


# ============================================================
# Bulk Client Actions (Delete / Archive)
# ============================================================

@login_required
def client_bulk_action(request):
    """Bulk delete/archive actions on selected entities (top-level)."""
    if request.method != "POST":
        return redirect("core:entity_list")

    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_list")

    action = request.POST.get("bulk_action")
    entity_ids = request.POST.getlist("entity_ids")

    if not entity_ids:
        messages.warning(request, "No entities selected.")
        return redirect("core:entity_list")

    entities = Entity.objects.filter(pk__in=entity_ids)

    # Ownership check: non-senior users can only act on their own entities
    if not request.user.is_senior:
        entities = entities.filter(
            Q(assigned_accountant=request.user) |
            Q(client__assigned_accountant=request.user)
        )

    if action == "archive":
        count = entities.update(is_archived=True)
        messages.success(request, f"Archived {count} entity/entities.")
        _log_action(request, "update", f"Archived {count} entity/entities")
    elif action == "unarchive":
        count = entities.update(is_archived=False)
        messages.success(request, f"Unarchived {count} entity/entities.")
        _log_action(request, "update", f"Unarchived {count} entity/entities")
    elif action == "delete":
        count = entities.count()
        for e in entities:
            _log_action(request, "delete", f"Deleted entity: {e.entity_name}")
        entities.delete()
        messages.success(request, f"Deleted {count} entity/entities and all associated data.")
    else:
        messages.error(request, "Invalid action.")

    return redirect("core:entity_list")


# Alias for backward compatibility
entity_bulk_action = client_bulk_action


# ============================================================
# Entity-level HandiLedger Import
# ============================================================

@login_required
def entity_import_handiledger(request, pk):
    """Import HandiLedger ZIP for a specific entity."""
    from .access_ledger_import import import_access_ledger_zip

    entity = get_entity_for_user(request, pk)

    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=pk)

    result = None

    if request.method == "POST":
        zip_file = request.FILES.get("zip_file")
        if not zip_file:
            messages.error(request, "Please select a ZIP file.")
        elif not zip_file.name.lower().endswith(".zip"):
            messages.error(request, "File must be a .zip file.")
        else:
            replace = request.POST.get("replace_existing") == "1"
            try:
                result = import_access_ledger_zip(
                    zip_file,
                    client=entity.client if entity.client else None,
                    entity=entity,
                    replace_existing=replace,
                )
                if result["errors"]:
                    messages.warning(
                        request,
                        f"Import completed with {len(result['errors'])} error(s)."
                    )
                else:
                    messages.success(
                        request,
                        f"Successfully imported: "
                        f"{result['years_imported']} years, "
                        f"{result['total_tb_lines']} TB lines, "
                        f"{result['total_dep_assets']} depreciation assets."
                    )
                _log_action(
                    request, "import",
                    f"Imported HandiLedger ZIP for {entity.entity_name}: "
                    f"{result['years_imported']} years",
                    entity,
                )
            except Exception as e:
                messages.error(request, "Import failed. Please check the file format and try again.")

    return render(request, "core/entity_import_handiledger.html", {
        "entity": entity,
        "result": result,
    })


# ============================================================
# Delete Unfinalised FY Data
# ============================================================

@login_required
def delete_unfinalised_fy(request, pk):
    """Delete all unfinalised financial years and their data for an entity."""
    entity = get_entity_for_user(request, pk)

    if request.method != "POST":
        return redirect("core:entity_detail", pk=pk)

    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=pk)

    unfinalised = entity.financial_years.exclude(status="finalised")
    count = unfinalised.count()
    if count == 0:
        messages.info(request, "No unfinalised financial years to delete.")
        return redirect("core:entity_detail", pk=pk)

    for fy in unfinalised:
        _log_action(request, "delete", f"Deleted unfinalised FY: {fy.year_label} for {entity.entity_name}", fy)

    unfinalised.delete()
    messages.success(request, f"Deleted {count} unfinalised financial year(s) and all associated data.")
    return redirect("core:entity_detail", pk=pk)


# ============================================================
# HTMX: Update TB Line Mapping (clickable AI suggestion)
# ============================================================

@login_required
def htmx_update_tb_mapping(request, pk):
    """HTMX endpoint to update the mapping of a trial balance line via dropdown."""
    line = get_object_or_404(TrialBalanceLine, pk=pk)
    get_financial_year_for_user(request, line.financial_year.pk)  # IDOR check
    if not request.user.can_edit:
        return HttpResponse("Permission denied", status=403)

    if request.method == "POST":
        mapping_id = request.POST.get("mapped_line_item")
        if mapping_id:
            try:
                mapping = AccountMapping.objects.get(pk=mapping_id)
                line.mapped_line_item = mapping
                line.save()

                # Also save to client account mapping for reuse
                ClientAccountMapping.objects.update_or_create(
                    entity=line.financial_year.entity,
                    client_account_code=line.account_code,
                    defaults={
                        "client_account_name": line.account_name,
                        "mapped_line_item": mapping,
                    },
                )
            except AccountMapping.DoesNotExist:
                pass

    # Return the updated row
    entity_type = line.financial_year.entity.entity_type
    coa_items = ChartOfAccount.objects.filter(
        entity_type=entity_type, is_active=True
    ).select_related("maps_to").order_by("section", "account_code")

    return render(request, "partials/tb_line_row_review.html", {
        "line": line,
        "coa_items": coa_items,
    })


# ============================================================
# Chart of Accounts API for searchable dropdown
# ============================================================

@login_required
def coa_search_api(request):
    """JSON API for searching chart of accounts — used by the review tab dropdown."""
    entity_type = request.GET.get("entity_type", "company")
    q = request.GET.get("q", "")

    qs = ChartOfAccount.objects.filter(
        entity_type=entity_type, is_active=True
    ).select_related("maps_to").order_by("section", "account_code")

    if q:
        qs = qs.filter(
            Q(account_code__icontains=q) | Q(account_name__icontains=q)
        )

    items = []
    for a in qs[:200]:
        items.append({
            "id": str(a.maps_to.pk) if a.maps_to else "",
            "code": a.account_code,
            "name": a.account_name,
            "section": a.get_section_display(),
            "mapping_label": a.maps_to.line_item_label if a.maps_to else "Unmapped",
        })
    return JsonResponse({"items": items})


# ============================================================
# XRM Pull (Xero Practice Manager)
# ============================================================

@login_required
@ratelimit(key='user', rate='10/m', method='POST')
def xrm_search(request, pk):
    """
    Search XPM for clients matching this entity's name.
    Returns a list of potential matches for the user to confirm.
    """
    entity = get_entity_for_user(request, pk)

    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)

    if not request.user.can_edit:
        return JsonResponse({"status": "error", "message": "Permission denied"}, status=403)

    from integrations.models import XPMConnection
    from integrations.xpm_sync import _ensure_valid_token, _xpm_get, _xml_text

    connection = XPMConnection.objects.filter(status="active").first()
    if not connection:
        return JsonResponse({
            "status": "error",
            "message": "No active XPM connection. Go to Integrations > XPM to connect first."
        })

    if not _ensure_valid_token(connection):
        return JsonResponse({
            "status": "error",
            "message": "XPM token has expired. Please reconnect via Integrations > XPM."
        })

    try:
        # Search XPM by entity name
        search_name = entity.entity_name or entity.trading_name or ""
        root = _xpm_get(connection, "client.api/search", params={"query": search_name})

        matches = []
        # XPM returns <Clients><Client>...</Client></Clients> or <Response><Clients>...
        clients_el = root.find(".//Clients")
        if clients_el is None:
            clients_el = root

        for client_el in clients_el.findall("Client"):
            c_uuid = _xml_text(client_el, "UUID")
            c_name = _xml_text(client_el, "Name")
            c_email = _xml_text(client_el, "Email")
            c_phone = _xml_text(client_el, "Phone")
            c_structure = _xml_text(client_el, "BusinessStructure")
            c_abn = _xml_text(client_el, "BusinessNumber")
            if c_uuid and c_name:
                matches.append({
                    "uuid": c_uuid,
                    "name": c_name,
                    "email": c_email,
                    "phone": c_phone,
                    "structure": c_structure,
                    "abn": c_abn,
                })

        if not matches:
            return JsonResponse({
                "status": "no_results",
                "message": f"No clients found in XPM matching '{search_name}'.",
                "matches": [],
            })

        return JsonResponse({
            "status": "success",
            "message": f"Found {len(matches)} match(es) in XPM.",
            "matches": matches,
            "search_term": search_name,
        })

    except Exception as e:
        import traceback
        logger.error(f"XRM search failed: {e}\n{traceback.format_exc()}")
        return JsonResponse({
            "status": "error",
            "message": "XPM search failed. Please try again later."
        })


@login_required
@ratelimit(key='user', rate='5/m', method='POST')
def xrm_pull(request, pk):
    """Pull entity data from Xero Practice Manager for a single entity."""
    entity = get_entity_for_user(request, pk)

    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)

    if not request.user.can_edit:
        return JsonResponse({"status": "error", "message": "Permission denied"}, status=403)

    # Accept xpm_client_id from POST body (set during confirmation)
    import json as json_mod
    try:
        body = json_mod.loads(request.body)
    except (json_mod.JSONDecodeError, ValueError):
        body = {}

    xpm_id_from_body = body.get("xpm_client_id", "")
    if xpm_id_from_body:
        # User confirmed a match — save the XPM Client ID to the entity
        entity.xpm_client_id = xpm_id_from_body
        entity.save(update_fields=["xpm_client_id"])

    if not entity.xpm_client_id:
        return JsonResponse({
            "status": "error",
            "message": "This entity does not have an XPM Client ID. "
                       "Use the XRM button to search and link a client first."
        })

    # Get active XPM connection
    from integrations.models import XPMConnection
    from integrations.xpm_sync import (
        _ensure_valid_token, _xpm_get, _xml_text, _xml_name,
        STRUCTURE_MAP, RELATIONSHIP_MAP,
    )
    import xml.etree.ElementTree as ET

    connection = XPMConnection.objects.filter(status="active").first()
    if not connection:
        return JsonResponse({
            "status": "error",
            "message": "No active XPM connection. Go to Integrations > XPM to connect first."
        })

    # Ensure token is valid (refresh if needed)
    if not _ensure_valid_token(connection):
        return JsonResponse({
            "status": "error",
            "message": "XPM token has expired and could not be refreshed. "
                       "Please reconnect via Integrations > XPM."
        })

    def _normalize_name(name):
        """Normalize a name for comparison: handle 'SURNAME, First' vs 'First Surname',
        strip extra whitespace, and lowercase."""
        if not name:
            return ""
        name = name.strip()
        # Handle 'SURNAME, FirstName' format
        if "," in name:
            parts = [p.strip() for p in name.split(",", 1)]
            if len(parts) == 2:
                name = f"{parts[1]} {parts[0]}"
        # Collapse multiple spaces, lowercase
        return " ".join(name.split()).lower()

    def _officer_exists(entity, name):
        """Check if an officer with a matching name already exists."""
        norm = _normalize_name(name)
        if not norm:
            return True  # Skip empty names
        for officer in EntityOfficer.objects.filter(entity=entity):
            if _normalize_name(officer.full_name) == norm:
                return True
        return False

    def _associate_exists(entity, name):
        """Check if a relationship/associate with a matching name already exists."""
        norm = _normalize_name(name)
        if not norm:
            return None
        for assoc in ClientAssociate.objects.filter(entity=entity):
            if _normalize_name(assoc.name) == norm:
                return assoc
        return None

    try:
        # Fetch detailed client data from XPM
        root = _xpm_get(connection, f"client.api/get/{entity.xpm_client_id}")
        client_el = root.find(".//Client")
        if client_el is None:
            client_el = root

        updated_fields = []

        # --- Entity Info ---
        abn = _xml_text(client_el, "BusinessNumber")
        acn = _xml_text(client_el, "CompanyNumber")
        tax_number = _xml_text(client_el, "TaxNumber")
        email = _xml_text(client_el, "Email")
        phone = _xml_text(client_el, "Phone")
        structure = _xml_text(client_el, "BusinessStructure")
        gst_registered = _xml_text(client_el, "GSTRegistered", "").lower() == "yes"

        # Address — XPM uses Address element with sub-elements
        address_el = client_el.find("Address")
        addr_line1 = addr_line2 = city = region = postcode = country = ""
        if address_el is not None:
            addr_line1 = _xml_text(address_el, "Address") or _xml_text(address_el, "Line1")
            addr_line2 = _xml_text(address_el, "Line2")
            city = _xml_text(address_el, "City")
            region = _xml_text(address_el, "Region")
            postcode = _xml_text(address_el, "PostCode")
            country = _xml_text(address_el, "Country")

        # Update entity fields (only overwrite if XPM has data)
        if abn and abn != entity.abn:
            entity.abn = abn; updated_fields.append("ABN")
        if acn and acn != entity.acn:
            entity.acn = acn; updated_fields.append("ACN")
        if tax_number and tax_number != entity.tfn:
            entity.tfn = tax_number; updated_fields.append("TFN")
        if email and email != entity.contact_email:
            entity.contact_email = email; updated_fields.append("Email")
        if phone and phone != entity.contact_phone:
            entity.contact_phone = phone; updated_fields.append("Phone")
        if gst_registered != entity.is_gst_registered:
            entity.is_gst_registered = gst_registered; updated_fields.append("GST Status")
        if structure:
            entity_type = STRUCTURE_MAP.get(structure, "")
            if entity_type and entity_type != entity.entity_type:
                entity.entity_type = entity_type; updated_fields.append("Entity Type")
        if addr_line1 and addr_line1 != entity.address_line_1:
            entity.address_line_1 = addr_line1; updated_fields.append("Address Line 1")
        if addr_line2 and addr_line2 != entity.address_line_2:
            entity.address_line_2 = addr_line2; updated_fields.append("Address Line 2")
        if city and city != entity.suburb:
            entity.suburb = city; updated_fields.append("Suburb")
        if region and region != entity.state:
            entity.state = region; updated_fields.append("State")
        if postcode and postcode != entity.postcode:
            entity.postcode = postcode; updated_fields.append("Postcode")
        if country and country != entity.country:
            entity.country = country; updated_fields.append("Country")

        entity.save()

        # --- Sync Contacts as Relationships (ClientAssociate) ---
        contacts_synced = 0
        contacts_el = client_el.find("Contacts")
        if contacts_el is not None:
            for contact_el in contacts_el.findall("Contact"):
                c_uuid = _xml_text(contact_el, "UUID")
                c_name = _xml_text(contact_el, "Name")
                c_email = _xml_text(contact_el, "Email")
                c_phone = _xml_text(contact_el, "Phone")
                c_mobile = _xml_text(contact_el, "Mobile")
                c_position = _xml_text(contact_el, "Position")
                if not c_name:
                    continue

                # Determine relationship type from position
                pos_lower = (c_position or "").lower()
                if "spouse" in pos_lower or "wife" in pos_lower or "husband" in pos_lower:
                    rel_type = "spouse"
                elif "child" in pos_lower or "son" in pos_lower or "daughter" in pos_lower:
                    rel_type = "child"
                elif "director" in pos_lower:
                    rel_type = "director"
                elif "shareholder" in pos_lower:
                    rel_type = "shareholder"
                elif "trustee" in pos_lower:
                    rel_type = "trustee"
                elif "beneficiary" in pos_lower:
                    rel_type = "beneficiary"
                elif "partner" in pos_lower:
                    rel_type = "partner_biz"
                elif "secretary" in pos_lower:
                    rel_type = "trustee"
                else:
                    rel_type = "other"

                # Find or create associate (deduplicate by UUID, then normalized name)
                assoc = None
                if c_uuid:
                    assoc = ClientAssociate.objects.filter(
                        entity=entity, xpm_contact_uuid=c_uuid
                    ).first()
                if not assoc:
                    assoc = _associate_exists(entity, c_name)
                if assoc:
                    changed = False
                    if c_uuid and not assoc.xpm_contact_uuid:
                        assoc.xpm_contact_uuid = c_uuid; changed = True
                    if c_email and not assoc.email:
                        assoc.email = c_email; changed = True
                    if (c_phone or c_mobile) and not assoc.phone:
                        assoc.phone = c_phone or c_mobile; changed = True
                    if c_position and not assoc.occupation:
                        assoc.occupation = c_position; changed = True
                    if changed:
                        assoc.save()
                else:
                    ClientAssociate.objects.create(
                        entity=entity,
                        name=c_name,
                        relationship_type=rel_type,
                        email=c_email,
                        phone=c_phone or c_mobile,
                        occupation=c_position,
                        xpm_contact_uuid=c_uuid,
                    )
                contacts_synced += 1

        # --- Sync Relationships ---
        relationships_synced = 0
        relationships_el = client_el.find("Relationships")
        if relationships_el is not None:
            for rel_el in relationships_el.findall("Relationship"):
                rel_type_xpm = _xml_text(rel_el, "Type")
                related_name_el = rel_el.find("RelatedClient")
                related_name = ""
                if related_name_el is not None:
                    name_el = related_name_el.find("Name")
                    if name_el is not None and name_el.text:
                        related_name = name_el.text.strip()
                if not related_name:
                    continue

                rel_type = RELATIONSHIP_MAP.get(rel_type_xpm, "related_entity")
                assoc = _associate_exists(entity, related_name)
                if assoc:
                    if rel_type and assoc.relationship_type == "other":
                        assoc.relationship_type = rel_type
                        assoc.save()
                else:
                    ClientAssociate.objects.create(
                        entity=entity,
                        name=related_name,
                        relationship_type=rel_type,
                    )
                relationships_synced += 1

        # --- Sync Officers from Contacts with officer-like positions ---
        officers_synced = 0
        if contacts_el is not None:
            for contact_el in contacts_el.findall("Contact"):
                c_name = _xml_text(contact_el, "Name")
                c_position = _xml_text(contact_el, "Position")
                if not c_name or not c_position:
                    continue
                pos_lower = c_position.lower()
                officer_role = None
                if "director" in pos_lower:
                    officer_role = "director"
                elif "partner" in pos_lower:
                    officer_role = "partner"
                elif "trustee" in pos_lower:
                    officer_role = "trustee"
                elif "secretary" in pos_lower:
                    officer_role = "secretary"
                elif "public officer" in pos_lower:
                    officer_role = "public_officer"
                if officer_role:
                    if not _officer_exists(entity, c_name):
                        EntityOfficer.objects.create(
                            entity=entity,
                            full_name=c_name,
                            role=officer_role,
                            title=c_position,
                        )
                        officers_synced += 1

        # Build summary
        summary_parts = []
        if updated_fields:
            summary_parts.append(f"Updated: {', '.join(updated_fields)}")
        if contacts_synced:
            summary_parts.append(f"{contacts_synced} contacts synced")
        if relationships_synced:
            summary_parts.append(f"{relationships_synced} relationships synced")
        if officers_synced:
            summary_parts.append(f"{officers_synced} officers added")
        if not summary_parts:
            summary_parts.append("Entity is already up to date")

        _log_action(
            request, "xrm_pull",
            f"XRM pull completed for {entity.entity_name}: "
            f"{len(updated_fields)} fields, {contacts_synced} contacts, "
            f"{relationships_synced} relationships, {officers_synced} officers",
            entity,
        )

        return JsonResponse({
            "status": "success",
            "message": f"XRM pull completed for {entity.entity_name}. "
                       + ". ".join(summary_parts) + "."
        })

    except Exception as e:
        import traceback
        logger.error(f"XRM pull failed for {entity.entity_name}: {e}\n{traceback.format_exc()}")
        return JsonResponse({
            "status": "error",
            "message": "XRM pull failed. Please try again later."
        })


@login_required
def trial_balance_download(request, pk):
    """
    Download comparative trial balance as Word (.docx) or PDF.
    Uses ?format=docx or ?format=pdf query parameter.
    """
    from io import BytesIO
    from collections import OrderedDict
    from decimal import Decimal

    fmt = request.GET.get("format", "pdf").lower()
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    tb_lines = TrialBalanceLine.objects.filter(
        financial_year=fy
    ).select_related('mapped_line_item').order_by('account_code')

    # Section ordering
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }

    sections = OrderedDict()
    grand_dr = Decimal('0')
    grand_cr = Decimal('0')
    grand_prior_dr = Decimal('0')
    grand_prior_cr = Decimal('0')

    for line in tb_lines:
        if line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = 'Unmapped'
        sections.setdefault(display_section, []).append(line)
        grand_dr += line.debit or Decimal('0')
        grand_cr += line.credit or Decimal('0')
        grand_prior_dr += line.prior_debit or Decimal('0')
        grand_prior_cr += line.prior_credit or Decimal('0')

    ordered = OrderedDict()
    seen = set()
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in seen and ds in sections:
            ordered[ds] = sections[ds]
            seen.add(ds)
    for key in sections:
        if key not in ordered:
            ordered[key] = sections[key]

    # Year labels
    current_year = str(fy.year_label)
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year = str(fy.prior_year.year_label)
    else:
        prior_year = 'Prior'

    # ABN formatting
    abn = entity.abn or ''
    abn_display = ''
    if abn:
        abn_clean = abn.replace(' ', '')
        if len(abn_clean) == 11:
            abn_display = f'ABN {abn_clean[:2]} {abn_clean[2:5]} {abn_clean[5:8]} {abn_clean[8:11]}'
        else:
            abn_display = f'ABN {abn}'

    safe_name = entity.entity_name.replace(' ', '_')

    if fmt == "docx":
        return _tb_download_word(
            fy, entity, ordered, current_year, prior_year,
            abn_display, grand_dr, grand_cr, grand_prior_dr, grand_prior_cr, safe_name
        )
    else:
        # Reuse existing trial_balance_pdf
        return trial_balance_pdf(request, pk)


def _tb_download_word(fy, entity, sections, current_year, prior_year,
                      abn_display, grand_dr, grand_cr, grand_prior_dr, grand_prior_cr, safe_name):
    """Generate a Word document for the comparative trial balance."""
    from io import BytesIO
    from decimal import Decimal
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    import os
    from django.conf import settings

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Logo
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'mcs_logo.png')
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_path, width=Cm(2.5))

    # Entity name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(entity.entity_name.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    # ABN
    if abn_display:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(abn_display)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Comparative Trial Balance as at {fy.end_date.strftime('%d %B %Y')}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(12)

    def fmt_val(val):
        if val and val != Decimal('0'):
            return f"{val:,.2f}"
        return ''

    # Column header table
    header_table = doc.add_table(rows=2, cols=6)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True

    # Header row 1
    cells = header_table.rows[0].cells
    cells[0].text = ''
    cells[1].text = ''
    cells[2].text = current_year
    cells[3].text = current_year
    cells[4].text = prior_year
    cells[5].text = prior_year

    # Header row 2
    cells = header_table.rows[1].cells
    cells[0].text = ''
    cells[1].text = ''
    cells[2].text = '$ Dr'
    cells[3].text = '$ Cr'
    cells[4].text = '$ Dr'
    cells[5].text = '$ Cr'

    # Style header rows
    for row_idx in range(2):
        for col_idx in range(6):
            cell = header_table.rows[row_idx].cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                if col_idx >= 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Section tables
    for section_name, lines in sections.items():
        # Section heading
        p = doc.add_paragraph()
        run = p.add_run(section_name)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.space_before = Pt(8)
        p.space_after = Pt(4)

        table = doc.add_table(rows=len(lines), cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, line in enumerate(lines):
            dr = line.debit or Decimal('0')
            cr = line.credit or Decimal('0')
            prior_dr = line.prior_debit or Decimal('0')
            prior_cr = line.prior_credit or Decimal('0')

            cells = table.rows[i].cells
            cells[0].text = line.account_code
            cells[1].text = line.account_name
            cells[2].text = fmt_val(dr)
            cells[3].text = fmt_val(cr)
            cells[4].text = fmt_val(prior_dr)
            cells[5].text = fmt_val(prior_cr)

            for col_idx in range(6):
                for paragraph in cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        if col_idx == 0:
                            run.bold = True
                    if col_idx >= 2:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Grand totals
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    totals_table = doc.add_table(rows=1, cols=6)
    totals_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = totals_table.rows[0].cells
    cells[0].text = ''
    cells[1].text = 'TOTALS'
    cells[2].text = f"{grand_dr:,.2f}"
    cells[3].text = f"{grand_cr:,.2f}"
    cells[4].text = f"{grand_prior_dr:,.2f}"
    cells[5].text = f"{grand_prior_cr:,.2f}"
    for col_idx in range(6):
        for paragraph in cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            if col_idx >= 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Net profit
    net_profit_current = grand_cr - grand_dr
    net_profit_prior = grand_prior_cr - grand_prior_dr
    profit_table = doc.add_table(rows=1, cols=6)
    profit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = profit_table.rows[0].cells
    cells[0].text = ''
    cells[1].text = 'Net Profit'
    cells[2].text = ''
    cells[3].text = f"{abs(net_profit_current):,.2f}"
    cells[4].text = ''
    cells[5].text = f"{abs(net_profit_prior):,.2f}"
    for col_idx in range(6):
        for paragraph in cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            if col_idx >= 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Comparative_TB_{safe_name}_{fy.year_label}.docx"
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
@require_POST
def delete_document(request, pk):
    """Delete a generated document."""
    from .models import GeneratedDocument
    doc = get_object_or_404(GeneratedDocument, pk=pk)
    fy_pk = doc.financial_year.pk
    # Delete the file from storage
    if doc.file:
        doc.file.delete(save=False)
    doc.delete()
    messages.success(request, "Document deleted.")
    return redirect("core:financial_year_detail", pk=fy_pk)
